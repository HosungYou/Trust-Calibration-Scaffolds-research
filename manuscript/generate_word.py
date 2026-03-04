#!/usr/bin/env python3
"""
Generate APA 7th Edition Word document for:
"Trust Calibration as the Missing Link in Educational AI Design"

Includes 5 figures, 2 tables, proper APA formatting.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SECTIONS_DIR = os.path.join(BASE_DIR, "sections")
FIGURES_DIR = os.path.join(BASE_DIR, "figures")
OUTPUT_PATH = os.path.join(BASE_DIR, "Trust_Calibration_Educational_AI_APA7.docx")

# ─── Helpers ───

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_styled_run(paragraph, text, bold=False, italic=False, size=12, font_name="Times New Roman"):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def add_body_paragraph(doc, text, first_line_indent=True):
    """Add an APA-style body paragraph with double spacing."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)

    # Parse markdown-like formatting
    process_inline_formatting(p, text)
    return p


def process_inline_formatting(paragraph, text):
    """Parse bold (**), italic (*), and em-dash (---) formatting."""
    # Replace --- with em dash
    text = text.replace("---", "\u2014")
    # Replace -- with en dash
    text = text.replace("--", "\u2013")

    # Pattern to find **bold** and *italic* markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)

    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            add_styled_run(paragraph, part[3:-3], bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            add_styled_run(paragraph, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*'):
            add_styled_run(paragraph, part[1:-1], italic=True)
        else:
            add_styled_run(paragraph, part)


def add_heading(doc, text, level=1):
    """Add an APA-style heading."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=2.0)

    if level == 1:
        # Level 1: Centered, Bold
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_run(p, text, bold=True)
    elif level == 2:
        # Level 2: Left-aligned, Bold
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_styled_run(p, text, bold=True)
    elif level == 3:
        # Level 3: Left-aligned, Bold Italic
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_styled_run(p, text, bold=True, italic=True)
    return p


def add_figure(doc, image_path, figure_num, caption_text):
    """Add a figure with APA 7th formatting: Figure N label, then caption, then image."""
    # Figure label
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_label, before=12, after=0, line_spacing=2.0)
    add_styled_run(p_label, f"Figure {figure_num}", bold=True, italic=True)

    # Caption
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_caption, before=0, after=6, line_spacing=2.0)
    process_inline_formatting(p_caption, caption_text)

    # Image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_img, before=0, after=12, line_spacing=1.0)
    run = p_img.add_run()
    run.add_picture(image_path, width=Inches(6.0))

    return p_img


def add_table_from_data(doc, table_num, title, headers, rows, col_widths=None):
    """Add an APA-style table."""
    # Table label
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_label, before=12, after=0, line_spacing=2.0)
    add_styled_run(p_label, f"Table {table_num}", bold=True, italic=True)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_title, before=0, after=6, line_spacing=2.0)
    process_inline_formatting(p_title, title)

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_styled_run(p, header, bold=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Top and bottom border for header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        tcPr.append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            process_inline_formatting(p, str(cell_text))
            for run in p.runs:
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths if specified
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Spacing after table
    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, before=6, after=6, line_spacing=2.0)

    return table


# ─── Title Page ───

def create_title_page(doc):
    # Running head (simplified - just add title info)
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Trust Calibration as the Missing Link in Educational AI Design:", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "A Critical Review and Two-Level Framework", bold=True)

    # Blank line
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Hosung You")

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Department of Artificial Intelligence Convergence Education")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Sungkyunkwan University")

    # Author note
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Author Note", bold=True)

    add_body_paragraph(doc, "Correspondence concerning this article should be addressed to Hosung You, Department of Artificial Intelligence Convergence Education, Sungkyunkwan University, Seoul, South Korea.")

    # Page break
    doc.add_page_break()


# ─── Abstract Page ───

def create_abstract_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=2.0)
    add_styled_run(p, "Abstract", bold=True)

    abstract_text = (
        "The rapid integration of generative AI into education raises a critical question: "
        "do learners trust AI systems appropriately? Bastani et al. (2025) demonstrated that "
        "unrestricted AI tutoring access produced 17% lower performance on unassisted assessments, "
        "attributable to uncalibrated overtrust and cognitive offloading. Yet existing frameworks "
        "in educational AI address trust formation and technology acceptance without addressing "
        "trust calibration\u2014the alignment between a learner\u2019s trust and the AI system\u2019s actual "
        "reliability. This paper presents a critical review of trust in educational AI research "
        "(2015\u20132026), analyzing 97 studies from Web of Science, Scopus, and supplementary "
        "HCI/human factors literature across trust conceptualizations, theoretical frameworks, "
        "measurement approaches, and oversight design. The review reveals a pronounced calibration "
        "gap: 84.5% of studies provided no explicit trust definition, only 7.2% measured trust "
        "empirically, and only 1 of 97 studies measured trust calibration accuracy. Nearly half "
        "(49.5%) employed no identifiable theoretical framework, and only 16.5% explicitly "
        "addressed calibration\u2014with 13 of those 16 papers originating from adjacent disciplines "
        "rather than educational AI databases. To address this gap, we propose an evidence-informed "
        "Two-Level Trust Calibration Framework integrating: (a) Level 1\u2014Learner Trust Dynamics, "
        "modeling how trust forms and miscalibrates through cognitive and metacognitive processes; "
        "(b) Level 2\u2014System-Mediated Trust Calibration, specifying six scaffolds grounded in "
        "established educational theory; connected through (c) an Adaptive Calibration Cycle in "
        "which scaffold intensity adjusts based on behavioral signals of trust calibration. The "
        "paper distinguishes trust calibration from self-regulated learning calibration, introduces "
        "the Trust\u2013AI Reliability Matrix for diagnostic assessment, and derives five design "
        "principles for trust-calibrated educational AI."
    )

    add_body_paragraph(doc, abstract_text, first_line_indent=False)

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Keywords: ", italic=True)
    add_styled_run(p, "trust calibration, educational AI, scaffolding, self-regulated learning, human oversight, generative AI")

    doc.add_page_break()


# ─── Section Processors ───

def read_section(filename):
    path = os.path.join(SECTIONS_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def process_section_text(doc, text, figure_map=None, table_insertions=None, figure_insertions=None):
    """Process markdown text into Word paragraphs with figure/table insertion points."""
    lines = text.strip().split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Handle headings
        if line.startswith('### '):
            heading_text = line[4:].strip()
            # Remove numbering like "5.1 "
            heading_text = re.sub(r'^\d+\.\d+\s+', '', heading_text)
            add_heading(doc, heading_text, level=2)
            i += 1
            continue

        if line.startswith('## '):
            heading_text = line[3:].strip()
            # Remove numbering like "5. "
            heading_text = re.sub(r'^\d+\.\s+', '', heading_text)
            add_heading(doc, heading_text, level=1)
            i += 1
            continue

        # Handle markdown tables
        if '|' in line and line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # Skip markdown tables (we insert our own formatted tables)
            continue

        # Handle bullet points / numbered lists
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line):
            bullet_text = re.sub(r'^[-*]\s+|^\d+\.\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
            # Add bullet
            add_styled_run(p, "\u2022 ")
            process_inline_formatting(p, bullet_text)
            i += 1
            continue

        # Check for figure insertion triggers
        if figure_insertions:
            for trigger, fig_info in list(figure_insertions.items()):
                if trigger in line:
                    # First add the paragraph containing the trigger
                    p = add_body_paragraph(doc, apply_figure_renumbering(line, figure_map))
                    # Then add the figure
                    fig_path = os.path.join(FIGURES_DIR, fig_info['file'])
                    if os.path.exists(fig_path):
                        add_figure(doc, fig_path, fig_info['num'], fig_info['caption'])
                    del figure_insertions[trigger]
                    i += 1
                    line = None
                    break
            if line is None:
                continue

        # Regular paragraph - apply figure renumbering
        processed_line = apply_figure_renumbering(line, figure_map) if figure_map else line
        add_body_paragraph(doc, processed_line)
        i += 1


def apply_figure_renumbering(text, figure_map):
    """Rename Figure references: Figure 2 -> Figure 4, Figure 3 -> Figure 5."""
    if not figure_map:
        return text
    for old, new in figure_map.items():
        text = text.replace(old, new)
    return text


# ─── Main Document Builder ───

def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Figure renumbering map
    figure_map = {
        "**Figure 2**": "**Figure 4**",
        "**Figure 3**": "**Figure 5**",
        "(Figure 3)": "(Figure 5)",
        "(Figure 2)": "(Figure 4)",
        "Figure 3)": "Figure 5)",
        "Figure 2 ": "Figure 4 ",
        "Figure 3 ": "Figure 5 ",
    }

    # ─── Title Page ───
    create_title_page(doc)

    # ─── Abstract ───
    create_abstract_page(doc)

    # ─── Section 1: Introduction ───
    text = read_section("01_introduction.md")
    process_section_text(doc, text, figure_map)

    # ─── Section 2: Conceptual Foundation ───
    text = read_section("02_conceptual_foundation.md")
    process_section_text(doc, text, figure_map)

    # ─── Section 3: Review Approach ───
    text = read_section("03_review_approach.md")
    process_section_text(doc, text, figure_map)

    # ─── Figure 1: PRISMA Flow ───
    fig1_path = os.path.join(FIGURES_DIR, "figure1_prisma_flow.png")
    if os.path.exists(fig1_path):
        add_figure(doc, fig1_path, 1,
                   "PRISMA-style flow diagram of the literature search and selection process. "
                   "Searches across Web of Science, Scopus, and PsycINFO yielded 147 papers, "
                   "of which 97 were coded using the 32-field coding scheme.")

    # ─── Section 4: Literature Landscape ───
    text = read_section("04_literature_landscape.md")

    # Split section 4 into subsections for figure insertion
    subsections = text.split("### 4.2")

    # 4.1 Trust Conceptualizations
    process_section_text(doc, subsections[0], figure_map)

    # ─── Figure 2: Literature Landscape ───
    fig2_path = os.path.join(FIGURES_DIR, "figure2_literature_landscape.png")
    if os.path.exists(fig2_path):
        add_figure(doc, fig2_path, 2,
                   "Literature landscape overview of 97 coded studies on trust in educational AI (2015\u20132026). "
                   "Panel (a) shows trust conceptualization distribution, (b) theoretical framework usage, "
                   "(c) trust measurement approaches, and (d) calibration coverage.")

    # 4.2 onward
    remaining_4 = "### 4.2" + subsections[1]

    # Split at 4.4 for calibration gap figure
    parts_4 = remaining_4.split("### 4.4")
    process_section_text(doc, parts_4[0], figure_map)

    # 4.4 The Calibration Gap
    if len(parts_4) > 1:
        process_section_text(doc, "### 4.4" + parts_4[1], figure_map)

    # ─── Figure 3: Calibration Gap ───
    fig3_path = os.path.join(FIGURES_DIR, "figure3_calibration_gap.png")
    if os.path.exists(fig3_path):
        add_figure(doc, fig3_path, 3,
                   "The calibration gap in educational AI trust research. "
                   "Only 16.5% of studies explicitly addressed calibration, and 13 of those 16 "
                   "originated from adjacent disciplines. Only 1 of 97 studies measured calibration accuracy.")

    # ─── Table 1: Literature Landscape Summary ───
    table1_headers = ["Dimension", "Category", "Count (N=97)", "Percentage"]
    table1_rows = [
        ["Trust Conceptualization", "Trust in AI (general)", "86", "88.7%"],
        ["", "Trust in automation", "6", "6.2%"],
        ["", "Institutional trust", "3", "3.1%"],
        ["", "Interpersonal trust", "2", "2.1%"],
        ["Trust Definition", "No definition", "82", "84.5%"],
        ["", "Implicit definition", "6", "6.2%"],
        ["", "Explicit definition", "9", "9.3%"],
        ["Theoretical Framework", "None", "48", "49.5%"],
        ["", "TAM/UTAUT", "22", "22.7%"],
        ["", "AIED frameworks", "13", "13.4%"],
        ["", "DOI theory", "13", "13.4%"],
        ["", "Trust-in-automation", "9", "9.3%"],
        ["", "SRL frameworks", "5", "5.2%"],
        ["", "SDT", "3", "3.1%"],
        ["Trust Measurement", "Not measured", "90", "92.8%"],
        ["", "Measured empirically", "7", "7.2%"],
        ["Calibration Coverage", "Explicitly addressed", "16", "16.5%"],
        ["", "Partially addressed", "42", "43.3%"],
        ["", "Not addressed", "39", "40.2%"],
        ["Calibration Accuracy", "Measured", "1", "1.0%"],
        ["", "Not measured", "96", "99.0%"],
    ]
    add_table_from_data(doc, 1,
                       "Summary of Literature Landscape Findings (N = 97)",
                       table1_headers, table1_rows,
                       col_widths=[1.8, 1.8, 1.2, 1.0])

    # ─── Section 5: Two-Level Framework ───
    text = read_section("05_two_level_framework.md")

    # Split at figure insertion points
    # Find the line with "**Figure 2** presents"
    parts_5 = text.split("**Figure 2** presents the Two-Level Framework visually")

    # Everything before Figure 2 reference
    process_section_text(doc, parts_5[0], figure_map)

    if len(parts_5) > 1:
        # Add the Figure 2 reference paragraph (now Figure 4)
        fig2_ref_text = "**Figure 4** presents the Two-Level Framework visually" + parts_5[1].split('\n')[0]
        add_body_paragraph(doc, "")  # spacer
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
        p.paragraph_format.first_line_indent = Inches(0.5)
        process_inline_formatting(p, fig2_ref_text)

        # ─── Figure 4: Two-Level Framework ───
        fig4_path = os.path.join(FIGURES_DIR, "figure4_two_level_framework.png")
        if os.path.exists(fig4_path):
            add_figure(doc, fig4_path, 4,
                       "The Two-Level Trust Calibration Framework. Level 2 (upper) specifies six Trust Calibration "
                       "Scaffolds and system-level oversight design. Level 1 (lower) models learner trust dynamics "
                       "through AI system interaction, internal calibration, and trust manifestation. The Adaptive "
                       "Calibration Cycle (center) connects the two levels through four stages: Performance, "
                       "Monitoring, Evaluation, and Adaptation. Bidirectional arrows indicate the sensing pathway "
                       "(behavioral data flowing upward) and the adjusting pathway (scaffold modifications flowing downward).")

        # Continue with text after Figure 2 reference line
        remaining_after_fig2 = '\n'.join(parts_5[1].split('\n')[1:])

        # Split at Figure 3 reference
        parts_fig3 = remaining_after_fig2.split("**Figure 3** introduces the Trust-Reliability")

        if len(parts_fig3) > 1:
            # Text between Figure 2 and Figure 3 references (empty or minimal)
            if parts_fig3[0].strip():
                process_section_text(doc, parts_fig3[0], figure_map)

            # Figure 3 reference paragraph (now Figure 5)
            fig3_ref_line = parts_fig3[1].split('\n')[0]
            fig3_ref_text = "**Figure 5** introduces the Trust-Reliability" + fig3_ref_line
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
            p.paragraph_format.first_line_indent = Inches(0.5)
            process_inline_formatting(p, fig3_ref_text)

            # ─── Figure 5: Trust-AI Reliability Matrix ───
            fig5_path = os.path.join(FIGURES_DIR, "figure5_trust_reliability_matrix.png")
            if os.path.exists(fig5_path):
                add_figure(doc, fig5_path, 5,
                           "The Trust\u2013AI Reliability Calibration Matrix. Four quadrants map learner trust "
                           "(high/low) against AI system reliability (high/low), producing four calibration "
                           "states: Calibrated Appropriate Use (I), Overtrust/\"Bastani Trap\" (II), "
                           "Undertrust/Missed Benefit (III), and Calibrated Avoidance (IV). Each quadrant "
                           "specifies the trust calibration status and recommended scaffold response.")

            # Remaining text after Figure 3 reference
            remaining_text = '\n'.join(parts_fig3[1].split('\n')[1:])
            if remaining_text.strip():
                process_section_text(doc, remaining_text, figure_map)
        else:
            # No Figure 3 split found - process remaining
            process_section_text(doc, remaining_after_fig2, figure_map)

    # ─── Table 2: Evidence Mapping ───
    table2_headers = ["Framework Component", "Primary Evidence Source", "Evidence Type", "Strength Assessment"]
    table2_rows = [
        ["Metacognitive Prompts", "Guo et al. (2022); Bannert et al. (2009)", "Direct (SRL, technology-enhanced learning)", "Strong (g = 0.50 for SRL outcomes)"],
        ["Desirable Difficulties", "Bjork & Bjork (2011)", "Adjacent (cognitive psychology)", "Strong (robust across learning domains)"],
        ["Productive Failure", "Kapur (2008, 2016)", "Adjacent (mathematics education)", "Strong (substantial advantages over conventional instruction)"],
        ["AI Uncertainty Transparency", "XAI literature; Lee & See (2004)", "Adjacent (HCI, human factors)", "Moderate (demonstrated in automation, limited in educational AI)"],
        ["Socratic Dialogue", "Chi et al. (2001)", "Adjacent (tutoring, dialogue-based learning)", "Moderate (evidence for questioning over telling, limited in AI contexts)"],
        ["Progressive Autonomy Release", "Wood et al. (1976); Pea (2004)", "Direct (scaffolding theory)", "Strong (foundational principle; d = 0.82, Hattie, 2009)"],
        ["Adaptive Calibration Cycle", "Zimmerman (2000); Winne & Hadwin (1998)", "Structural analogy (SRL theory)", "Moderate (SRL cycle established; system-level application is novel)"],
        ["Trust Calibration in Ed. AI", "Sparse", "Gap identified", "Weak (primary gap motivating this framework)"],
        ["Two-Level Integration", "Novel contribution", "Conceptual synthesis", "Framework-level (no direct precedent)"],
    ]
    add_table_from_data(doc, 2,
                       "Evidence Mapping for Two-Level Trust Calibration Framework Components",
                       table2_headers, table2_rows,
                       col_widths=[1.5, 1.8, 1.5, 1.8])

    # ─── Section 6: Implications ───
    text = read_section("06_implications.md")
    # Apply figure renumbering
    process_section_text(doc, text, figure_map)

    # ─── Section 7: Conclusion ───
    text = read_section("07_conclusion.md")
    process_section_text(doc, text, figure_map)

    # ─── References ───
    doc.add_page_break()
    add_heading(doc, "References", level=1)

    references = get_references()
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
        process_inline_formatting(p, ref)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

    # Report stats
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    img_count = sum(1 for p in doc.paragraphs for r in p.runs
                    if r._element.findall(qn('w:drawing')))
    print(f"Paragraphs: {para_count}, Tables: {table_count}, Images: {img_count}")


def get_references():
    """Return APA 7th formatted references."""
    return [
        "Al-Adwan, A. S., Li, N., Al-Adwan, A., Abbasi, G. A., Albelbisi, N. A., & Habibi, A. (2023). Extending the technology acceptance model (TAM) to predict university students' intentions to use metaverse-based learning platforms. *Education and Information Technologies*, *28*(9), 11823\u201311850.",
        "Anderson, T., & Shattuck, J. (2012). Design-based research: A decade of progress in education research? *Educational Researcher*, *41*(1), 16\u201325.",
        "Arrieta, A. B., D\u00edaz-Rodr\u00edguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., ... & Herrera, F. (2020). Explainable artificial intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. *Information Fusion*, *58*, 82\u2013115.",
        "Azevedo, R. (2005). Using hypermedia as a metacognitive tool for enhancing student learning? The role of self-regulated learning. *Educational Psychologist*, *40*(4), 199\u2013209.",
        "Baidoo-Anu, D., & Ansah, L. O. (2023). Education in the era of generative artificial intelligence (AI): Understanding the potential benefits of ChatGPT in promoting teaching and learning. *Journal of AI*, *7*(1), 52\u201362.",
        "Bannert, M., Sonnenberg, C., Mengelkamp, C., & Pieger, E. (2009). Effects of a metacognitive support device in learning environments. *Computers in Human Behavior*, *31*, 585\u2013596.",
        "Bastani, H., Bastani, O., Sungu, A., Ge, H., Kabakcı, \u00d6., & Mariman, R. (2025). Generative AI can harm learning. *Proceedings of the National Academy of Sciences*, *122*(2), e2412467122.",
        "Belland, B. R. (2014). Scaffolding: Definition, current debates, and future directions. In J. M. Spector et al. (Eds.), *Handbook of research on educational communications and technology* (pp. 505\u2013518). Springer.",
        "Bjork, E. L., & Bjork, R. A. (2011). Making things hard on yourself, but in a good way: Creating desirable difficulties to enhance learning. In M. A. Gernsbacher et al. (Eds.), *Psychology and the real world* (pp. 56\u201364). Worth Publishers.",
        "Chi, M. T. H., Siler, S. A., Jeong, H., Yamauchi, T., & Hausmann, R. G. (2001). Learning from human tutoring. *Cognitive Science*, *25*(4), 471\u2013533.",
        "Cotton, D. R. E., Cotton, P. A., & Shipway, J. R. (2024). Chatting and cheating: Ensuring academic integrity in the era of ChatGPT. *Innovations in Education and Teaching International*, *61*(2), 228\u2013239.",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. *MIS Quarterly*, *13*(3), 319\u2013340.",
        "de Visser, E. J., Peeters, M. M. M., Jung, M. F., Kohn, S., Shaw, T. H., Pak, R., & Neerincx, M. A. (2020). Towards a theory of longitudinal trust calibration in human\u2013robot teams. *International Journal of Social Robotics*, *12*, 459\u2013478.",
        "Dunlosky, J., & Rawson, K. A. (2012). Overconfidence produces underachievement: Inaccurate self evaluations undermine students' learning and retention. *Learning and Instruction*, *22*(4), 271\u2013280.",
        "Grant, M. J., & Booth, A. (2009). A typology of reviews: An analysis of 14 review types and associated methodologies. *Health Information and Libraries Journal*, *26*(2), 91\u2013108.",
        "Greene, J. A., & Azevedo, R. (2007). A theoretical review of Winne and Hadwin's model of self-regulated learning: New perspectives and directions. *Review of Educational Research*, *77*(3), 334\u2013372.",
        "Guo, L., Wang, W., Yin, Z., & Zheng, Q. (2022). The effects of metacognitive scaffolding on students' self-regulated learning in technology-enhanced environments: A meta-analysis. *Education and Information Technologies*, *27*(5), 6619\u20136637.",
        "Hannafin, M., Land, S., & Oliver, K. (1999). Open learning environments: Foundations, methods, and models. In C. Reigeluth (Ed.), *Instructional-design theories and models* (Vol. 2, pp. 115\u2013140). Lawrence Erlbaum.",
        "Hattie, J. (2009). *Visible learning: A synthesis of over 800 meta-analyses relating to achievement*. Routledge.",
        "Hoff, K. A., & Bashir, M. (2015). Trust in automation: Integrating empirical evidence on factors that influence trust. *Human Factors*, *57*(3), 407\u2013434.",
        "Hofstede, G. (2001). *Culture's consequences: Comparing values, behaviors, institutions, and organizations across nations* (2nd ed.). Sage.",
        "Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., ... & Fung, P. (2023). Survey of hallucination in natural language generation. *ACM Computing Surveys*, *55*(12), 1\u201338.",
        "Kapur, M. (2008). Productive failure. *Cognition and Instruction*, *26*(3), 379\u2013424.",
        "Kapur, M. (2016). Examining productive failure, productive success, and constructive failure as design principles for learning. *Instructional Science*, *44*(3), 287\u2013306.",
        "Kasneci, E., Se\u00dfler, K., K\u00fcchemann, S., Bannert, M., Dementieva, D., Fischer, F., ... & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. *Learning and Individual Differences*, *103*, 102274.",
        "Lankton, N. K., McKnight, D. H., & Tripp, J. (2015). Technology, humanness, and trust: Rethinking trust in technology. *Journal of the Association for Information Systems*, *16*(10), 880\u2013918.",
        "Lee, J. D., & See, K. A. (2004). Trust in automation: Designing for appropriate reliance. *Human Factors*, *46*(1), 50\u201380.",
        "Lee, S., & Bosch, N. (2025). AI reliance and metacognitive calibration: A study on ChatGPT-assisted decision making. *Proceedings of the ACM Conference on Learning at Scale*, Article 12.",
        "Lee, S., Chung, S., & Bosch, N. (2025). Metacognitive sensitivity as a trust calibration mechanism in human-AI collaboration. *International Journal of Human-Computer Studies*, *185*, 103258.",
        "Li, X., Hess, T. J., & Valacich, J. S. (2008). Why do we trust new technology? A study of initial trust formation with organizational information systems. *Journal of Strategic Information Systems*, *17*(1), 39\u201371.",
        "Li, Y., Wang, C., & Zhang, P. (2024). Effects of AI confidence calibration on user trust and reliance in AI-assisted decision making. *International Journal of Human-Computer Studies*, *182*, 103167.",
        "Long, D., & Magerko, B. (2020). What is AI literacy? Competencies and design considerations. In *Proceedings of the 2020 CHI Conference on Human Factors in Computing Systems* (pp. 1\u201316). ACM.",
        "Ma, S., Lei, Y., Wang, X., Zheng, C., Shi, C., Yin, M., & Ma, X. (2024). Are you sure? Calibrating trust in AI assistants through self-confidence prompts. In *Proceedings of the 2024 CHI Conference on Human Factors in Computing Systems*. ACM.",
        "Mayer, R. C., Davis, J. H., & Schoorman, F. D. (1995). An integrative model of organizational trust. *Academy of Management Review*, *20*(3), 709\u2013734.",
        "McKenney, S., & Reeves, T. C. (2019). *Conducting educational design research* (2nd ed.). Routledge.",
        "Mehrotra, S., Jonker, C. M., & Tielman, M. L. (2024). More than meets the eye: A systematic literature review on appropriate trust in human-AI interaction. *International Journal of Human-Computer Studies*, *186*, 103252.",
        "Nelson, T. O., & Narens, L. (1990). Metamemory: A theoretical framework and new findings. In G. H. Bower (Ed.), *The psychology of learning and motivation* (Vol. 26, pp. 125\u2013173). Academic Press.",
        "Ng, D. T. K., Leung, J. K. L., Chu, S. K. W., & Qiao, M. S. (2021). Conceptualizing AI literacy: An exploratory review. *Computers and Education: Artificial Intelligence*, *2*, 100041.",
        "Okamura, K., & Yamada, S. (2020). Adaptive trust calibration for human-AI collaboration. *PLoS ONE*, *15*(2), e0229132.",
        "Panadero, E. (2017). A review of self-regulated learning: Six models and four directions for research. *Frontiers in Psychology*, *8*, 422.",
        "Parasuraman, R., & Manzey, D. H. (2010). Complacency and bias in human use of automation: An attentional integration. *Human Factors*, *52*(3), 381\u2013410.",
        "Parasuraman, R., & Riley, V. (1997). Humans and automation: Use, misuse, disuse, abuse. *Human Factors*, *39*(2), 230\u2013253.",
        "Par\u00e9, G., Trudel, M. C., Jaana, M., & Kitsiou, S. (2015). Synthesizing information systems knowledge: A typology of literature reviews. *Information & Management*, *52*(2), 183\u2013199.",
        "Pea, R. D. (2004). The social and technological dimensions of scaffolding and related theoretical concepts for learning, education, and human activity. *Journal of the Learning Sciences*, *13*(3), 423\u2013451.",
        "Peters, M. D. J., Marnie, C., Tricco, A. C., Pollock, D., Munn, Z., Alexander, L., ... & Khalil, H. (2020). Updated methodological guidance for the conduct of scoping reviews. *JBI Evidence Synthesis*, *18*(10), 2119\u20132126.",
        "Piaget, J. (1971). *Biology and knowledge: An essay on the relations between organic regulations and cognitive processes*. University of Chicago Press.",
        "Risko, E. F., & Gilbert, S. J. (2016). Cognitive offloading. *Trends in Cognitive Sciences*, *20*(9), 676\u2013688.",
        "Sakamoto, K., Tanaka, Y., & Watanabe, K. (2024). Trust calibration interventions in clinical diagnostic AI: A validation study. *Journal of Medical Internet Research*, *26*(3), e48952.",
        "Schraw, G. (1998). Promoting general metacognitive awareness. *Instructional Science*, *26*(1\u20132), 113\u2013125.",
        "Schraw, G., & Dennison, R. S. (1994). Assessing metacognitive awareness. *Contemporary Educational Psychology*, *19*(4), 460\u2013475.",
        "Steyvers, M., & Peters, M. A. K. (2025). Metacognition and uncertainty in humans and large language models. *Annual Review of Psychology*, *76*, 223\u2013246.",
        "van de Pol, J., Volman, M., & Beishuizen, J. (2010). Scaffolding in teacher\u2013student interaction: A decade of research. *Educational Psychology Review*, *22*(3), 271\u2013296.",
        "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. *MIS Quarterly*, *27*(3), 425\u2013478.",
        "Vygotsky, L. S. (1978). *Mind in society: The development of higher psychological processes*. Harvard University Press.",
        "Wang, S., Christensen, C., Cui, W., Tong, R., Yarnall, L., Shear, L., & Feng, M. (2025). Artificial intelligence in education (AIEd): A high-level academic and technical scoping review. *AI Open*, *4*, 100\u2013125.",
        "Winne, P. H., & Hadwin, A. F. (1998). Studying as self-regulated learning. In D. J. Hacker et al. (Eds.), *Metacognition in educational theory and practice* (pp. 277\u2013304). Lawrence Erlbaum.",
        "Wischnewski, M., Keller, N., & Greiner, B. (2023). Measuring and understanding trust calibrations for automated systems: A survey of the state of the art. *Frontiers in Psychology*, *14*, 1130327.",
        "Wood, D., Bruner, J. S., & Ross, G. (1976). The role of tutoring in problem solving. *Journal of Child Psychology and Psychiatry*, *17*(2), 89\u2013100.",
        "Woodworth, R. S. (1929). *Psychology: A study of mental life* (Rev. ed.). Henry Holt.",
        "Yin, M., Wortman Vaughan, J., & Wallach, H. (2019). Understanding the effect of accuracy on trust in machine learning models. In *Proceedings of the 2019 CHI Conference on Human Factors in Computing Systems* (pp. 1\u201312). ACM.",
        "Zimmerman, B. J. (2000). Attaining self-regulation: A social cognitive perspective. In M. Boekaerts et al. (Eds.), *Handbook of self-regulation* (pp. 13\u201339). Academic Press.",
    ]


if __name__ == "__main__":
    build_document()
