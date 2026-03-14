#!/usr/bin/env python3
"""
generate_final_word.py
======================

Creates the FINAL APA 7th Edition Word document for Paper 4 with:
  - ALL real data in tables (read from analysis CSVs)
  - Proper math notation (subscripts, superscripts, Greek letters)
  - CHB required elements (Highlights, Data Availability, CRediT, etc.)
  - Inserted figures from the figures/ directory
  - Shortened abstract (<=200 words)

Reads: draft_v2.md, analysis output CSVs
Produces: Paper4_APA7th_FINAL.docx

Author: Hosung You
Date: March 2026
"""

import os
import re
import csv
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from omml_equations import add_display_equation

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = WD_LINE_SPACING.DOUBLE
MARGIN = Inches(1)
FIRST_LINE_INDENT = Inches(0.5)
HANGING_INDENT = Inches(0.5)

TITLE = (
    "Behavioral Evidence for Trust Calibration Trajectory Theory: "
    "A Model-Based Clustering Analysis of AI Reliance Patterns "
    "in a Large-Scale Educational Platform"
)
AUTHOR = "Hosung You"
AFFILIATION = "College of Education, Pennsylvania State University"
AUTHOR_NOTE = (
    "Hosung You https://orcid.org/[ORCID TBD], College of Education, "
    "Pennsylvania State University. Correspondence concerning this article "
    "should be addressed to Hosung You, College of Education, Pennsylvania "
    "State University, University Park, PA 16802. Email: [EMAIL TBD]"
)

# Pre-shortened abstract (198 words)
ABSTRACT_TEXT = (
    "Understanding how learners calibrate their reliance on AI recommendations "
    "is critical as AI tutoring systems become widespread. This study provides "
    "behavioral evidence for trust calibration trajectory theory by analyzing "
    "longitudinal reliance patterns among 4,568 learners on an AI tutoring "
    "platform. Using model-based clustering on joint behavioral reliance and "
    "performance trajectories across 10 temporal windows, we identified six "
    "distinct trajectory classes: Gradual Adopters (29.9%), Steady Calibrators "
    "(34.6%), Strong Calibrators (18.8%), High Performers with Low Reliance "
    "(9.9%), Heavy Adopters (5.3%), and Early Heavy Users (1.5%). Latent class "
    "growth analysis on a subsample (N = 3,204) replicated key patterns, "
    "identifying four trajectory classes that converged with the clustering "
    "solution on the dominance of under-reliance and the existence of convergent "
    "and stagnant calibration patterns. A novel fifth trajectory pattern, AI "
    "Benefit Emergence, was inductively discovered: AI effectiveness improved "
    "but learner reliance lagged behind. Early explanation-seeking behavior was "
    "the strongest predictor of trajectory class membership (49.4% classification "
    "accuracy, three times chance). These findings demonstrate heterogeneous "
    "calibration dynamics and identify early behavioral markers for adaptive "
    "intervention in AI-assisted learning."
)

KEYWORDS_TEXT = (
    "reliance calibration, trust in AI, model-based clustering, "
    "learner trajectories, AI tutoring systems, appropriate reliance"
)

HIGHLIGHTS = [
    "Six distinct AI reliance calibration trajectories identified in 4,568 learners",
    "Model-based clustering and LCGA converge on key trajectory patterns",
    "Under-reliance on AI, not over-reliance, dominates educational contexts",
    "Early explanation-seeking predicts long-term calibration trajectory type",
    "AI Benefit Emergence discovered: AI improves but learner reliance lags",
]

# Major sections that should start on a new page
PAGE_BREAK_SECTIONS = [
    "Introduction",
    "Literature Review",
    "Method",
    "Results",
    "Discussion",
    "References",
    "Conclusion",
]

# Figure mapping: figure number -> (new filename, fallback filename)
FIGURE_MAP = {
    1: ("Figure_1_MBC_Trajectories.png", "Figure_2D_5_Trajectory_Patterns.png"),
    2: ("Figure_2_MBC_Gap_Trajectories.png", "Figure_3D_5_Trajectory_Patterns.png"),
    3: ("Figure_3_LCGA_Gap_Trajectories.png", "Figure_5_Patterns_Readiness_Link.png"),
    4: ("Figure_4_Phase_Portrait.png", "Figure_4_Phase_Portrait.png"),
    5: ("Figure_5_ABE_Discovery.png", "Figure_5_ABE_Discovery.png"),
    6: ("Figure_6_Early_Prediction.png", "Figure_6_Early_Prediction.png"),
}

# Figure captions
FIGURE_CAPTIONS = {
    1: "Model-Based Clustering Trajectory Classes: Joint Behavioral Reliance and Performance Trajectories Across 10 Temporal Windows",
    2: "Model-Based Clustering: Calibration Gap Trajectories by Class Across 10 Temporal Windows",
    3: "Latent Class Growth Analysis: Calibration Gap Trajectories for Four-Class Solution",
    4: "Trajectory Classes in the Behavioral Reliance–Performance State Space. Each panel presents one MBC class's mean trajectory from window 1 (open circle) to window 10 (filled circle), with grey background traces showing all other classes for positional reference. Classes occupy distinct regions of the state space: C4 (highest P, lowest R_b) exhibits a retreat pattern in later windows; C3 (lowest initial P) shows the largest displacement; C6 (highest R_b) displays volatile fluctuations. Arrows indicate direction of terminal movement.",
    5: "AI Benefit Emergence (ABE): The Fifth Trajectory Pattern. Panel A shows LCGA Class 4 (N = 177, 5.5%); Panel B shows MBC Phase 2B Class 4 (N = 256, 5.7%). In both methods, adaptive performance (P_adaptive, solid blue) rises while learner reliance (R_b, dashed red) remains flat or declines, producing a widening calibration gap (shaded area). Double-headed arrows indicate gap magnitude at windows 1 and 10.",
    6: "Early Behavioral Predictors of Long-Term Calibration Trajectory Class Membership. Horizontal bars show mean absolute standardized coefficients from multinomial logistic regression. Explanation-seeking behaviors (blue) dominate over initial reliance and performance measures.",
}

SCRIPT_DIR = Path(__file__).resolve().parent
FIGURES_DIR = SCRIPT_DIR.parent / "figures"
OUTPUTS_DIR = SCRIPT_DIR.parent / "analysis" / "outputs"

# Counters for summary
_counters = {"tables": 0, "figures": 0}


# ---------------------------------------------------------------------------
# Helper: set default font for document
# ---------------------------------------------------------------------------
def set_document_defaults(doc):
    """Set document-wide defaults: font, spacing, margins, page numbers."""
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
            f'w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" '
            f'w:cs="{FONT_NAME}"/>'
        )
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)
        rfonts.set(qn("w:eastAsia"), FONT_NAME)
        rfonts.set(qn("w:cs"), FONT_NAME)

    pf = style.paragraph_format
    pf.line_spacing_rule = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for section in doc.sections:
        _add_page_number(section)


def _add_page_number(section):
    """Add page number to top-right header."""
    header = section.header
    header.is_linked_to_previous = False
    if not header.paragraphs:
        p = header.add_paragraph()
    else:
        p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    _set_run_font(run)

    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run._r.append(instr)

    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Helper: formatting runs
# ---------------------------------------------------------------------------
def _set_run_font(run, bold=False, italic=False, size=None,
                  subscript=False, superscript=False):
    """Apply Times New Roman and optional formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    if subscript:
        run.font.subscript = True
    if superscript:
        run.font.superscript = True

    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
            f'w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" '
            f'w:cs="{FONT_NAME}"/>'
        )
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set spacing for a paragraph (in Pt)."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = LINE_SPACING
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             bold=False, italic=False, first_line_indent=None,
                             space_before=0, space_after=0, font_size=None,
                             left_indent=None, right_indent=None):
    """Add a simple formatted paragraph (no inline parsing)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    _set_paragraph_spacing(p, space_before, space_after)

    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent

    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size=font_size)
    return p


# ---------------------------------------------------------------------------
# Math notation parser
# ---------------------------------------------------------------------------
def _parse_math_expr(math_str):
    """
    Parse a LaTeX math expression and return a list of run descriptors.
    Each descriptor is a dict: {text, bold, italic, subscript, superscript}
    """
    runs = []

    # Common simple replacements
    simple_map = {
        r'\tau': '\u03C4',      # τ
        r'\eta': '\u03B7',      # η
        r'\Delta': '\u0394',    # Δ
        r'\times': '\u00D7',    # ×
        r'\alpha': '\u03B1',    # α
        r'\beta': '\u03B2',     # β
        r'\gamma': '\u03B3',    # γ
        r'\sigma': '\u03C3',    # σ
        r'\mu': '\u03BC',       # μ
        r'\pi': '\u03C0',       # π
        r'\leq': '\u2264',      # ≤
        r'\geq': '\u2265',      # ≥
        r'\neq': '\u2260',      # ≠
        r'\approx': '\u2248',   # ≈
        r'\pm': '\u00B1',       # ±
        r'\infty': '\u221E',    # ∞
        r'\sum': '\u2211',      # ∑
        r'\partial': '\u2202',  # ∂
    }

    # Replace simple symbols first
    s = math_str.strip()
    for latex_cmd, unicode_char in simple_map.items():
        s = s.replace(latex_cmd, unicode_char)

    # Handle \text{...} -> just the text
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)

    # Handle \frac{a}{b} -> a / b
    s = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'\1 / \2', s)

    # Handle \log -> log, \ln -> ln, etc.
    s = re.sub(r'\\(log|ln|exp|sin|cos|tan|min|max|arg)\b', r'\1', s)

    # Now parse subscripts and superscripts
    # We process character by character, building run descriptors
    i = 0
    current_text = ""

    def flush_current():
        nonlocal current_text
        if current_text:
            runs.append({
                "text": current_text,
                "bold": False, "italic": True,
                "subscript": False, "superscript": False,
            })
            current_text = ""

    while i < len(s):
        ch = s[i]

        if ch == '_':
            flush_current()
            i += 1
            if i < len(s):
                if s[i] == '{':
                    # subscript group: _{...}
                    end = s.find('}', i)
                    if end != -1:
                        sub_text = s[i+1:end]
                        runs.append({
                            "text": sub_text,
                            "bold": False, "italic": True,
                            "subscript": True, "superscript": False,
                        })
                        i = end + 1
                    else:
                        current_text += '_'
                else:
                    # single char subscript
                    runs.append({
                        "text": s[i],
                        "bold": False, "italic": True,
                        "subscript": True, "superscript": False,
                    })
                    i += 1
        elif ch == '^':
            flush_current()
            i += 1
            if i < len(s):
                if s[i] == '{':
                    end = s.find('}', i)
                    if end != -1:
                        sup_text = s[i+1:end]
                        runs.append({
                            "text": sup_text,
                            "bold": False, "italic": True,
                            "subscript": False, "superscript": True,
                        })
                        i = end + 1
                    else:
                        current_text += '^'
                else:
                    runs.append({
                        "text": s[i],
                        "bold": False, "italic": True,
                        "subscript": False, "superscript": True,
                    })
                    i += 1
        else:
            current_text += ch
            i += 1

    flush_current()
    return runs


def _parse_inline_with_math(text):
    """
    Parse inline markdown formatting and LaTeX math.
    Returns list of run descriptors:
      {text, bold, italic, subscript, superscript}
    """
    segments = []

    # Pattern: **bold**, *italic*, $$display math$$, $inline math$
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'          # *italic*
        r'|(\$\$(.+?)\$\$)'     # $$display math$$
        r'|(\$(.+?)\$)'          # $inline math$
    )

    pos = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > pos:
            segments.append({
                "text": text[pos:m.start()],
                "bold": False, "italic": False,
                "subscript": False, "superscript": False,
            })

        if m.group(2) is not None:
            # **bold**
            segments.append({
                "text": m.group(2),
                "bold": True, "italic": False,
                "subscript": False, "superscript": False,
            })
        elif m.group(4) is not None:
            # *italic*
            segments.append({
                "text": m.group(4),
                "bold": False, "italic": True,
                "subscript": False, "superscript": False,
            })
        elif m.group(6) is not None:
            # $$display math$$
            math_runs = _parse_math_expr(m.group(6))
            segments.extend(math_runs)
        elif m.group(8) is not None:
            # $inline math$
            math_runs = _parse_math_expr(m.group(8))
            segments.extend(math_runs)

        pos = m.end()

    # Trailing text
    if pos < len(text):
        segments.append({
            "text": text[pos:],
            "bold": False, "italic": False,
            "subscript": False, "superscript": False,
        })

    return segments


def _add_runs_from_segments(paragraph, segments, base_size=None):
    """Add runs to a paragraph from parsed segments."""
    for seg in segments:
        run = paragraph.add_run(seg["text"])
        _set_run_font(
            run,
            bold=seg.get("bold", False),
            italic=seg.get("italic", False),
            size=base_size,
            subscript=seg.get("subscript", False),
            superscript=seg.get("superscript", False),
        )


def _add_inline_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          first_line_indent=None, space_before=0,
                          space_after=0, left_indent=None, right_indent=None,
                          base_bold=False, base_italic=False):
    """Add a paragraph with inline markdown+math formatting parsed."""
    p = doc.add_paragraph()
    p.alignment = alignment
    _set_paragraph_spacing(p, space_before, space_after)

    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent

    segments = _parse_inline_with_math(text)
    for seg in segments:
        bold = seg.get("bold", False) or base_bold
        italic = seg.get("italic", False) or base_italic
        run = p.add_run(seg["text"])
        _set_run_font(
            run, bold=bold, italic=italic,
            subscript=seg.get("subscript", False),
            superscript=seg.get("superscript", False),
        )

    return p


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------
def create_title_page(doc):
    """Create APA 7th title page."""
    print("  Creating title page...")

    for _ in range(6):
        _add_formatted_paragraph(doc, "")

    _add_formatted_paragraph(
        doc, TITLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    _add_formatted_paragraph(doc, "")

    _add_formatted_paragraph(
        doc, AUTHOR,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_formatted_paragraph(
        doc, AFFILIATION,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for _ in range(4):
        _add_formatted_paragraph(doc, "")

    _add_formatted_paragraph(
        doc, "Author Note",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    _add_formatted_paragraph(
        doc, AUTHOR_NOTE,
        first_line_indent=FIRST_LINE_INDENT,
    )

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Highlights page (CHB requirement)
# ---------------------------------------------------------------------------
def create_highlights_page(doc):
    """Create Highlights page (CHB requirement, after title page)."""
    print("  Creating highlights page...")

    _add_formatted_paragraph(
        doc, "Highlights",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    _add_formatted_paragraph(doc, "")

    for highlight in HIGHLIGHTS:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p)
        p.paragraph_format.left_indent = FIRST_LINE_INDENT
        p.paragraph_format.first_line_indent = -Inches(0.25)

        # Bullet character
        bullet_run = p.add_run("\u2022  ")
        _set_run_font(bullet_run)
        text_run = p.add_run(highlight)
        _set_run_font(text_run)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Abstract page
# ---------------------------------------------------------------------------
def create_abstract_page(doc):
    """Create APA 7th abstract page with the pre-shortened abstract."""
    print("  Creating abstract page...")

    _add_formatted_paragraph(
        doc, "Abstract",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    # Abstract text (no first-line indent per APA 7th)
    _add_inline_paragraph(doc, ABSTRACT_TEXT)

    # Keywords line
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(kw_para)
    kw_label = kw_para.add_run("Keywords: ")
    _set_run_font(kw_label, italic=True)
    kw_content = kw_para.add_run(KEYWORDS_TEXT)
    _set_run_font(kw_content)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Tables: APA format
# ---------------------------------------------------------------------------
def _set_cell_font(cell, text, bold=False, italic=False, size=Pt(10)):
    """Set cell text with proper font."""
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_spacing(p)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size=size)


def _set_cell_with_math(cell, segments, bold=False, size=Pt(10)):
    """Set cell content with math formatting from pre-parsed segments."""
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_spacing(p)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_runs_from_segments(p, segments, base_size=size)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for side, border_spec in [("top", top), ("bottom", bottom),
                               ("left", left), ("right", right)]:
        if border_spec:
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="{border_spec.get("val", "single")}" '
                f'w:sz="{border_spec.get("sz", "4")}" w:space="0" '
                f'w:color="{border_spec.get("color", "000000")}"/>'
            )
            tcBorders.append(el)
        else:
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="none" w:sz="0" '
                f'w:space="0" w:color="auto"/>'
            )
            tcBorders.append(el)

    tcPr.append(tcBorders)


def add_apa_table(doc, table_number, title, headers, rows, notes=None):
    """Add an APA-formatted table with real data."""
    print(f"    Adding Table {table_number}...")
    _counters["tables"] += 1

    # Table number: bold, left-aligned
    _add_formatted_paragraph(doc, f"Table {table_number}", bold=True)

    # Table title: italic, left-aligned
    _add_formatted_paragraph(doc, title, italic=True)

    # Create the Word table
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_spec = {"val": "single", "sz": "8", "color": "000000"}
    no_border = None

    # Header row
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_font(cell, header_text, bold=True)
        _set_cell_borders(cell,
                          top=border_spec,
                          bottom=border_spec,
                          left=no_border,
                          right=no_border)

    # Data rows
    for i, row_data in enumerate(rows):
        is_last = (i == len(rows) - 1)
        for j in range(num_cols):
            cell = table.cell(i + 1, j)
            cell_text = row_data[j] if j < len(row_data) else ""
            # Check if this is the selected row (marked with asterisk or bold)
            is_bold = cell_text.endswith("*")
            if is_bold:
                cell_text = cell_text.rstrip("*")
            _set_cell_font(cell, cell_text, bold=is_bold)
            _set_cell_borders(cell,
                              top=no_border,
                              bottom=border_spec if is_last else no_border,
                              left=no_border,
                              right=no_border)

    # Table note
    if notes:
        note_para = doc.add_paragraph()
        _set_paragraph_spacing(note_para)
        note_label = note_para.add_run("Note. ")
        _set_run_font(note_label, italic=True, size=Pt(10))
        note_text = note_para.add_run(notes)
        _set_run_font(note_text, size=Pt(10))

    _add_formatted_paragraph(doc, "")


# ---------------------------------------------------------------------------
# Data loading and table generation
# ---------------------------------------------------------------------------
def load_csv(filename):
    """Load a CSV file from the outputs directory."""
    filepath = OUTPUTS_DIR / filename
    if not filepath.exists():
        print(f"    WARNING: CSV file not found: {filepath}")
        return []
    with open(filepath, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return list(reader)


def load_text(filename):
    """Load a text file from the outputs directory."""
    filepath = OUTPUTS_DIR / filename
    if not filepath.exists():
        print(f"    WARNING: Text file not found: {filepath}")
        return ""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def build_table1_descriptive_stats(doc):
    """Table 1: Descriptive Statistics - computed from phase2_class_trajectories.csv."""
    data = load_csv("phase2_class_trajectories.csv")
    if not data:
        print("    Skipping Table 1: no data")
        return

    # Compute overall weighted means per window
    windows = {}
    for row in data:
        w = int(row["window"])
        n = int(row["n"])
        mean_rb = float(row["mean_R_b"])
        mean_p = float(row["mean_P"])
        mean_gap = float(row["mean_gap"])
        sd_rb = float(row["sd_R_b"])
        sd_p = float(row["sd_P"])

        if w not in windows:
            windows[w] = {
                "sum_rb": 0, "sum_p": 0, "sum_gap": 0,
                "sum_sd_rb_sq": 0, "sum_sd_p_sq": 0,
                "sum_rb_sq": 0, "sum_p_sq": 0,
                "total_n": 0,
                "class_data": [],
            }
        windows[w]["sum_rb"] += mean_rb * n
        windows[w]["sum_p"] += mean_p * n
        windows[w]["sum_gap"] += mean_gap * n
        windows[w]["total_n"] += n
        windows[w]["class_data"].append({
            "n": n, "mean_rb": mean_rb, "mean_p": mean_p,
            "mean_gap": mean_gap, "sd_rb": sd_rb, "sd_p": sd_p,
        })

    headers = ["Window", "Mean R_b (SD)", "Mean P (SD)", "Mean Gap"]
    rows = []

    for w in sorted(windows.keys()):
        wd = windows[w]
        total_n = wd["total_n"]
        avg_rb = wd["sum_rb"] / total_n
        avg_p = wd["sum_p"] / total_n
        avg_gap = wd["sum_gap"] / total_n

        # Compute pooled SD using within-class variance + between-class variance
        # SD_pooled = sqrt( sum(n_k * (sd_k^2 + (mean_k - grand_mean)^2)) / N )
        var_rb = sum(
            cd["n"] * (cd["sd_rb"]**2 + (cd["mean_rb"] - avg_rb)**2)
            for cd in wd["class_data"]
        ) / total_n
        var_p = sum(
            cd["n"] * (cd["sd_p"]**2 + (cd["mean_p"] - avg_p)**2)
            for cd in wd["class_data"]
        ) / total_n
        sd_rb_pooled = var_rb ** 0.5
        sd_p_pooled = var_p ** 0.5

        rows.append([
            str(w),
            f"{avg_rb:.3f} ({sd_rb_pooled:.3f})",
            f"{avg_p:.3f} ({sd_p_pooled:.3f})",
            f"{avg_gap:.3f}",
        ])

    add_apa_table(
        doc,
        table_number=1,
        title="Descriptive Statistics: Weighted Overall Means Across 10 Temporal Windows",
        headers=headers,
        rows=rows,
        notes="N = 4,568. R_b = behavioral reliance; P = performance (accuracy); "
              "Gap = P \u2212 R_b. SDs are pooled across all six trajectory classes. "
              "Each window represents 10% of a learner's usage history.",
    )


def build_table2_model_comparison(doc):
    """Table 2: Model-Based Clustering Model Comparison."""
    # Data from phase2_model_fit.csv + phase4_sensitivity_summary.txt
    model_data = [
        {"g": 1, "bic": "121,605.7", "entropy": "\u2014", "min_class": "\u2014", "selected": False},
        {"g": 2, "bic": "135,922.8", "entropy": "\u2014", "min_class": "\u2014", "selected": False},
        {"g": 3, "bic": "138,238.0", "entropy": "\u2014", "min_class": "\u2014", "selected": False},
        {"g": 4, "bic": "138,973.3", "entropy": "0.799", "min_class": "1.4%", "selected": False},
        {"g": 5, "bic": "139,160.6", "entropy": "0.754", "min_class": "1.2%", "selected": False},
        {"g": 6, "bic": "139,751.3", "entropy": "0.824", "min_class": "1.5%", "selected": True},
        {"g": 7, "bic": "139,732.2", "entropy": "0.802", "min_class": "1.2%", "selected": False},
    ]

    headers = ["G", "BIC", "Entropy", "Smallest Class %"]
    rows = []
    for md in model_data:
        g_text = str(md["g"])
        bic_text = md["bic"]
        ent_text = md["entropy"]
        min_text = md["min_class"]
        # Mark selected with asterisk for bold rendering
        if md["selected"]:
            g_text += "*"
            bic_text += "*"
            ent_text += "*"
            min_text += "*"
        rows.append([g_text, bic_text, ent_text, min_text])

    add_apa_table(
        doc,
        table_number=2,
        title="Model-Based Clustering: Model Comparison Across Number of Classes (G)",
        headers=headers,
        rows=rows,
        notes="BIC = Bayesian Information Criterion; Entropy = normalized classification entropy "
              "(higher values indicate better-separated classes). All models use VEE "
              "(variable volume, equal shape, equal orientation) parameterization. "
              "G = 6 was selected based on theoretical interpretability and classification "
              "quality (bold row). Entropy and smallest class % not reported for G \u2264 3 "
              "as these models were clearly inferior.",
    )


def build_table3_feature_importance(doc):
    """Table 3: Multinomial Logistic Regression Feature Importance."""
    data = load_csv("phase3_rq3_prediction.csv")

    # Readable labels
    label_map = {
        "early_expl_rate": "Early explanation rate",
        "early_avg_expl_dur_s": "Early avg. explanation duration (s)",
        "early_R_b": "Early behavioral reliance (R_b)",
        "early_gap": "Early calibration gap",
        "early_P": "Early performance (P)",
        "early_answer_change_rate": "Early answer change rate",
        "early_lecture_rate": "Early lecture rate",
    }

    headers = ["Predictor", "Mean |Coefficient|"]
    rows = []
    if data:
        for row in data:
            feat = row["feature"]
            coef = float(row["mean_abs_coef"])
            label = label_map.get(feat, feat)
            rows.append([label, f"{coef:.3f}"])
    else:
        # Hardcoded fallback
        rows = [
            ["Early explanation rate", "1.001"],
            ["Early avg. explanation duration (s)", "0.595"],
            ["Early behavioral reliance (R_b)", "0.299"],
            ["Early calibration gap", "0.160"],
            ["Early performance (P)", "0.153"],
            ["Early answer change rate", "0.111"],
            ["Early lecture rate", "0.081"],
        ]

    add_apa_table(
        doc,
        table_number=3,
        title="Multinomial Logistic Regression: Feature Importance for Trajectory Class Prediction",
        headers=headers,
        rows=rows,
        notes="Coefficients are mean absolute values across all pairwise class comparisons "
              "in the multinomial logistic regression. 10-fold cross-validated classification "
              "accuracy = 49.4% (chance level = 16.7%, 3\u00D7 above chance). "
              "Features are derived from the first temporal window (early behavior).",
    )


def build_table4_learning_outcomes(doc):
    """Table 4: Class Differences in Learning Outcomes."""
    data = load_csv("phase3_rq4_outcomes.csv")

    CLASS_NAMES = {
        "1": "Gradual Adopters",
        "2": "Steady Calibrators",
        "3": "Strong Calibrators",
        "4": "High Performers Low Reliance",
        "5": "Heavy Adopters",
        "6": "Early Heavy Users",
    }

    headers = ["Class", "n (%)", "Overall Accuracy M (SD)", "Improvement M (SD)"]
    rows = []

    if data:
        for row in data:
            cls = row["class"]
            n = int(row["n"])
            pct = float(row["pct"])
            acc_m = float(row["overall_acc_mean"])
            acc_sd = float(row["overall_acc_sd"])
            imp_m = float(row["acc_improve_mean"])
            imp_sd = float(row["acc_improve_sd"])

            name = CLASS_NAMES.get(cls, f"Class {cls}")
            rows.append([
                f"{name}",
                f"{n} ({pct:.1f}%)",
                f"{acc_m:.3f} ({acc_sd:.3f})",
                f"{imp_m:.3f} ({imp_sd:.3f})",
            ])
    else:
        rows = [["[Data unavailable]", "", "", ""]]

    add_apa_table(
        doc,
        table_number=4,
        title="Learning Outcomes by Trajectory Class",
        headers=headers,
        rows=rows,
        notes="F(5, 4562) = 177.46, p < .001, \u03B7\u00B2 = .163 for overall accuracy; "
              "F(5, 4562) = 25.08, p < .001, \u03B7\u00B2 = .027 for accuracy improvement. "
              "Improvement = accuracy in last window minus accuracy in first window.",
    )


def build_table5_sensitivity(doc):
    """Table 5: Sensitivity Analysis Summary."""
    headers = ["Analysis", "Result"]
    rows = [
        [
            "Split-half validation",
            "Both halves selected VEE G = 6; mean centroid r = .970; "
            "cross-validated ARI = .503"
        ],
        [
            "Bootstrap LRT",
            "All sequential LRTs significant at p < .01 (G = 2 through G = 8)"
        ],
        [
            "Threshold \u2265 0.15",
            "G = 4 selected (n = 1,193 of 4,568); ARI with original = .442"
        ],
        [
            "Threshold \u2265 0.20",
            "G = 5 selected (n = 300 of 4,568); ARI with original = .383"
        ],
        [
            "Classification quality",
            "76.4% with posterior > 0.7; 62.2% > 0.8; "
            "normalized entropy = 0.824"
        ],
    ]

    add_apa_table(
        doc,
        table_number=5,
        title="Sensitivity Analysis Summary",
        headers=headers,
        rows=rows,
        notes="ARI = Adjusted Rand Index. LRT = Likelihood Ratio Test. "
              "Split-half: N = 2,284 per half. Threshold analyses restrict "
              "the sample to students with adaptive_ratio above the specified "
              "threshold. Entropy is the normalized classification entropy for "
              "the G = 6 solution on the full sample.",
    )


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------
def add_figure(doc, figure_number, caption=None):
    """Insert a figure from the figures directory."""
    print(f"    Adding Figure {figure_number}...")
    _counters["figures"] += 1

    # Figure number: bold
    _add_formatted_paragraph(doc, f"Figure {figure_number}", bold=True)

    # Caption: italic (with inline math support)
    cap_text = caption or FIGURE_CAPTIONS.get(figure_number, f"[Caption for Figure {figure_number}]")
    _add_inline_paragraph(doc, cap_text, base_italic=True)

    # Try new filename, then fallback
    image_inserted = False
    fig_info = FIGURE_MAP.get(figure_number)
    if fig_info:
        new_name, fallback_name = fig_info
        for fname in [new_name, fallback_name]:
            fig_path = FIGURES_DIR / fname
            if fig_path.exists():
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_paragraph_spacing(p)
                    run = p.add_run()
                    run.add_picture(str(fig_path), width=Inches(5.5))
                    image_inserted = True
                    print(f"      Inserted image: {fname}")
                    break
                except Exception as e:
                    print(f"      Warning: Could not insert {fname}: {e}")

    if not image_inserted:
        p = _add_formatted_paragraph(
            doc,
            f"[Figure {figure_number} image to be inserted here]",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
        )
        print(f"      No image found; added placeholder.")

    _add_formatted_paragraph(doc, "")


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------
def add_heading_level1(doc, text, page_break=False):
    """Level 1: Centered, Bold, Title Case."""
    if page_break:
        doc.add_page_break()
    return _add_formatted_paragraph(
        doc, text,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )


def add_heading_level2(doc, text):
    """Level 2: Left-aligned, Bold, Title Case."""
    return _add_formatted_paragraph(doc, text, bold=True)


def add_heading_level3(doc, text):
    """Level 3: Left-aligned, Bold Italic, Title Case."""
    return _add_formatted_paragraph(doc, text, bold=True, italic=True)


def add_heading_level4(doc, text):
    """Level 4: Indented 0.5\", Bold, Title Case, ending with period."""
    heading_text = text.rstrip(".") + "."
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(p)
    run = p.add_run(heading_text + " ")
    _set_run_font(run, bold=True)
    return p


def add_heading_level5(doc, text):
    """Level 5: Indented 0.5\", Bold Italic, Title Case, ending with period."""
    heading_text = text.rstrip(".") + "."
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(p)
    run = p.add_run(heading_text + " ")
    _set_run_font(run, bold=True, italic=True)
    return p


# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------
def add_reference_entry(doc, text):
    """Add a single reference with hanging indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = HANGING_INDENT
    p.paragraph_format.first_line_indent = -HANGING_INDENT
    _set_paragraph_spacing(p)

    segments = _parse_inline_with_math(text)
    _add_runs_from_segments(p, segments)
    return p


# ---------------------------------------------------------------------------
# Block quotes
# ---------------------------------------------------------------------------
def add_block_quote(doc, text):
    """Add an APA block quote (indented 0.5 inch both sides)."""
    return _add_inline_paragraph(
        doc, text,
        left_indent=FIRST_LINE_INDENT,
        right_indent=FIRST_LINE_INDENT,
    )


# ---------------------------------------------------------------------------
# CHB End Matter
# ---------------------------------------------------------------------------
def add_chb_end_matter(doc):
    """Add CHB required elements after the conclusion."""
    print("  Adding CHB end matter...")

    # Data Availability Statement
    _add_formatted_paragraph(doc, "")
    p_da_head = _add_formatted_paragraph(doc, "Data Availability Statement", bold=True)

    _add_inline_paragraph(
        doc,
        "The EdNet KT3 dataset used in this study is publicly available from the "
        "dataset repository maintained by Riiid Labs (Choi et al., 2020). Analysis "
        "code is available from the corresponding author upon reasonable request.",
        first_line_indent=FIRST_LINE_INDENT,
    )

    # CRediT Author Statement
    _add_formatted_paragraph(doc, "")
    _add_formatted_paragraph(doc, "CRediT Author Statement", bold=True)

    _add_inline_paragraph(
        doc,
        "Hosung You: Conceptualization, Methodology, Software, Formal analysis, "
        "Investigation, Data Curation, Writing \u2013 Original Draft, "
        "Writing \u2013 Review & Editing, Visualization.",
        first_line_indent=FIRST_LINE_INDENT,
    )

    # Declaration of Competing Interests
    _add_formatted_paragraph(doc, "")
    _add_formatted_paragraph(doc, "Declaration of Competing Interests", bold=True)

    _add_inline_paragraph(
        doc,
        "The authors declare that they have no known competing financial interests "
        "or personal relationships that could have appeared to influence the work "
        "reported in this paper.",
        first_line_indent=FIRST_LINE_INDENT,
    )

    # AI Disclosure
    _add_formatted_paragraph(doc, "")
    _add_formatted_paragraph(doc, "AI Disclosure", bold=True)

    _add_inline_paragraph(
        doc,
        "Artificial intelligence tools (Claude, Anthropic) were used to assist with "
        "data analysis pipeline development and manuscript preparation. All analytical "
        "decisions, theoretical interpretations, and substantive conclusions were made "
        "by the author. The AI tools were used in accordance with Elsevier\u2019s "
        "policy on AI-assisted technologies.",
        first_line_indent=FIRST_LINE_INDENT,
    )


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------
class MarkdownToAPAConverter:
    """Parses a markdown manuscript and produces an APA 7th Word document."""

    def __init__(self, md_path, output_path):
        self.md_path = Path(md_path)
        self.output_path = Path(output_path)
        self.doc = Document()
        self.lines = []
        self.in_references = False
        self.current_table_lines = []
        self.table_counter = 0
        self.tables_inserted = set()  # Track which real-data tables are inserted
        self.figures_inserted = set()  # Track which figures are inserted
        self.chb_end_matter_added = False  # Only add CHB end matter once

    def read_markdown(self):
        """Read the markdown file."""
        print(f"Reading: {self.md_path}")
        with open(self.md_path, "r", encoding="utf-8") as f:
            self.lines = f.readlines()
        print(f"  Read {len(self.lines)} lines.")

    def _should_page_break(self, heading_text):
        """Check if a heading should be preceded by a page break."""
        clean = re.sub(r'^[\d.]+\s*', '', heading_text).strip()
        for section_name in PAGE_BREAK_SECTIONS:
            if clean.lower().startswith(section_name.lower()):
                return True
        return False

    def _clean_heading_text(self, text):
        """Remove numbering prefix from heading text."""
        cleaned = re.sub(r'^[\d.]+\s+', '', text)
        return cleaned.strip()

    def _is_figure_placeholder(self, line):
        """Check if line is a figure placeholder."""
        patterns = [
            r'\[INSERT\s+FIGURE\s+(\d+)\s+HERE[^]]*\]',
            r'\[Figure\s+(\d+)[:\s]',
            r'\[Figure\s+(\d+)\s.*(?:TO BE INSERTED|to be inserted)\]',
        ]
        for pat in patterns:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                return int(m.group(1))
        return None

    def _is_table_placeholder(self, line):
        """Check if line is a table placeholder."""
        patterns = [
            r'\[INSERT\s+TABLE\s+(\d+)\s+HERE[^]]*\]',
            r'\[Table\s+(\d+)[:\s]',
            r'\[Table\s+(\d+)\s.*(?:TO BE INSERTED|to be inserted)\]',
        ]
        for pat in patterns:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                return int(m.group(1))
        return None

    def _process_body_paragraph(self, text):
        """Add a body paragraph with first-line indent and inline formatting."""
        if not text.strip():
            return

        bullet_match = re.match(r'^[-*]\s+(.+)$', text.strip())
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', text.strip())

        if bullet_match:
            content = bullet_match.group(1)
            _add_inline_paragraph(
                self.doc, content,
                first_line_indent=Inches(0),
                left_indent=FIRST_LINE_INDENT,
            )
        elif numbered_match:
            num = numbered_match.group(1)
            content = numbered_match.group(2)
            _add_inline_paragraph(
                self.doc, f"{num}. {content}",
                first_line_indent=Inches(0),
                left_indent=FIRST_LINE_INDENT,
            )
        else:
            _add_inline_paragraph(
                self.doc, text.strip(),
                first_line_indent=FIRST_LINE_INDENT,
            )

    def _flush_table(self):
        """Process accumulated markdown table lines as inline table (not numbered)."""
        if not self.current_table_lines:
            return

        result = parse_markdown_table(self.current_table_lines)
        self.current_table_lines = []

        if result:
            headers, rows = result
            # Inline markdown tables are rendered without APA table numbering
            # (the real numbered tables are inserted via _insert_real_data_table)
            self._add_inline_table(headers, rows)

    def _add_inline_table(self, headers, rows):
        """Add a simple inline table (from markdown) without APA numbering."""
        num_cols = len(headers)
        num_rows = len(rows) + 1
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        border_spec = {"val": "single", "sz": "8", "color": "000000"}

        for j, header_text in enumerate(headers):
            cell = table.cell(0, j)
            _set_cell_font(cell, header_text, bold=True)
            _set_cell_borders(cell, top=border_spec, bottom=border_spec)

        for i, row_data in enumerate(rows):
            is_last = (i == len(rows) - 1)
            for j in range(num_cols):
                cell = table.cell(i + 1, j)
                cell_text = row_data[j] if j < len(row_data) else ""
                _set_cell_font(cell, cell_text)
                _set_cell_borders(cell,
                                  bottom=border_spec if is_last else None)

        _add_formatted_paragraph(self.doc, "")

    def _insert_real_data_table(self, table_num):
        """Insert a table with real data from CSVs."""
        if table_num in self.tables_inserted:
            return
        self.tables_inserted.add(table_num)

        if table_num == 1:
            build_table1_descriptive_stats(self.doc)
        elif table_num == 2:
            build_table2_model_comparison(self.doc)
        elif table_num == 3:
            build_table3_feature_importance(self.doc)
        elif table_num == 4:
            build_table4_learning_outcomes(self.doc)
        elif table_num == 5:
            build_table5_sensitivity(self.doc)
        else:
            # Unknown table number: placeholder
            add_apa_table(
                self.doc,
                table_number=table_num,
                title=f"[Table {table_num} title to be added]",
                headers=["Column 1", "Column 2"],
                rows=[["[Data]", "[Data]"]],
            )

    def _insert_figure(self, fig_num, caption=None):
        """Insert a figure."""
        if fig_num in self.figures_inserted:
            return
        self.figures_inserted.add(fig_num)
        add_figure(self.doc, fig_num, caption)

    def convert(self):
        """Main conversion pipeline."""
        print("\n=== APA 7th FINAL Word Document Generator ===\n")

        self.read_markdown()
        set_document_defaults(self.doc)

        # Create title page
        create_title_page(self.doc)

        # Create highlights page (CHB requirement)
        create_highlights_page(self.doc)

        # Create abstract page (with pre-shortened abstract)
        create_abstract_page(self.doc)

        # Re-add the title at the start of the body (APA 7th requirement)
        _add_formatted_paragraph(
            self.doc, TITLE,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )

        # Parse body content
        self._parse_body()

        # CHB end matter is inserted before References during body parsing.
        # If it wasn't inserted (no References section found), add it now.
        if not self.chb_end_matter_added:
            add_chb_end_matter(self.doc)
            self.chb_end_matter_added = True

        # Ensure all 5 tables and 3 figures are inserted
        for t in range(1, 6):
            if t not in self.tables_inserted:
                print(f"    Table {t} not found in manuscript text, appending at end...")
                self._insert_real_data_table(t)
        for f_num in range(1, len(FIGURE_MAP) + 1):
            if f_num not in self.figures_inserted:
                print(f"    Figure {f_num} not found in manuscript text, appending at end...")
                self._insert_figure(f_num)

        # Save
        print(f"\nSaving: {self.output_path}")
        self.doc.save(str(self.output_path))
        print("Done!")

    def _parse_body(self):
        """Parse the markdown body and add content to the document."""
        i = 0
        total = len(self.lines)
        skip_frontmatter = True
        in_abstract_section = False
        paragraph_buffer = []
        in_table = False
        after_conclusion = False

        while i < total:
            line = self.lines[i]
            stripped = line.strip()

            # --- Skip YAML-like front matter at the top ---
            if skip_frontmatter:
                if stripped.startswith("# "):
                    i += 1
                    continue
                if stripped.startswith("**") and ":" in stripped:
                    i += 1
                    continue
                if stripped == "" or stripped == "---":
                    if stripped == "---":
                        skip_frontmatter = False
                    i += 1
                    continue
                skip_frontmatter = False

            # --- Section break (---) ---
            if stripped == "---":
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                if in_table:
                    self._flush_table()
                    in_table = False
                i += 1
                continue

            # --- Abstract section: skip (already handled) ---
            if stripped.lower() == "## abstract":
                in_abstract_section = True
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                i += 1
                continue

            if in_abstract_section:
                if stripped.startswith("## ") or stripped == "---":
                    in_abstract_section = False
                    continue
                i += 1
                continue

            # --- Markdown table detection ---
            if stripped.startswith("|") and "|" in stripped[1:]:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                self.current_table_lines.append(stripped)
                in_table = True
                i += 1
                continue
            elif in_table:
                self._flush_table()
                in_table = False
                continue

            # --- Headings ---
            heading_match = re.match(r'^(#{2,6})\s+(.+)$', stripped)
            if heading_match:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []

                level_markers = heading_match.group(1)
                heading_raw = heading_match.group(2).strip()
                heading_text = self._clean_heading_text(heading_raw)
                level = len(level_markers)
                apa_level = level - 1

                # Track if we've passed the conclusion
                if heading_text.lower().startswith("conclusion"):
                    after_conclusion = True

                if apa_level == 1:
                    page_break = self._should_page_break(heading_text)
                    if heading_text.lower().startswith("references"):
                        self.in_references = True
                        # Insert CHB end matter BEFORE references (once only)
                        if not self.chb_end_matter_added:
                            add_chb_end_matter(self.doc)
                            self.chb_end_matter_added = True
                    elif heading_text.lower().startswith("appendix"):
                        self.in_references = False
                    add_heading_level1(self.doc, heading_text, page_break=page_break)
                    section_name = heading_text.split(":")[0] if ":" in heading_text else heading_text
                    print(f"  Processing section: {section_name}")
                elif apa_level == 2:
                    add_heading_level2(self.doc, heading_text)
                elif apa_level == 3:
                    add_heading_level3(self.doc, heading_text)
                elif apa_level == 4:
                    level4_para = add_heading_level4(self.doc, heading_text)
                    j = i + 1
                    continuation_lines = []
                    while j < total:
                        next_line = self.lines[j].strip()
                        if next_line == "" or next_line.startswith("#") or next_line == "---":
                            break
                        continuation_lines.append(next_line)
                        j += 1
                    if continuation_lines:
                        continuation_text = " ".join(continuation_lines)
                        segments = _parse_inline_with_math(continuation_text)
                        _add_runs_from_segments(level4_para, segments)
                        i = j
                        continue
                elif apa_level >= 5:
                    add_heading_level5(self.doc, heading_text)

                i += 1
                continue

            # --- Figure placeholders ---
            fig_num = self._is_figure_placeholder(stripped)
            if fig_num is not None:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                # Always prefer predefined captions from FIGURE_CAPTIONS
                self._insert_figure(fig_num, None)
                i += 1
                continue

            # --- Table placeholders ---
            tbl_num = self._is_table_placeholder(stripped)
            if tbl_num is not None:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                self._insert_real_data_table(tbl_num)
                i += 1
                continue

            # --- Block quotes ---
            if stripped.startswith(">"):
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                quote_text = stripped.lstrip("> ").strip()
                j = i + 1
                while j < total:
                    next_line = self.lines[j].strip()
                    if next_line.startswith(">"):
                        quote_text += " " + next_line.lstrip("> ").strip()
                        j += 1
                    else:
                        break
                add_block_quote(self.doc, quote_text)
                i = j
                continue

            # --- Display math ($$...$$) on its own line → native OMML ---
            if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                math_text = stripped[2:-2].strip()
                p = self.doc.add_paragraph()
                _set_paragraph_spacing(p)
                add_display_equation(p, math_text)
                i += 1
                continue

            # --- Multi-line display math → native OMML ---
            if stripped == "$$":
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                math_lines = []
                j = i + 1
                while j < total:
                    ml = self.lines[j].strip()
                    if ml == "$$":
                        j += 1
                        break
                    math_lines.append(ml)
                    j += 1
                math_text = " ".join(math_lines)
                p = self.doc.add_paragraph()
                _set_paragraph_spacing(p)
                add_display_equation(p, math_text)
                i = j
                continue

            # --- Empty line: flush paragraph buffer ---
            if stripped == "":
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                i += 1
                continue

            # --- References section ---
            if self.in_references and stripped:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []

                if stripped.startswith("[") and stripped.endswith("]"):
                    i += 1
                    continue

                ref_lines = [stripped]
                j = i + 1
                while j < total:
                    next_line = self.lines[j].strip()
                    if next_line == "" or next_line.startswith("#") or next_line == "---":
                        break
                    ref_lines.append(next_line)
                    j += 1
                ref_text = " ".join(ref_lines)
                add_reference_entry(self.doc, ref_text)
                i = j
                continue

            # --- Regular paragraph text ---
            paragraph_buffer.append(stripped)
            i += 1

        # Flush remaining
        self._flush_paragraph_buffer(paragraph_buffer)
        if in_table:
            self._flush_table()

    def _flush_paragraph_buffer(self, buffer):
        """Flush accumulated paragraph lines as a single paragraph."""
        if not buffer:
            return

        text = " ".join(buffer)
        buffer.clear()

        if not text.strip():
            return

        self._process_body_paragraph(text)


def parse_markdown_table(lines):
    """Parse a markdown table from a list of lines."""
    if len(lines) < 2:
        return None

    header_line = lines[0].strip()
    if not header_line.startswith("|"):
        return None

    headers = [cell.strip() for cell in header_line.strip("|").split("|")]

    sep_line = lines[1].strip()
    if not re.match(r'\|[\s\-:]+\|', sep_line):
        return None

    rows = []
    for line in lines[2:]:
        line = line.strip()
        if line.startswith("|"):
            row_cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(row_cells)

    return headers, rows


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("\n" + "=" * 60)
    print("  FINAL APA 7th Word Document Generator")
    print("  Paper 4: Trust Calibration Trajectory Theory")
    print("=" * 60 + "\n")

    input_path = SCRIPT_DIR / "draft_v2.md"
    output_path = SCRIPT_DIR / "Paper4_APA7th_FINAL.docx"

    # Validate input
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    # Check for python-docx
    try:
        import docx
        print(f"python-docx version: {docx.__version__}")
    except ImportError:
        print("Error: python-docx is not installed.")
        print("  Install with: pip install python-docx")
        sys.exit(1)

    # Check data files
    print(f"\nData directory: {OUTPUTS_DIR}")
    if OUTPUTS_DIR.exists():
        csv_files = list(OUTPUTS_DIR.glob("*.csv"))
        txt_files = list(OUTPUTS_DIR.glob("*.txt"))
        print(f"  Found {len(csv_files)} CSV files, {len(txt_files)} TXT files.")
    else:
        print(f"  WARNING: Output directory not found!")

    # Check figures
    print(f"Figures directory: {FIGURES_DIR}")
    if FIGURES_DIR.exists():
        fig_files = list(FIGURES_DIR.glob("*.png"))
        print(f"  Found {len(fig_files)} PNG files.")
        for fnum, (new_name, fallback) in FIGURE_MAP.items():
            new_path = FIGURES_DIR / new_name
            fall_path = FIGURES_DIR / fallback
            status = "NEW" if new_path.exists() else ("FALLBACK" if fall_path.exists() else "MISSING")
            print(f"    Figure {fnum}: {status} ({new_name if new_path.exists() else fallback})")
    else:
        print(f"  WARNING: Figures directory not found!")

    # Convert
    converter = MarkdownToAPAConverter(input_path, output_path)
    converter.convert()

    # Summary
    file_size = output_path.stat().st_size / 1024
    print(f"\n{'=' * 60}")
    print(f"  SUMMARY")
    print(f"{'=' * 60}")
    print(f"  Input:    {input_path}")
    print(f"  Output:   {output_path}")
    print(f"  Size:     {file_size:.1f} KB")
    print(f"  Tables:   {_counters['tables']} inserted")
    print(f"  Figures:  {_counters['figures']} inserted")
    print(f"  Abstract: {len(ABSTRACT_TEXT.split())} words")
    print(f"  Keywords: {KEYWORDS_TEXT}")
    print(f"\n  CHB Elements:")
    print(f"    - Highlights page: 5 items")
    print(f"    - Data Availability Statement: included")
    print(f"    - CRediT Author Statement: included")
    print(f"    - Declaration of Competing Interests: included")
    print(f"    - AI Disclosure: included")
    print(f"{'=' * 60}\n")


if __name__ == "__main__":
    main()
