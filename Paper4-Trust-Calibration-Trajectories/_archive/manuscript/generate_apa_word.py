#!/usr/bin/env python3
"""
generate_apa_word.py
====================

Converts a Markdown manuscript (draft_v2.md) to an APA 7th Edition formatted
Word document (Paper4_APA7th_v2.docx).

APA 7th formatting includes:
  - Times New Roman 12pt throughout
  - Double-spaced, 1-inch margins
  - Page numbers top-right
  - Title page, abstract, body, references with proper heading levels
  - APA table formatting (horizontal rules only)
  - Hanging indent references
  - Inline bold, italic, and math notation parsing

Usage:
    python generate_apa_word.py
    python generate_apa_word.py --input draft_v2.md --output Paper4_APA7th_v2.docx

Author: Hosung You
Date: March 2026
"""

import os
import re
import sys
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = WD_LINE_SPACING.DOUBLE
MARGIN = Inches(1)  # 2.54 cm
FIRST_LINE_INDENT = Inches(0.5)
HANGING_INDENT = Inches(0.5)

TITLE = (
    "Behavioral Evidence for Trust Calibration Trajectory Theory: "
    "A Model-Based Clustering Analysis of AI Reliance Patterns "
    "in a Large-Scale Educational Platform"
)
AUTHOR = "Hosung You"
AFFILIATION = "College of Education, Pennsylvania State University"
AUTHOR_NOTE = (
    "Hosung You https://orcid.org/[ORCID TBD], College of Education, "
    "Pennsylvania State University. Correspondence concerning this article "
    "should be addressed to Hosung You, College of Education, Pennsylvania "
    "State University, University Park, PA 16802. Email: [EMAIL TBD]"
)

# Major sections that should start on a new page
PAGE_BREAK_SECTIONS = [
    "Introduction",
    "Literature Review",
    "Method",
    "Results",
    "Discussion",
    "References",
    "Conclusion",
]

# Figure mapping: figure reference text -> filename in ../figures/
FIGURE_MAP = {
    "1": "Figure_2D_5_Trajectory_Patterns.png",
    "2": "Figure_3D_5_Trajectory_Patterns.png",
    "3": "Figure_2D_5_Trajectory_Patterns.png",  # fallback
    "4": "Figure_3D_5_Trajectory_Patterns.png",  # fallback
    "5": "Figure_5_Patterns_Readiness_Link.png",
}

SCRIPT_DIR = Path(__file__).resolve().parent
FIGURES_DIR = SCRIPT_DIR.parent / "figures"


# ---------------------------------------------------------------------------
# Helper: set default font for document
# ---------------------------------------------------------------------------
def set_document_defaults(doc):
    """Set document-wide defaults: font, spacing, margins, page numbers."""
    # -- Margins --
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # -- Default font via styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)

    # Set East-Asian fallback font via XML
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
                           f'w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" '
                           f'w:cs="{FONT_NAME}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)
        rfonts.set(qn("w:eastAsia"), FONT_NAME)
        rfonts.set(qn("w:cs"), FONT_NAME)

    # -- Paragraph defaults: double spacing, no extra before/after --
    pf = style.paragraph_format
    pf.line_spacing_rule = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # -- Page numbers (top-right, starting from page 1) --
    for section in doc.sections:
        _add_page_number(section)


def _add_page_number(section):
    """Add page number to top-right header of the given section."""
    header = section.header
    header.is_linked_to_previous = False
    if not header.paragraphs:
        p = header.add_paragraph()
    else:
        p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Set font for the header paragraph
    run = p.add_run()
    _set_run_font(run)

    # Insert PAGE field
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run._r.append(instr)

    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Helper: formatting runs
# ---------------------------------------------------------------------------
def _set_run_font(run, bold=False, italic=False, size=None):
    """Apply Times New Roman and optional bold/italic to a run."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Ensure East-Asian font name is also set
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
                           f'w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" '
                           f'w:cs="{FONT_NAME}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _set_paragraph_spacing(paragraph, space_before=0, space_after=0):
    """Set spacing for a paragraph (in Pt)."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = LINE_SPACING
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             bold=False, italic=False, first_line_indent=None,
                             space_before=0, space_after=0, font_size=None,
                             left_indent=None, right_indent=None):
    """Add a simple formatted paragraph (no inline parsing)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    _set_paragraph_spacing(p, space_before, space_after)

    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent

    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size=font_size)
    return p


# ---------------------------------------------------------------------------
# Inline formatting parser
# ---------------------------------------------------------------------------
def _parse_inline(text):
    """
    Parse inline markdown formatting and return a list of segments.
    Each segment is a dict with keys: text, bold, italic.

    Handles:
      - **bold**
      - *italic*
      - $math$ (rendered as italic)
      - $$display math$$ (rendered as italic)
    """
    segments = []

    # Regex pattern to match inline formatting tokens
    # Order matters: ** before *, $$ before $
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(\$\$(.+?)\$\$)'      # $$display math$$
        r'|(\$(.+?)\$)'           # $inline math$
    )

    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            segments.append({
                "text": text[pos:m.start()],
                "bold": False,
                "italic": False,
            })

        if m.group(2) is not None:
            # **bold**
            segments.append({"text": m.group(2), "bold": True, "italic": False})
        elif m.group(4) is not None:
            # *italic*
            segments.append({"text": m.group(4), "bold": False, "italic": True})
        elif m.group(6) is not None:
            # $$display math$$ -> italic
            segments.append({"text": m.group(6), "bold": False, "italic": True})
        elif m.group(8) is not None:
            # $inline math$ -> italic
            segments.append({"text": m.group(8), "bold": False, "italic": True})

        pos = m.end()

    # Trailing plain text
    if pos < len(text):
        segments.append({"text": text[pos:], "bold": False, "italic": False})

    return segments


def _add_runs_from_segments(paragraph, segments):
    """Add runs to a paragraph from parsed segments."""
    for seg in segments:
        run = paragraph.add_run(seg["text"])
        _set_run_font(run, bold=seg["bold"], italic=seg["italic"])


def _add_inline_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          first_line_indent=None, space_before=0,
                          space_after=0, left_indent=None, right_indent=None,
                          base_bold=False, base_italic=False):
    """Add a paragraph with inline markdown formatting parsed."""
    p = doc.add_paragraph()
    p.alignment = alignment
    _set_paragraph_spacing(p, space_before, space_after)

    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent

    segments = _parse_inline(text)
    for seg in segments:
        bold = seg["bold"] or base_bold
        italic = seg["italic"] or base_italic
        run = p.add_run(seg["text"])
        _set_run_font(run, bold=bold, italic=italic)

    return p


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------
def create_title_page(doc):
    """Create APA 7th title page."""
    print("  Creating title page...")

    # Add some blank lines to push title to upper half
    for _ in range(6):
        _add_formatted_paragraph(doc, "")

    # Title: bold, centered
    _add_formatted_paragraph(
        doc, TITLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    # Blank line
    _add_formatted_paragraph(doc, "")

    # Author name: centered
    _add_formatted_paragraph(
        doc, AUTHOR,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Affiliation: centered
    _add_formatted_paragraph(
        doc, AFFILIATION,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Push down a bit for Author Note
    for _ in range(4):
        _add_formatted_paragraph(doc, "")

    # Author Note heading
    _add_formatted_paragraph(
        doc, "Author Note",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    # Author note content with first-line indent
    _add_formatted_paragraph(
        doc, AUTHOR_NOTE,
        first_line_indent=FIRST_LINE_INDENT,
    )

    # Page break after title page
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Abstract page
# ---------------------------------------------------------------------------
def create_abstract_page(doc, abstract_text, keywords_text):
    """Create APA 7th abstract page."""
    print("  Creating abstract page...")

    # "Abstract" heading: centered, bold
    _add_formatted_paragraph(
        doc, "Abstract",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    # Abstract text: single paragraph, no first-line indent
    _add_inline_paragraph(doc, abstract_text)

    # Keywords line: italic label, then keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(kw_para)
    kw_label = kw_para.add_run("Keywords: ")
    _set_run_font(kw_label, italic=True)
    kw_content = kw_para.add_run(keywords_text)
    _set_run_font(kw_content)

    # Page break after abstract
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------
def add_heading_level1(doc, text, page_break=False):
    """Level 1: Centered, Bold, Title Case."""
    if page_break:
        doc.add_page_break()
    p = _add_formatted_paragraph(
        doc, text,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    return p


def add_heading_level2(doc, text):
    """Level 2: Left-aligned, Bold, Title Case."""
    p = _add_formatted_paragraph(doc, text, bold=True)
    return p


def add_heading_level3(doc, text):
    """Level 3: Left-aligned, Bold Italic, Title Case."""
    p = _add_formatted_paragraph(doc, text, bold=True, italic=True)
    return p


def add_heading_level4(doc, text):
    """
    Level 4: Indented 0.5", Bold, Title Case, ending with period.
    Text continues on same line (caller must handle continuation).
    Returns the paragraph so caller can append text.
    """
    # Ensure text ends with period
    heading_text = text.rstrip(".")
    heading_text += "."
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(p)
    run = p.add_run(heading_text + " ")
    _set_run_font(run, bold=True)
    return p


def add_heading_level5(doc, text):
    """Level 5: Indented 0.5", Bold Italic, Title Case, ending with period."""
    heading_text = text.rstrip(".")
    heading_text += "."
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _set_paragraph_spacing(p)
    run = p.add_run(heading_text + " ")
    _set_run_font(run, bold=True, italic=True)
    return p


# ---------------------------------------------------------------------------
# Tables: APA format
# ---------------------------------------------------------------------------
def _set_cell_font(cell, text, bold=False, italic=False):
    """Set cell text with proper font."""
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_spacing(p)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic, size=Pt(10))


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell. Each border is a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for side, border_spec in [("top", top), ("bottom", bottom),
                               ("left", left), ("right", right)]:
        if border_spec:
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="{border_spec.get("val", "single")}" '
                f'w:sz="{border_spec.get("sz", "4")}" w:space="0" '
                f'w:color="{border_spec.get("color", "000000")}"/>'
            )
            tcBorders.append(el)
        else:
            # No border
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="none" w:sz="0" '
                f'w:space="0" w:color="auto"/>'
            )
            tcBorders.append(el)

    tcPr.append(tcBorders)


def add_apa_table(doc, table_number, title, headers, rows, notes=None):
    """
    Add an APA-formatted table.

    Args:
        doc: Document object
        table_number: int, e.g. 1
        title: str, table title (will be italic)
        headers: list of str, column headers
        rows: list of list of str, data rows
        notes: str or None, table notes
    """
    print(f"    Adding Table {table_number}...")

    # Table number: bold, left-aligned
    _add_formatted_paragraph(doc, f"Table {table_number}", bold=True)

    # Table title: italic, left-aligned
    _add_formatted_paragraph(doc, title, italic=True)

    # Create the Word table
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_spec = {"val": "single", "sz": "8", "color": "000000"}
    no_border = None

    # Header row
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_font(cell, header_text, bold=True)
        # Top and bottom borders on header row
        _set_cell_borders(cell,
                          top=border_spec,
                          bottom=border_spec,
                          left=no_border,
                          right=no_border)

    # Data rows
    for i, row_data in enumerate(rows):
        is_last = (i == len(rows) - 1)
        for j in range(num_cols):
            cell = table.cell(i + 1, j)
            cell_text = row_data[j] if j < len(row_data) else ""
            _set_cell_font(cell, cell_text)
            # Bottom border only on last row
            _set_cell_borders(cell,
                              top=no_border,
                              bottom=border_spec if is_last else no_border,
                              left=no_border,
                              right=no_border)

    # Table note
    if notes:
        note_para = doc.add_paragraph()
        _set_paragraph_spacing(note_para)
        note_label = note_para.add_run("Note. ")
        _set_run_font(note_label, italic=True, size=Pt(10))
        note_text = note_para.add_run(notes)
        _set_run_font(note_text, size=Pt(10))

    # Blank line after table
    _add_formatted_paragraph(doc, "")


def parse_markdown_table(lines):
    """
    Parse a markdown table from a list of lines.
    Returns (headers, rows) or None if not a valid table.
    """
    if len(lines) < 2:
        return None

    # First line should be header
    header_line = lines[0].strip()
    if not header_line.startswith("|"):
        return None

    headers = [cell.strip() for cell in header_line.strip("|").split("|")]

    # Second line should be separator (|---|---|...)
    sep_line = lines[1].strip()
    if not re.match(r'\|[\s\-:]+\|', sep_line):
        return None

    # Remaining lines are data rows
    rows = []
    for line in lines[2:]:
        line = line.strip()
        if line.startswith("|"):
            row_cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(row_cells)

    return headers, rows


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------
def add_figure_placeholder(doc, figure_number, caption=None):
    """Add a figure placeholder or insert the actual image if available."""
    print(f"    Adding Figure {figure_number}...")

    # Figure number: bold
    _add_formatted_paragraph(doc, f"Figure {figure_number}", bold=True)

    # Figure title/caption: italic
    if caption:
        _add_formatted_paragraph(doc, caption, italic=True)

    # Try to find and insert the image
    fig_key = str(figure_number)
    filename = FIGURE_MAP.get(fig_key)
    image_inserted = False

    if filename:
        fig_path = FIGURES_DIR / filename
        if fig_path.exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_spacing(p)
                run = p.add_run()
                run.add_picture(str(fig_path), width=Inches(5.5))
                image_inserted = True
                print(f"      Inserted image: {filename}")
            except Exception as e:
                print(f"      Warning: Could not insert image {filename}: {e}")

    if not image_inserted:
        # Add placeholder box
        p = _add_formatted_paragraph(
            doc,
            f"[Figure {figure_number} image to be inserted here]",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
        )
        # Add a simple border-like indicator
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="6" w:color="999999"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="6" w:color="999999"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="6" w:color="999999"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="6" w:color="999999"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        print(f"      No image found; added placeholder text.")

    # Blank line after figure
    _add_formatted_paragraph(doc, "")


# ---------------------------------------------------------------------------
# References section
# ---------------------------------------------------------------------------
def add_reference_entry(doc, text):
    """Add a single reference with hanging indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = HANGING_INDENT
    p.paragraph_format.first_line_indent = -HANGING_INDENT
    _set_paragraph_spacing(p)

    # Parse inline formatting (italic journal names, etc.)
    segments = _parse_inline(text)
    _add_runs_from_segments(p, segments)
    return p


# ---------------------------------------------------------------------------
# Block quotes
# ---------------------------------------------------------------------------
def add_block_quote(doc, text):
    """Add an APA block quote (indented 0.5 inch both sides)."""
    p = _add_inline_paragraph(
        doc, text,
        left_indent=FIRST_LINE_INDENT,
        right_indent=FIRST_LINE_INDENT,
    )
    return p


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------
class MarkdownToAPAConverter:
    """Parses a markdown manuscript and produces an APA 7th Word document."""

    def __init__(self, md_path, output_path):
        self.md_path = Path(md_path)
        self.output_path = Path(output_path)
        self.doc = Document()
        self.lines = []
        self.abstract_text = ""
        self.keywords_text = ""
        self.in_references = False
        self.current_table_lines = []
        self.table_counter = 0

    def read_markdown(self):
        """Read the markdown file."""
        print(f"Reading: {self.md_path}")
        with open(self.md_path, "r", encoding="utf-8") as f:
            self.lines = f.readlines()
        print(f"  Read {len(self.lines)} lines.")

    def _extract_abstract_and_keywords(self):
        """Extract abstract and keywords from the markdown."""
        in_abstract = False
        abstract_lines = []
        keywords = ""

        for line in self.lines:
            stripped = line.strip()
            if stripped.lower().startswith("## abstract") or stripped.lower() == "## abstract":
                in_abstract = True
                continue
            if in_abstract:
                if stripped.startswith("## ") or stripped == "---":
                    break
                if stripped.lower().startswith("**keywords:**"):
                    keywords = stripped.replace("**Keywords:**", "").replace(
                        "**keywords:**", "").strip()
                    continue
                if stripped:
                    abstract_lines.append(stripped)

        self.abstract_text = " ".join(abstract_lines)
        self.keywords_text = keywords

    def _should_page_break(self, heading_text):
        """Check if a heading should be preceded by a page break."""
        clean = re.sub(r'^[\d.]+\s*', '', heading_text).strip()
        for section_name in PAGE_BREAK_SECTIONS:
            if clean.lower().startswith(section_name.lower()):
                return True
        return False

    def _clean_heading_text(self, text):
        """Remove numbering prefix from heading text (e.g., '1. Introduction' -> 'Introduction')."""
        # Remove patterns like "1. ", "2.1 ", "3.4.2 ", etc.
        cleaned = re.sub(r'^[\d.]+\s+', '', text)
        return cleaned.strip()

    def _is_figure_placeholder(self, line):
        """Check if line is a figure placeholder like [INSERT FIGURE X HERE] or [Figure X ...]."""
        patterns = [
            r'\[INSERT\s+FIGURE\s+(\d+)\s+HERE[^]]*\]',
            r'\[Figure\s+(\d+)[:\s]',
            r'\[Figure\s+(\d+)\s.*(?:TO BE INSERTED|to be inserted)\]',
        ]
        for pat in patterns:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                return int(m.group(1))
        return None

    def _is_table_placeholder(self, line):
        """Check if line is a table placeholder like [INSERT TABLE X HERE] or [Table X ...]."""
        patterns = [
            r'\[INSERT\s+TABLE\s+(\d+)\s+HERE[^]]*\]',
            r'\[Table\s+(\d+)[:\s]',
            r'\[Table\s+(\d+)\s.*(?:TO BE INSERTED|to be inserted)\]',
        ]
        for pat in patterns:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                return int(m.group(1))
        return None

    def _process_body_paragraph(self, text):
        """Add a body paragraph with first-line indent and inline formatting."""
        if not text.strip():
            return

        # Check for bullet/list items
        bullet_match = re.match(r'^[-*]\s+(.+)$', text.strip())
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', text.strip())

        if bullet_match:
            content = bullet_match.group(1)
            _add_inline_paragraph(
                self.doc, content,
                first_line_indent=Inches(0),
                left_indent=FIRST_LINE_INDENT,
            )
        elif numbered_match:
            num = numbered_match.group(1)
            content = numbered_match.group(2)
            _add_inline_paragraph(
                self.doc, f"{num}. {content}",
                first_line_indent=Inches(0),
                left_indent=FIRST_LINE_INDENT,
            )
        else:
            _add_inline_paragraph(
                self.doc, text.strip(),
                first_line_indent=FIRST_LINE_INDENT,
            )

    def _flush_table(self):
        """Process accumulated markdown table lines."""
        if not self.current_table_lines:
            return

        result = parse_markdown_table(self.current_table_lines)
        self.current_table_lines = []

        if result:
            headers, rows = result
            self.table_counter += 1
            add_apa_table(
                self.doc,
                table_number=self.table_counter,
                title="[Table title to be added]",
                headers=headers,
                rows=rows,
            )

    def convert(self):
        """Main conversion pipeline."""
        print("\n=== APA 7th Word Document Generator ===\n")

        self.read_markdown()
        set_document_defaults(self.doc)

        # Extract abstract and keywords
        self._extract_abstract_and_keywords()

        # Create title page
        create_title_page(self.doc)

        # Create abstract page
        create_abstract_page(self.doc, self.abstract_text, self.keywords_text)

        # Re-add the title at the start of the body (APA 7th requirement)
        _add_formatted_paragraph(
            self.doc, TITLE,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )

        # Parse body content
        self._parse_body()

        # Save
        print(f"\nSaving: {self.output_path}")
        self.doc.save(str(self.output_path))
        print("Done!")

    def _parse_body(self):
        """Parse the markdown body and add content to the document."""
        i = 0
        total = len(self.lines)
        skip_frontmatter = True  # Skip title/author/date/status lines at top
        in_abstract_section = False
        paragraph_buffer = []
        in_table = False

        while i < total:
            line = self.lines[i]
            stripped = line.strip()

            # --- Skip YAML-like front matter at the top ---
            if skip_frontmatter:
                if stripped.startswith("# "):
                    # This is the main title - skip it (we have title page)
                    i += 1
                    continue
                if stripped.startswith("**") and ":" in stripped:
                    # Metadata lines like **Authors:**, **Date:**, etc.
                    i += 1
                    continue
                if stripped == "" or stripped == "---":
                    if stripped == "---":
                        skip_frontmatter = False
                    i += 1
                    continue
                skip_frontmatter = False

            # --- Section break (---) ---
            if stripped == "---":
                # Flush any pending paragraph
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                # Flush any pending table
                if in_table:
                    self._flush_table()
                    in_table = False
                i += 1
                continue

            # --- Abstract section: skip (already handled) ---
            if stripped.lower() == "## abstract":
                in_abstract_section = True
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                i += 1
                continue

            if in_abstract_section:
                if stripped.startswith("## ") or stripped == "---":
                    in_abstract_section = False
                    # Don't increment i; re-process this line
                    continue
                i += 1
                continue

            # --- Markdown table detection ---
            if stripped.startswith("|") and "|" in stripped[1:]:
                # Flush paragraph buffer first
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                self.current_table_lines.append(stripped)
                in_table = True
                i += 1
                continue
            elif in_table:
                self._flush_table()
                in_table = False
                # Don't increment; re-process this line
                continue

            # --- Headings ---
            heading_match = re.match(r'^(#{2,6})\s+(.+)$', stripped)
            if heading_match:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []

                level_markers = heading_match.group(1)
                heading_raw = heading_match.group(2).strip()
                heading_text = self._clean_heading_text(heading_raw)
                level = len(level_markers)  # ## = 2, ### = 3, etc.

                # Map markdown heading levels to APA heading levels
                # ## -> Level 1, ### -> Level 2, #### -> Level 3,
                # ##### -> Level 4, ###### -> Level 5
                apa_level = level - 1  # ## (2) -> APA 1, ### (3) -> APA 2, etc.

                if apa_level == 1:
                    page_break = self._should_page_break(heading_text)
                    # Special case: "References" section
                    if heading_text.lower().startswith("references"):
                        self.in_references = True
                    elif heading_text.lower().startswith("appendix"):
                        self.in_references = False
                    add_heading_level1(self.doc, heading_text, page_break=page_break)
                    section_name = heading_text.split(":")[0] if ":" in heading_text else heading_text
                    print(f"  Processing section: {section_name}")
                elif apa_level == 2:
                    add_heading_level2(self.doc, heading_text)
                elif apa_level == 3:
                    add_heading_level3(self.doc, heading_text)
                elif apa_level == 4:
                    # Level 4 heading - text continues on same line
                    # Collect following paragraph text
                    level4_para = add_heading_level4(self.doc, heading_text)
                    # Look ahead for continuation text
                    j = i + 1
                    continuation_lines = []
                    while j < total:
                        next_line = self.lines[j].strip()
                        if next_line == "" or next_line.startswith("#") or next_line == "---":
                            break
                        continuation_lines.append(next_line)
                        j += 1
                    if continuation_lines:
                        continuation_text = " ".join(continuation_lines)
                        segments = _parse_inline(continuation_text)
                        _add_runs_from_segments(level4_para, segments)
                        i = j
                        continue
                elif apa_level >= 5:
                    add_heading_level5(self.doc, heading_text)

                i += 1
                continue

            # --- Figure placeholders ---
            fig_num = self._is_figure_placeholder(stripped)
            if fig_num is not None:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                # Try to extract a caption from the bracket text
                caption_match = re.search(r'\[(?:INSERT\s+)?(?:Figure\s+\d+)[:\s—–-]*(.+?)\]',
                                          stripped, re.IGNORECASE)
                caption = caption_match.group(1).strip(" —–-]") if caption_match else None
                if caption and caption.upper() in ("HERE", "TO BE INSERTED"):
                    caption = None
                add_figure_placeholder(self.doc, fig_num, caption)
                i += 1
                continue

            # --- Table placeholders ---
            tbl_num = self._is_table_placeholder(stripped)
            if tbl_num is not None:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                # Extract placeholder text for note
                caption_match = re.search(r'\[(?:INSERT\s+)?(?:Table\s+\d+)[:\s—–-]*(.+?)\]',
                                          stripped, re.IGNORECASE)
                caption = caption_match.group(1).strip(" —–-]") if caption_match else None
                if caption and caption.upper() in ("HERE", "TO BE INSERTED"):
                    caption = f"[Table {tbl_num} to be inserted]"
                add_apa_table(
                    self.doc,
                    table_number=tbl_num,
                    title=caption or f"[Table {tbl_num} title to be added]",
                    headers=["Column 1", "Column 2", "Column 3"],
                    rows=[["[Data]", "[Data]", "[Data]"]],
                    notes="[Table notes to be added]",
                )
                i += 1
                continue

            # --- Block quotes ---
            if stripped.startswith(">"):
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                quote_text = stripped.lstrip("> ").strip()
                # Collect multi-line block quote
                j = i + 1
                while j < total:
                    next_line = self.lines[j].strip()
                    if next_line.startswith(">"):
                        quote_text += " " + next_line.lstrip("> ").strip()
                        j += 1
                    else:
                        break
                add_block_quote(self.doc, quote_text)
                i = j
                continue

            # --- Display math ($$...$$) on its own line ---
            if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                math_text = stripped[2:-2].strip()
                _add_formatted_paragraph(
                    self.doc, math_text,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    italic=True,
                )
                i += 1
                continue

            # --- Multi-line display math ---
            if stripped == "$$":
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                math_lines = []
                j = i + 1
                while j < total:
                    ml = self.lines[j].strip()
                    if ml == "$$":
                        j += 1
                        break
                    math_lines.append(ml)
                    j += 1
                math_text = " ".join(math_lines)
                _add_formatted_paragraph(
                    self.doc, math_text,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    italic=True,
                )
                i = j
                continue

            # --- Empty line: flush paragraph buffer ---
            if stripped == "":
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []
                i += 1
                continue

            # --- References section: each non-empty line after blank is a reference ---
            if self.in_references and stripped:
                self._flush_paragraph_buffer(paragraph_buffer)
                paragraph_buffer = []

                # Skip lines like "[FULL LIST TBD ...]"
                if stripped.startswith("[") and stripped.endswith("]"):
                    i += 1
                    continue

                # Collect the full reference (may span multiple lines)
                ref_lines = [stripped]
                j = i + 1
                while j < total:
                    next_line = self.lines[j].strip()
                    if next_line == "" or next_line.startswith("#") or next_line == "---":
                        break
                    ref_lines.append(next_line)
                    j += 1
                ref_text = " ".join(ref_lines)
                add_reference_entry(self.doc, ref_text)
                i = j
                continue

            # --- Regular paragraph text ---
            # Accumulate lines until blank line (multi-line paragraph)
            paragraph_buffer.append(stripped)
            i += 1

        # Flush any remaining buffer
        self._flush_paragraph_buffer(paragraph_buffer)
        if in_table:
            self._flush_table()

    def _flush_paragraph_buffer(self, buffer):
        """Flush accumulated paragraph lines as a single paragraph."""
        if not buffer:
            return

        text = " ".join(buffer)
        buffer.clear()

        if not text.strip():
            return

        self._process_body_paragraph(text)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown manuscript to APA 7th Word document."
    )
    parser.add_argument(
        "--input", "-i",
        default=None,
        help="Input markdown file (default: draft_v2.md in script directory)",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Output .docx file (default: Paper4_APA7th_v2.docx in script directory)",
    )
    args = parser.parse_args()

    # Resolve paths
    if args.input:
        input_path = Path(args.input).resolve()
    else:
        input_path = SCRIPT_DIR / "draft_v2.md"

    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = SCRIPT_DIR / "Paper4_APA7th_v2.docx"

    # Validate input
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        print("  Make sure draft_v2.md exists in the manuscript directory,")
        print("  or specify an input file with --input.")

        # Fallback: try draft_v1.md
        fallback = SCRIPT_DIR / "draft_v1.md"
        if fallback.exists():
            print(f"\n  Found fallback: {fallback.name}")
            response = input("  Use draft_v1.md instead? [y/N]: ").strip().lower()
            if response == "y":
                input_path = fallback
            else:
                sys.exit(1)
        else:
            sys.exit(1)

    # Check for python-docx
    try:
        import docx
        print(f"python-docx version: {docx.__version__}")
    except ImportError:
        print("Error: python-docx is not installed.")
        print("  Install with: pip install python-docx")
        sys.exit(1)

    # Check figures directory
    if FIGURES_DIR.exists():
        fig_files = list(FIGURES_DIR.glob("*.png"))
        print(f"Figures directory: {FIGURES_DIR}")
        print(f"  Found {len(fig_files)} PNG files.")
    else:
        print(f"Warning: Figures directory not found: {FIGURES_DIR}")
        print("  Figure placeholders will be used instead of images.")

    # Convert
    converter = MarkdownToAPAConverter(input_path, output_path)
    converter.convert()

    # Summary
    print(f"\n=== Summary ===")
    print(f"  Input:  {input_path}")
    print(f"  Output: {output_path}")
    print(f"  Size:   {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
