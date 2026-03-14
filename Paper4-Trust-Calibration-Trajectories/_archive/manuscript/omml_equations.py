#!/usr/bin/env python3
"""
omml_equations.py
=================
Convert LaTeX math expressions to OMML (Office Math Markup Language) XML
for native Word equations via python-docx.

Supports: fractions, subscripts, superscripts, \\text{}, \\hat{},
          \\sum with limits, \\log, Greek letters, parentheses, operators.

Usage:
    from omml_equations import add_display_equation
    p = doc.add_paragraph()
    add_display_equation(p, r"R_b(\\tau) = \\frac{a}{b}")

Author: Hosung You
Date: March 2026
"""

from lxml import etree

# ---------------------------------------------------------------------------
# Namespaces
# ---------------------------------------------------------------------------
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"

NSMAP = {"m": MATH_NS, "w": WORD_NS}


def _m(tag):
    """Element name in math namespace."""
    return f"{{{MATH_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Greek & Symbol Maps
# ---------------------------------------------------------------------------
GREEK = {
    "alpha": "\u03B1", "beta": "\u03B2", "gamma": "\u03B3", "delta": "\u03B4",
    "epsilon": "\u03B5", "zeta": "\u03B6", "eta": "\u03B7", "theta": "\u03B8",
    "iota": "\u03B9", "kappa": "\u03BA", "lambda": "\u03BB", "mu": "\u03BC",
    "nu": "\u03BD", "xi": "\u03BE", "pi": "\u03C0", "rho": "\u03C1",
    "sigma": "\u03C3", "tau": "\u03C4", "upsilon": "\u03C5", "phi": "\u03C6",
    "chi": "\u03C7", "psi": "\u03C8", "omega": "\u03C9",
    "Alpha": "\u0391", "Beta": "\u0392", "Gamma": "\u0393", "Delta": "\u0394",
    "Epsilon": "\u0395", "Zeta": "\u0396", "Eta": "\u0397", "Theta": "\u0398",
    "Lambda": "\u039B", "Xi": "\u039E", "Pi": "\u03A0", "Sigma": "\u03A3",
    "Phi": "\u03A6", "Psi": "\u03A8", "Omega": "\u03A9",
}

SYMBOLS = {
    "times": "\u00D7", "cdot": "\u00B7", "pm": "\u00B1", "mp": "\u2213",
    "leq": "\u2264", "geq": "\u2265", "neq": "\u2260", "approx": "\u2248",
    "infty": "\u221E", "partial": "\u2202", "ldots": "\u2026",
}

NARY_OPS = {"sum": "\u2211", "prod": "\u220F", "int": "\u222B"}

FUNCTIONS = {"log", "ln", "exp", "sin", "cos", "tan", "min", "max", "lim", "arg"}


# ---------------------------------------------------------------------------
# Tokenizer
# ---------------------------------------------------------------------------
class _Tokenizer:
    """Tokenize a LaTeX math string into (type, value) pairs."""

    def __init__(self, text):
        self.tokens = []
        self._tokenize(text)
        self.idx = 0

    # ---- tokenizer core ----
    def _tokenize(self, s):
        i = 0
        n = len(s)
        while i < n:
            ch = s[i]

            # --- backslash commands ---
            if ch == "\\":
                j = i + 1
                if j < n and not s[j].isalpha():
                    # escaped char: \_ \{ \} \\ etc.
                    self.tokens.append(("CHAR", s[j]))
                    i = j + 1
                    continue
                while j < n and s[j].isalpha():
                    j += 1
                cmd = s[i + 1 : j]

                # Special: \text{...} -> single TEXT token
                if cmd == "text":
                    k = j
                    while k < n and s[k] == " ":
                        k += 1
                    if k < n and s[k] == "{":
                        content, end = self._extract_braced(s, k)
                        content = content.replace("\\_", "_")
                        self.tokens.append(("TEXT", content))
                        i = end
                        continue

                self.tokens.append(("CMD", cmd))
                i = j

            elif ch == "{":
                self.tokens.append(("LBRACE", ch)); i += 1
            elif ch == "}":
                self.tokens.append(("RBRACE", ch)); i += 1
            elif ch == "_":
                self.tokens.append(("SUB", ch)); i += 1
            elif ch == "^":
                self.tokens.append(("SUP", ch)); i += 1
            elif ch == "(":
                self.tokens.append(("LPAREN", ch)); i += 1
            elif ch == ")":
                self.tokens.append(("RPAREN", ch)); i += 1
            elif ch in "+-=<>":
                self.tokens.append(("OP", ch)); i += 1
            elif ch == " ":
                i += 1  # skip whitespace in math
            elif ch.isdigit():
                j = i
                while j < n and (s[j].isdigit() or s[j] == "."):
                    j += 1
                self.tokens.append(("NUM", s[i:j])); i = j
            elif ch.isalpha():
                self.tokens.append(("LETTER", ch)); i += 1
            elif ch == ",":
                self.tokens.append(("COMMA", ch)); i += 1
            else:
                self.tokens.append(("CHAR", ch)); i += 1

    @staticmethod
    def _extract_braced(s, start):
        """Extract content of {...} starting at `start` (pointing to '{').
        Returns (content_str, index_after_closing_brace)."""
        depth = 1
        k = start + 1
        n = len(s)
        while k < n and depth > 0:
            if s[k] == "\\":
                k += 2  # skip escape
            elif s[k] == "{":
                depth += 1; k += 1
            elif s[k] == "}":
                depth -= 1; k += 1
            else:
                k += 1
        return s[start + 1 : k - 1], k

    # ---- iterator interface ----
    def peek(self):
        return self.tokens[self.idx] if self.idx < len(self.tokens) else None

    def consume(self):
        tok = self.tokens[self.idx]
        self.idx += 1
        return tok

    def has_more(self):
        return self.idx < len(self.tokens)


# ---------------------------------------------------------------------------
# Parser / OMML Builder
# ---------------------------------------------------------------------------
class _Builder:
    """Recursive-descent parser: LaTeX tokens → OMML lxml elements."""

    # ---- public entry points ----

    def display(self, latex_str):
        """Return an <m:oMathPara> element for a display equation."""
        tok = _Tokenizer(latex_str)
        para = etree.Element(_m("oMathPara"), nsmap=NSMAP)
        omath = etree.SubElement(para, _m("oMath"))
        self._expr(tok, omath)
        return para

    def inline(self, latex_str):
        """Return an <m:oMath> element for an inline equation."""
        tok = _Tokenizer(latex_str)
        omath = etree.Element(_m("oMath"), nsmap=NSMAP)
        self._expr(tok, omath)
        return omath

    # ---- grammar rules ----

    def _expr(self, tok, parent):
        """Parse elements until end, RBRACE, or RPAREN."""
        while tok.has_more():
            t = tok.peek()
            if t is None or t[0] in ("RBRACE", "RPAREN"):
                break
            self._element(tok, parent)

    def _element(self, tok, parent):
        t = tok.peek()
        if t is None:
            return
        tt, tv = t

        if tt == "CMD":
            self._cmd(tok, parent, tv)
        elif tt == "TEXT":
            tok.consume()
            run = self._run(tv, plain=True)
            self._wrap_scripts(tok, parent, run)
        elif tt == "LETTER":
            tok.consume()
            run = self._run(tv, italic=True)
            self._wrap_scripts(tok, parent, run)
        elif tt == "NUM":
            tok.consume()
            run = self._run(tv)
            self._wrap_scripts(tok, parent, run)
        elif tt == "OP":
            tok.consume()
            # Use Unicode minus for '-'
            ch = "\u2212" if tv == "-" else tv
            parent.append(self._run(f" {ch} "))
        elif tt == "COMMA":
            tok.consume()
            parent.append(self._run(", "))
        elif tt == "CHAR":
            tok.consume()
            parent.append(self._run(tv))
        elif tt == "LPAREN":
            d = self._delimited(tok)
            self._wrap_scripts(tok, parent, d)
        elif tt == "LBRACE":
            tok.consume()
            self._expr(tok, parent)
            if tok.has_more() and tok.peek()[0] == "RBRACE":
                tok.consume()
        elif tt in ("SUB", "SUP"):
            # orphan script
            self._wrap_scripts(tok, parent, self._run(""))
        else:
            tok.consume()

    def _cmd(self, tok, parent, cmd):
        if cmd == "frac":
            self._frac(tok, parent)
        elif cmd == "hat":
            elem = self._accent(tok)
            self._wrap_scripts(tok, parent, elem)
        elif cmd in NARY_OPS:
            self._nary(tok, parent)
        elif cmd in FUNCTIONS:
            self._func(tok, parent)
        elif cmd in GREEK:
            tok.consume()
            run = self._run(GREEK[cmd], italic=True)
            self._wrap_scripts(tok, parent, run)
        elif cmd in SYMBOLS:
            tok.consume()
            parent.append(self._run(SYMBOLS[cmd]))
        elif cmd == "left" or cmd == "right":
            tok.consume()  # skip \left / \right
            if tok.has_more():
                t2 = tok.peek()
                if t2 and t2[0] in ("LPAREN", "RPAREN", "CHAR"):
                    tok.consume()  # skip the delimiter char
        else:
            tok.consume()
            parent.append(self._run(cmd))

    # ---- compound structures ----

    def _frac(self, tok, parent):
        tok.consume()  # 'frac'
        f = etree.SubElement(parent, _m("f"))
        num = etree.SubElement(f, _m("num"))
        self._braced(tok, num)
        den = etree.SubElement(f, _m("den"))
        self._braced(tok, den)

    def _accent(self, tok):
        tok.consume()  # 'hat'
        acc = etree.Element(_m("acc"))
        pr = etree.SubElement(acc, _m("accPr"))
        ch = etree.SubElement(pr, _m("chr"))
        ch.set(_m("val"), "\u0302")
        e = etree.SubElement(acc, _m("e"))
        self._braced(tok, e)
        return acc

    def _nary(self, tok, parent):
        cmd = tok.consume()[1]
        char = NARY_OPS[cmd]

        nary = etree.SubElement(parent, _m("nary"))
        pr = etree.SubElement(nary, _m("naryPr"))
        ch = etree.SubElement(pr, _m("chr"))
        ch.set(_m("val"), char)
        # limLoc: subSup (limits beside, not above/below) for inline feel
        lim = etree.SubElement(pr, _m("limLoc"))
        lim.set(_m("val"), "undOvr")

        sub = etree.SubElement(nary, _m("sub"))
        sup = etree.SubElement(nary, _m("sup"))
        e = etree.SubElement(nary, _m("e"))

        if tok.has_more() and tok.peek()[0] == "SUB":
            tok.consume()
            self._braced(tok, sub)
        if tok.has_more() and tok.peek()[0] == "SUP":
            tok.consume()
            self._braced(tok, sup)

        # body: consume all remaining in current scope
        self._expr(tok, e)

    def _func(self, tok, parent):
        cmd = tok.consume()[1]
        func = etree.SubElement(parent, _m("func"))
        etree.SubElement(func, _m("funcPr"))
        fn = etree.SubElement(func, _m("fName"))
        fn.append(self._run(cmd, plain=True))
        e = etree.SubElement(func, _m("e"))

        if tok.has_more() and tok.peek()[0] == "LPAREN":
            d = self._delimited(tok)
            e.append(d)
        elif tok.has_more():
            self._element(tok, e)

    def _delimited(self, tok):
        tok.consume()  # LPAREN
        d = etree.Element(_m("d"))
        pr = etree.SubElement(d, _m("dPr"))
        bc = etree.SubElement(pr, _m("begChr"))
        bc.set(_m("val"), "(")
        ec = etree.SubElement(pr, _m("endChr"))
        ec.set(_m("val"), ")")
        e = etree.SubElement(d, _m("e"))
        while tok.has_more():
            if tok.peek()[0] == "RPAREN":
                tok.consume()
                break
            self._element(tok, e)
        return d

    # ---- helpers ----

    def _braced(self, tok, parent):
        """Parse {group} or single element."""
        if tok.has_more() and tok.peek()[0] == "LBRACE":
            tok.consume()
            self._expr(tok, parent)
            if tok.has_more() and tok.peek()[0] == "RBRACE":
                tok.consume()
        elif tok.has_more():
            self._element(tok, parent)

    def _wrap_scripts(self, tok, parent, base):
        """Attach sub/superscripts to base if present."""
        has_sub = has_sup = False
        sub_el = sup_el = None

        if tok.has_more() and tok.peek()[0] == "SUB":
            has_sub = True
            tok.consume()
            sub_el = etree.Element(_m("sub"))
            self._braced(tok, sub_el)

        if tok.has_more() and tok.peek()[0] == "SUP":
            has_sup = True
            tok.consume()
            sup_el = etree.Element(_m("sup"))
            self._braced(tok, sup_el)

        if has_sub and has_sup:
            w = etree.SubElement(parent, _m("sSubSup"))
            e = etree.SubElement(w, _m("e"))
            e.append(base)
            w.append(sub_el)
            w.append(sup_el)
        elif has_sub:
            w = etree.SubElement(parent, _m("sSub"))
            e = etree.SubElement(w, _m("e"))
            e.append(base)
            w.append(sub_el)
        elif has_sup:
            w = etree.SubElement(parent, _m("sSup"))
            e = etree.SubElement(w, _m("e"))
            e.append(base)
            w.append(sup_el)
        else:
            parent.append(base)

    @staticmethod
    def _run(text, italic=False, plain=False):
        """Create <m:r> with text and optional style."""
        r = etree.Element(_m("r"))
        rpr = etree.SubElement(r, _m("rPr"))
        sty = etree.SubElement(rpr, _m("sty"))
        if plain:
            sty.set(_m("val"), "p")
        elif italic:
            sty.set(_m("val"), "i")
        else:
            sty.set(_m("val"), "p")  # numbers/operators plain
        t = etree.SubElement(r, _m("t"))
        t.text = text
        t.set(f"{{{XML_NS}}}space", "preserve")
        return r


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
_builder = _Builder()


def latex_to_omml_display(latex_str):
    """Convert LaTeX to an <m:oMathPara> element (display equation)."""
    return _builder.display(latex_str)


def latex_to_omml_inline(latex_str):
    """Convert LaTeX to an <m:oMath> element (inline equation)."""
    return _builder.inline(latex_str)


def add_display_equation(paragraph, latex_str):
    """Insert a native Word display equation into a python-docx paragraph.

    Args:
        paragraph: A python-docx Paragraph object.
        latex_str: LaTeX math string (without $$ delimiters).
    """
    omml = latex_to_omml_display(latex_str)
    paragraph._element.append(omml)


def add_inline_equation(paragraph, latex_str):
    """Insert a native Word inline equation into a python-docx paragraph.

    Args:
        paragraph: A python-docx Paragraph object.
        latex_str: LaTeX math string (without $ delimiters).
    """
    omml = latex_to_omml_inline(latex_str)
    paragraph._element.append(omml)


# ---------------------------------------------------------------------------
# CLI test
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    test_cases = [
        r"R_b(\tau) = \frac{\text{Number of adaptive\_offer episodes in window } \tau}{\text{Total number of episodes in window } \tau}",
        r"P(\tau) = \frac{\text{Number of correct responses in window } \tau}{\text{Total number of question-answering episodes in window } \tau}",
        r"\text{Gap}(\tau) = R_b(\tau) - P(\tau)",
        r"P_{\text{adaptive}}(\tau) = \frac{\text{Correct responses on adaptive questions in window } \tau}{\text{Total adaptive questions in window } \tau}",
        r"P_{\text{non-adaptive}}(\tau) = \frac{\text{Correct responses on self-selected questions in window } \tau}{\text{Total self-selected questions in window } \tau}",
        r"\text{AI\_benefit}(\tau) = P_{\text{adaptive}}(\tau) - P_{\text{non-adaptive}}(\tau)",
        r"\text{Gap}_{\text{new}}(\tau) = R_b(\tau) - P_{\text{adaptive}}(\tau)",
        r"E = 1 - \frac{-\sum_{i=1}^{N}\sum_{k=1}^{K} \hat{p}_{ik} \log(\hat{p}_{ik})}{N \log(K)}",
    ]

    print("Testing LaTeX → OMML conversion")
    print("=" * 60)
    for i, tex in enumerate(test_cases, 1):
        omml = latex_to_omml_display(tex)
        xml_str = etree.tostring(omml, pretty_print=True).decode()
        print(f"\nEquation {i}:")
        print(f"  LaTeX: {tex[:80]}...")
        print(f"  OMML elements: {len(list(omml.iter()))} nodes")
        # Print first few lines of XML
        for line in xml_str.split("\n")[:5]:
            print(f"    {line}")
        print("    ...")

    # Quick test: create a Word doc with all equations
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        for i, tex in enumerate(test_cases, 1):
            # Label
            p = doc.add_paragraph()
            p.add_run(f"Equation ({i}):")
            p.paragraph_format.space_after = Pt(0)

            # Equation
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            add_display_equation(p, tex)

            # Spacer
            doc.add_paragraph()

        out_path = "test_omml_equations.docx"
        doc.save(out_path)
        print(f"\nTest document saved: {out_path}")
    except ImportError:
        print("\npython-docx not available; skipping Word test.")
