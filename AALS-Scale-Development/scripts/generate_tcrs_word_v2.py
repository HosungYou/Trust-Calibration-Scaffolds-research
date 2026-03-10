#!/usr/bin/env python3
"""
Generate APA 7th Edition Word document for TCRS manuscript.
Includes proper Tables and Figure embedded in body text.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
MANUSCRIPT_DIR = os.path.join(PROJECT_DIR, "manuscript")
FIGURES_DIR = os.path.join(PROJECT_DIR, "figures")
OUTPUT_PATH = os.path.join(MANUSCRIPT_DIR, "TCRS_Theoretical_Background_APA7.docx")
FIGURE_FILE = os.path.join(FIGURES_DIR, "Figure_1_Trust_Calibration_Process_Model_v3.png")


# ─── APA 7 Formatting Helpers ───

def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 4, "val": "single"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def format_apa_table(table):
    """Apply APA 7 formatting: no vertical borders, horizontal lines at top, header bottom, table bottom."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders first
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 0, "val": "none"},
                bottom={"sz": 0, "val": "none"},
                left={"sz": 0, "val": "none"},
                right={"sz": 0, "val": "none"})

    # Top border on header row
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"sz": 8, "val": "single"}, bottom={"sz": 4, "val": "single"})

    # Bottom border on last row
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 8, "val": "single"})


def set_cell_text(cell, text, bold=False, italic=False, size=10, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def add_run(paragraph, text, bold=False, italic=False, size=12, superscript=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    return run


def add_paragraph(doc, first_line_indent=True):
    """Create an APA-style body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    return p


def add_body_text(doc, text, first_line_indent=True):
    """Add body text with inline formatting."""
    p = add_paragraph(doc, first_line_indent)
    # Simple inline formatting: handle *italic* and **bold**
    import re
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*'):
            add_run(p, part[1:-1], italic=True)
        else:
            # Handle em-dash and en-dash
            part = part.replace('---', '\u2014').replace('--', '\u2013')
            add_run(p, part)
    return p


def add_heading(doc, text, level=1):
    """Add APA-style heading."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 2.0

    if level == 1:  # Centered, Bold
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True)
    elif level == 2:  # Flush Left, Bold
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=True)
    elif level == 3:  # Flush Left, Bold Italic
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=True, italic=True)
    elif level == 4:  # Indented, Bold, Period.
        pf.first_line_indent = Inches(0.5)
        add_run(p, text + ".", bold=True)
    return p


def add_bullet(doc, text, indent_level=0):
    """Add bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.5 + indent_level * 0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    add_run(p, "\u2022 ")
    import re
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*'):
            add_run(p, part[1:-1], italic=True)
        else:
            add_run(p, part)
    return p


def add_numbered(doc, number, text):
    """Add numbered list item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    add_run(p, f"{number}. ")
    import re
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*'):
            add_run(p, part[1:-1], italic=True)
        else:
            add_run(p, part)
    return p


def add_table_title(doc, number, title):
    """Add APA 7 table title: 'Table N' in bold, title in italic on next line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    add_run(p, f"Table {number}", bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing = 2.0
    add_run(p2, title, italic=True)
    return p2


def add_table_note(doc, note_text):
    """Add APA 7 table note."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    add_run(p, "Note. ", italic=True, size=10)
    add_run(p, note_text, size=10)
    return p


def add_figure_apa(doc, image_path, fig_num, caption):
    """Add APA 7 figure: label, caption, then image."""
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_before = Pt(12)
    p_label.paragraph_format.space_after = Pt(0)
    p_label.paragraph_format.line_spacing = 2.0
    add_run(p_label, f"Figure {fig_num}", bold=True, italic=True)

    p_caption = doc.add_paragraph()
    p_caption.paragraph_format.space_before = Pt(0)
    p_caption.paragraph_format.space_after = Pt(6)
    p_caption.paragraph_format.line_spacing = 2.0
    import re
    parts = re.split(r'(\*.*?\*)', caption)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            add_run(p_caption, part[1:-1], italic=True)
        else:
            add_run(p_caption, part)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(12)
    run = p_img.add_run()
    run.add_picture(image_path, width=Inches(6.5))


# ─── Table Generators ───

def add_table_selective_borrowing(doc):
    """Table 1: Selective Theoretical Borrowing"""
    add_table_title(doc, 1, "Selective Theoretical Borrowing: Scope of Mechanism Extraction")

    headers = ["Target Construct Function", "Borrowed Mechanism (Source)",
               "What Is NOT Borrowed", "Borrowing Rationale"]
    data = [
        ["Monitoring capacity",
         "Self-monitoring of trust state (Winne & Hadwin, 1998, Phases 1\u20133)",
         "Knowledge of cognition, metacognitive regulation strategies, SRL Phase 4 adaptation",
         "Trust calibration begins with self-monitoring; broader SRL processes are outside scope"],
        ["Evaluative capacity",
         "Trust-reliability matching assessment (Lee & See, 2004)",
         "Trust formation, trust repair, dispositional trust, organizational trust antecedents",
         "Calibration evaluates already-formed trust; how trust forms or repairs is beyond scope"],
        ["Regulatory capacity",
         "Intentional verification and adjustment behaviors (Bandura, 2001)",
         "Intentionality, forethought, collective agency, proxy agency, moral agency",
         "Calibration requires behavioral adjustment; broader agentic properties are not measured"],
    ]

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(1.3), Inches(1.7), Inches(1.7), Inches(1.6)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            set_cell_text(table.rows[i + 1].cells[j], cell_text, size=10)

    format_apa_table(table)
    add_table_note(doc,
        "Each theory serves as an explanatory resource for one functional capacity, "
        "not as a structural pillar of the model. See Davis (1989), Ajzen (1991), and "
        "Lazarus & Folkman (1984) for precedent models using similar functional decomposition.")


def add_table_2(doc):
    """Table 2: Comparison of Existing AI Literacy Scales"""
    add_table_title(doc, 2, "Comparison of Existing AI Literacy Scales and Trust Calibration Coverage")

    headers = ["Scale", "Authors (Year)", "Items", "Key Dimensions", "Trust Calibration\nCoverage"]
    data = [
        ["SNAIL", "Laupichler et al. (2023)", "31", "Technical understanding,\ncritical appraisal, practical application", "None"],
        ["MAILS", "Carolus et al. (2023)", "27", "AI knowledge, meta-competencies\n(persuasion, emotion regulation)", "Partial\n(meta-competencies)"],
        ["MAILS-Short", "Carolus et al. (2024)", "10", "Abbreviated MAILS", "Partial"],
        ["AILQ", "Ng et al. (2024)", "~30", "Affective, Behavioural,\nCognitive, Ethical", "None"],
        ["SAIL4ALL", "Pinski et al. (2025)", "Varies", "Performance-based\nAI knowledge", "None"],
        ["GAIL Scales", "Multiple (2024\u20132025)", "Varies", "Awareness, usage, evaluation,\nethics of GenAI", "Partial\n(evaluation)"],
        ["Collaborative\nAI Metacognition", "2025", "Varies", "Planning, monitoring,\nevaluating in human-AI interaction", "Partial\n(monitoring)"],
        ["TCRS\n(Proposed)", "Present study", "48", "Calibration Awareness,\nCalibration Judgment,\nCalibration Action", "Full coverage:\ntrust state awareness,\nevaluation, regulation"],
    ]

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    widths = [Inches(1.1), Inches(1.2), Inches(0.5), Inches(2.0), Inches(1.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            bold = (i == len(data) - 1)  # Bold last row (TCRS)
            set_cell_text(table.rows[i + 1].cells[j], cell_text, bold=bold, size=10)

    format_apa_table(table)
    add_table_note(doc,
        "SNAIL = Scale for the Assessment of Non-experts\u2019 AI Literacy; "
        "MAILS = Meta AI Literacy Scale; AILQ = AI Literacy Questionnaire; "
        "SAIL4ALL = Scale of AI Literacy for All; GAIL = Generative AI Literacy; "
        "TCRS = Trust Calibration Readiness Scale. "
        "\"Full coverage\" indicates that the scale addresses all three elements required for trust calibration: "
        "(a) awareness of one\u2019s own trust state, (b) evaluation of trust appropriateness, "
        "and (c) behavioral regulation of trust.")


def add_table_3(doc):
    """Table 3: TCRS Subscale Structure and Theoretical Foundations"""
    add_table_title(doc, 3, "TCRS Subscale Structure, Functional Capacities, and Explanatory Resources")

    headers = ["Subscale", "Functional\nCapacity", "Explanatory\nResource", "Facets", "Items", "Response\nFormat"]
    data = [
        ["Calibration\nAwareness\n(CA-Aw)",
         "Monitoring:\n\u201CWhere am I?\u201D",
         "Self-monitoring in SRL\n(Winne & Hadwin, 1998,\nPhases 1\u20133; Efklides, 2006)",
         "Trust state awareness\nAI system understanding\nTask context recognition\nAffective monitoring\nTemporal drift awareness\nSocial awareness",
         "16",
         "7-point\nagreement\n(1 = Strongly\nDisagree to\n7 = Strongly\nAgree)"],
        ["Calibration\nJudgment\n(CA-Jd)",
         "Evaluative:\n\u201CIs my trust\nappropriate?\u201D",
         "Trust dynamics\n(Lee & See, 2004);\nEpistemic vigilance\n(Sperber et al., 2010)",
         "Appropriateness evaluation\nSituational comparison\nDiscrepancy detection\nTemporal drift judgment\nProportional trust\nOvergeneralization detection",
         "16",
         "7-point\nagreement\n(1 = Strongly\nDisagree to\n7 = Strongly\nAgree)"],
        ["Calibration\nAction\n(CA-Ac)",
         "Regulatory:\n\u201CAdjust toward\nappropriate\u201D",
         "Self-regulation\n(Bandura, 2001,\nself-reactiveness);\nSRL Phase 4\n(Winne & Hadwin, 1998)",
         "Verification behaviors\nTrust adjustment\nStrategic regulation\nSocial verification\nProactive calibration\nMeta-strategic regulation",
         "16",
         "7-point\nfrequency\n(1 = Never\nto 7 = Always)"],
    ]

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(1.0), Inches(1.0), Inches(1.3), Inches(1.5), Inches(0.5), Inches(1.0)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            set_cell_text(table.rows[i + 1].cells[j], cell_text, size=10)

    format_apa_table(table)
    add_table_note(doc,
        "SRL = Self-Regulated Learning. Source theories serve as explanatory resources for each "
        "functional capacity, not as constructs being measured in their entirety (see Table 1). "
        "The fourth process component (Calibration "
        "Outcome) is not directly measured because the TCRS assesses readiness rather than accuracy. "
        "Calibration Awareness and Calibration Judgment use an agreement response format because "
        "they assess self-perceptions; Calibration Action uses a frequency format because it "
        "assesses behavioral occurrence (Krosnick, 1999).")


def add_table_4(doc):
    """Table 4: Full 48-Item Pool"""
    add_table_title(doc, 4, "Trust Calibration Readiness Scale (TCRS): Full 48-Item Pool")

    headers = ["Item", "Facet", "Item Text", "Dir."]

    # --- CA-Aw items ---
    aw_items = [
        ["Aw01", "Trust state", "When I use an AI tool for a learning task, I am aware of how much I trust its output.", "+"],
        ["Aw02", "Trust state", "I notice when I am relying too heavily on AI-generated content.", "+"],
        ["Aw03", "Trust state", "I rarely think about whether I am trusting AI too much or too little.", "R"],
        ["Aw04", "Trust state", "I notice when my level of trust in an AI tool changes during a task.", "+"],
        ["Aw05", "Temporal drift", "I notice shifts in how much I trust an AI tool as I use it over time.", "+"],
        ["Aw06", "Affective", "I notice my emotional reactions (e.g., surprise, unease) when I receive an AI-generated response.", "+"],
        ["Aw07", "System understanding", "Before trusting an AI tool\u2019s output, I pay attention to what the AI was designed to do.", "+"],
        ["Aw08", "System understanding", "I think about how an AI tool produces its responses when I use it for learning.", "+"],
        ["Aw09", "System understanding", "I don\u2019t usually think about the limitations of an AI tool when using its suggestions.", "R"],
        ["Aw10", "Task context", "Before using AI for a task, I pause to think about the situation.", "+"],
        ["Aw11", "Task context", "I am mindful that my trust in AI may need to vary across different situations.", "+"],
        ["Aw12", "Task context", "I treat AI outputs the same regardless of the type of task.", "R"],
        ["Aw13", "Task context", "When using AI for learning, I pay attention to how important accuracy is for the specific task.", "+"],
        ["Aw14", "Task context", "I am aware that AI tools perform differently depending on the subject area.", "+"],
        ["Aw15", "Temporal drift", "I am aware that my trust in an AI tool may change the more I use it, without my noticing.", "+"],
        ["Aw16", "Social", "I notice whether other people around me (classmates, instructors) seem to trust AI more or less than I do.", "+"],
    ]

    jd_items = [
        ["Jd01", "Appropriateness", "When using AI for a task, I evaluate whether my level of trust matches what the situation requires.", "+"],
        ["Jd02", "Appropriateness", "I weigh whether I should rely more or less on AI based on the specific situation.", "+"],
        ["Jd03", "Appropriateness", "I find it difficult to judge if I am trusting an AI tool the right amount.", "R"],
        ["Jd04", "Appropriateness", "I think carefully about whether AI output is reliable enough before I use it for learning.", "+"],
        ["Jd05", "Comparison", "I consider how well AI-generated content aligns with what I already know about a topic.", "+"],
        ["Jd06", "Comparison", "When AI gives me an answer, I think about whether other sources would give similar results.", "+"],
        ["Jd07", "Comparison", "I evaluate AI suggestions differently depending on how much the task outcome matters.", "+"],
        ["Jd08", "Comparison", "I accept AI outputs without thinking about how well they fit the specific context.", "R"],
        ["Jd09", "Discrepancy", "When something about an AI\u2019s output doesn\u2019t seem right, I think about why it might be wrong.", "+"],
        ["Jd10", "Discrepancy", "I notice inconsistencies in AI-generated content that suggest it may be unreliable.", "+"],
        ["Jd11", "Discrepancy", "I often overlook errors or inconsistencies in AI-generated content.", "R"],
        ["Jd12", "Discrepancy", "When an AI tool gives an unexpected answer, I think about whether the output is trustworthy.", "+"],
        ["Jd13", "Discrepancy", "I think about whether trusting AI in a given situation could lead to a poor outcome.", "+"],
        ["Jd14", "Temporal drift", "I think about whether I may have become too comfortable relying on an AI tool over repeated use.", "+"],
        ["Jd15", "Proportional", "I evaluate whether the amount of trust I place in AI is proportional to the evidence I have about its accuracy.", "+"],
        ["Jd16", "Overgeneralization", "I think about whether my positive experience with AI in one area might make me trust it too much in a different area.", "+"],
    ]

    ac_items = [
        ["Ac01", "Verification", "I cross-reference important AI outputs with textbooks, articles, or other reliable sources.", "+"],
        ["Ac02", "Verification", "When an AI tool gives me a factual claim, I look it up to confirm it is correct.", "+"],
        ["Ac03", "Verification", "I use AI suggestions without verifying them.", "R"],
        ["Ac04", "Verification", "I ask follow-up questions to an AI tool to test the quality of its responses.", "+"],
        ["Ac05", "Adjustment", "I adjust my reliance on AI based on how well it has performed in similar tasks.", "+"],
        ["Ac06", "Adjustment", "After finding an error in an AI output, I change how much I trust it for future tasks.", "+"],
        ["Ac07", "Adjustment", "Even when AI gives me wrong answers, I continue to use it the same way.", "R"],
        ["Ac08", "Adjustment", "I adjust how much I depend on AI tools depending on the difficulty of the task.", "+"],
        ["Ac09", "Strategy", "I have specific strategies for evaluating whether AI-generated content is trustworthy.", "+"],
        ["Ac10", "Strategy", "I intentionally test AI tools on questions I already know the answer to.", "+"],
        ["Ac11", "Strategy", "I set personal rules for when to accept and when to question AI outputs.", "+"],
        ["Ac12", "Strategy", "I lack a systematic approach for deciding whether to trust AI outputs.", "R"],
        ["Ac13", "Strategy", "I reflect on whether my use of AI helped or hurt my learning after completing a task.", "+"],
        ["Ac14", "Social", "I discuss AI-generated content with classmates or instructors to get their perspective on its quality.", "+"],
        ["Ac15", "Proactive", "I deliberately approach some AI outputs with skepticism, even when they look correct, to practice critical evaluation.", "+"],
        ["Ac16", "Meta-strategic", "I revise my approach to evaluating AI outputs when I discover that my previous method was insufficient.", "+"],
    ]

    all_items = aw_items + jd_items + ac_items
    total_rows = len(all_items) + 1 + 3  # +1 header, +3 subscale separator headers

    table = doc.add_table(rows=1 + 3 + len(all_items), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(0.5), Inches(1.0), Inches(4.2), Inches(0.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    row_idx = 1

    # CA-Aw section
    merged = table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[3])
    set_cell_text(merged, "Calibration Awareness (CA-Aw): 1 = Strongly Disagree to 7 = Strongly Agree",
                  bold=True, italic=True, size=10)
    row_idx += 1

    for item in aw_items:
        for j, val in enumerate(item):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else None
            set_cell_text(table.rows[row_idx].cells[j], val, size=9.5, alignment=alignment)
        row_idx += 1

    # CA-Jd section
    merged = table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[3])
    set_cell_text(merged, "Calibration Judgment (CA-Jd): 1 = Strongly Disagree to 7 = Strongly Agree",
                  bold=True, italic=True, size=10)
    row_idx += 1

    for item in jd_items:
        for j, val in enumerate(item):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else None
            set_cell_text(table.rows[row_idx].cells[j], val, size=9.5, alignment=alignment)
        row_idx += 1

    # CA-Ac section
    merged = table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[3])
    set_cell_text(merged, "Calibration Action (CA-Ac): 1 = Never to 7 = Always",
                  bold=True, italic=True, size=10)
    row_idx += 1

    for item in ac_items:
        for j, val in enumerate(item):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else None
            set_cell_text(table.rows[row_idx].cells[j], val, size=9.5, alignment=alignment)
        row_idx += 1

    format_apa_table(table)

    # Add inner borders for subscale headers
    for sep_row in [1, 1 + 1 + len(aw_items), 1 + 2 + len(aw_items) + len(jd_items)]:
        for cell in table.rows[sep_row].cells:
            set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"})

    add_table_note(doc,
        "Dir. = Direction; + = positively worded; R = reverse-coded. "
        "Items Aw03, Aw09, Aw12, Jd03, Jd08, Jd11, Ac03, Ac07, and Ac12 are reverse-coded (9 items, 19%). "
        "Calibration Awareness and Calibration Judgment items use an agreement scale; "
        "Calibration Action items use a frequency scale.")


def add_table_5(doc):
    """Table 5: Planned Validity Evidence Framework"""
    add_table_title(doc, 5, "Planned Convergent and Discriminant Validity Evidence")

    headers = ["Instrument", "TCRS Subscale", "Expected r", "Type", "Rationale"]
    data = [
        ["MAILS Short\n(Carolus et al., 2024)", "CA-Aw\nCA-Jd\nCA-Ac",
         ".45\u2013.55\n.35\u2013.45\n.30\u2013.40", "Convergent",
         "AI literacy overlaps with\ncalibration awareness;\nless relevant to behavioral\nregulation"],
        ["Critical Thinking\nDisposition Scale\n(Sosu, 2013)", "CA-Jd\nCA-Ac\nCA-Aw",
         ".45\u2013.60\n.30\u2013.40\n.25\u2013.35", "Convergent",
         "Evaluative thinking overlaps\nwith trust judgment;\nfalsification test: if\nCA-Jd \u00D7 CT > .60,\nJudgment may not be distinct"],
        ["S-TIAS (2025)", "All subscales", "< .30", "Discriminant",
         "Trust attitude \u2260 trust\ncalibration readiness"],
        ["General Self-Efficacy\n(Schwarzer &\nJerusalem, 1995)", "All subscales", "< .30", "Discriminant",
         "Self-efficacy language was\nremoved from TCRS items;\nlow r confirms success"],
        ["MC-SDS Short Form\n(Crowne &\nMarlowe, 1960)", "All subscales", "< .30", "Response\nbias control",
         "If any subscale r > .30,\nthis signals social\ndesirability concern"],
    ]

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(1.2), Inches(0.9), Inches(0.8), Inches(0.8), Inches(2.3)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i + 1].cells[j], val, size=10)

    format_apa_table(table)
    add_table_note(doc,
        "MAILS = Meta AI Literacy Scale; S-TIAS = Short Trust in Automation Scale; "
        "MC-SDS = Marlowe-Crowne Social Desirability Scale. "
        "Expected correlation ranges are based on theoretical predictions from the process model. "
        "The falsification test for Calibration Judgment: if CA-Jd \u00D7 Critical Thinking r \u2265 .60, "
        "a two-subscale solution (Awareness + Action) would be reported.")


def add_table_6(doc):
    """Table 6: Planned CFA Model Comparisons"""
    add_table_title(doc, 6, "Planned Confirmatory Factor Analysis Model Comparisons")

    headers = ["Model", "Structure", "Description"]
    data = [
        ["Model 1", "Single-factor", "All items load on one general trust calibration readiness factor"],
        ["Model 2", "Two correlated factors", "Calibration Cognition (Awareness + Judgment) and Calibration Action"],
        ["Model 3", "Three correlated factors", "Calibration Awareness, Calibration Judgment, and Calibration Action (as theorized)"],
        ["Model 4", "Bifactor", "General trust calibration readiness factor + three specific group factors"],
    ]

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(0.8), Inches(1.5), Inches(4.0)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i + 1].cells[j], val, size=10)

    format_apa_table(table)
    add_table_note(doc,
        "Model fit will be evaluated using: \u03C7\u00B2, CFI > .95, TLI > .95, "
        "RMSEA < .06, SRMR < .08 (Hu & Bentler, 1999). "
        "If the bifactor model (Model 4) is supported, \u03C9h and \u03C9s will be computed "
        "(Rodriguez et al., 2016).")


# ─── Document Builder ───

def create_title_page(doc):
    """APA 7 title page."""
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    add_run(p, "Trust Calibration in Educational AI:", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    add_run(p, "A Process Model and Scale Development", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0

    for text in ["Hosung You",
                 "Department of Artificial Intelligence Convergence Education",
                 "Sungkyunkwan University"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        add_run(p, text)

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    add_run(p, "Author Note", bold=True)

    add_body_text(doc,
        "Correspondence concerning this article should be addressed to "
        "Hosung You, Department of Artificial Intelligence Convergence Education, "
        "Sungkyunkwan University, Seoul, South Korea.")

    doc.add_page_break()


def build_abstract(doc):
    """Abstract page."""
    add_heading(doc, "Abstract", level=1)

    add_body_text(doc,
        "As artificial intelligence becomes increasingly integrated into educational settings, "
        "the question of whether learners can appropriately calibrate their trust in AI systems "
        "has become critical. A recent systematic review found that 83.5% of studies examining "
        "trust in educational AI fell below a calibration threshold, revealing a pervasive gap "
        "between learner trust and AI reliability. Yet no instrument currently exists to measure "
        "learners\u2019 readiness to calibrate their trust in AI tools. This study addresses this "
        "gap through two contributions. First, we propose a process model of trust calibration "
        "that synthesizes metacognitive monitoring theory, trust in automation research, and "
        "human agency theory into a theorized sequential framework: Awareness (recognizing one\u2019s "
        "trust state) \u2192 Judgment (evaluating trust appropriateness) \u2192 Action (actively adjusting "
        "trust and verifying AI outputs). Second, we develop and validate the Trust Calibration "
        "Readiness Scale (TCRS), a self-report instrument measuring learners\u2019 preparedness at "
        "each stage of this calibration process. In Study 1 (N = [TBC]), exploratory factor "
        "analysis was conducted to determine the scale\u2019s factor structure. In Study 2 (N = [TBC]), "
        "confirmatory factor analysis and convergent/discriminant validity evidence were examined. "
        "The results [TO BE COMPLETED]. The TCRS offers educators and researchers a practical "
        "tool for assessing and supporting learners\u2019 capacity to engage in appropriate trust "
        "calibration when using AI for learning.",
        first_line_indent=False)

    p = add_paragraph(doc)
    p.paragraph_format.first_line_indent = Inches(0.5)
    add_run(p, "Keywords: ", italic=True)
    add_run(p, "trust calibration, AI literacy, educational AI, scale development, metacognition, self-regulated learning")

    doc.add_page_break()


def build_introduction(doc):
    """Introduction section."""
    add_heading(doc, "Introduction", level=1)

    # 1.1
    add_heading(doc, "The Trust Calibration Problem in Educational AI", level=2)

    add_body_text(doc,
        "The integration of artificial intelligence into education has accelerated dramatically, "
        "with AI tools now supporting writing, problem-solving, assessment, and personalized "
        "learning at unprecedented scale. As learners interact with AI systems of varying "
        "reliability across diverse tasks, a fundamental question emerges: Can learners "
        "appropriately calibrate their trust in these systems?")

    add_body_text(doc,
        "Trust calibration\u2014the correspondence between a learner\u2019s level of trust and an AI "
        "system\u2019s actual reliability\u2014is critical because miscalibration carries significant "
        "consequences. Overtrust leads to misuse: uncritical acceptance of inaccurate AI outputs, "
        "reduced critical thinking, and degraded learning outcomes (Lee & See, 2004; Parasuraman "
        "& Riley, 1997). Distrust leads to disuse: rejection of helpful AI support, increased "
        "cognitive load, and missed learning opportunities (Lee & See, 2004). Appropriate trust\u2014"
        "calibrated trust that matches the AI\u2019s actual capabilities and limitations\u2014enables "
        "learners to leverage AI\u2019s strengths while maintaining intellectual autonomy.")

    add_body_text(doc,
        "The scope of this problem is substantial. A systematic review of 97 studies examining "
        "trust in educational AI (2015\u20132026) found that 83.5% fell below a calibration threshold, "
        "meaning that the vast majority of research documented situations where learner trust was "
        "misaligned with AI reliability ([Author], 2026). This \u201Ccalibration gap\u201D is not merely an "
        "academic concern; it has direct implications for whether AI-enhanced education fulfills "
        "its promise or undermines the learning it aims to support.")

    add_body_text(doc,
        "Despite the growing recognition of this problem, the field lacks a measurement instrument "
        "that captures learners\u2019 readiness to calibrate their trust in AI. Existing AI literacy "
        "scales focus on general knowledge and skills (Carolus et al., 2023; Laupichler et al., "
        "2023; Ng et al., 2024) but do not address the specific cognitive and behavioral processes "
        "involved in trust calibration. Trust measurement instruments capture trust attitudes "
        "(Jian et al., 2000) but not the capacity to calibrate those attitudes. This study "
        "addresses this gap by developing the Trust Calibration Readiness Scale (TCRS).")

    # 1.2
    add_heading(doc, "Distinguishing Trust from Trust Calibration", level=2)

    add_body_text(doc,
        "A critical distinction must be drawn between trust and trust calibration, as the two "
        "are often conflated in educational AI research.")

    add_body_text(doc,
        "Trust is an attitude: a learner\u2019s willingness to be vulnerable to an AI system based "
        "on expectations about its future behavior (Lee & See, 2004; Mayer et al., 1995). Trust "
        "can be measured as a level (high or low) and is influenced by factors such as perceived "
        "reliability, system transparency, and prior experience.")

    add_body_text(doc,
        "Trust calibration is a process: the degree to which a learner\u2019s trust level matches the "
        "AI system\u2019s actual reliability (Lee & See, 2004). Calibration is not about having more "
        "or less trust; it is about having the *right amount* of trust for the situation. A learner "
        "who highly trusts a highly reliable AI system is well-calibrated; a learner who highly "
        "trusts an unreliable AI system is miscalibrated (overtrusting).")

    add_body_text(doc,
        "This distinction has a direct measurement implication. Existing trust scales (e.g., "
        "Trust in Automation Scale; Jian et al., 2000) measure trust level but cannot assess "
        "whether that trust is appropriately calibrated. Measuring calibration accuracy would "
        "require comparing self-reported trust against an objective reliability benchmark\u2014a "
        "design that is infeasible for a general-purpose survey instrument. We therefore adopt a "
        "third approach: measuring trust calibration readiness\u2014the knowledge, awareness, and "
        "behavioral disposition that enable a learner to engage in the calibration process. This "
        "approach is analogous to measuring metacognitive awareness (Schraw & Dennison, 1994) "
        "rather than metacognitive accuracy, or measuring self-regulated learning strategies "
        "(Pintrich et al., 1993) rather than actual self-regulation outcomes.")

    # 1.3
    add_heading(doc, "Existing AI Literacy Scales and Their Limitations", level=2)

    add_body_text(doc,
        "The field of AI literacy measurement has grown rapidly, with a systematic review "
        "identifying 16 distinct scales validated across 22 studies (Lintner, 2024). These "
        "instruments represent important advances but leave the trust calibration domain "
        "unaddressed. Table 2 summarizes the major existing scales and their coverage of trust "
        "calibration elements.")

    add_body_text(doc,
        "General AI literacy scales measure knowledge and skills related to AI understanding "
        "and use. The Scale for the Assessment of Non-experts\u2019 AI Literacy (SNAIL; Laupichler "
        "et al., 2023) measures technical understanding, critical appraisal, and practical "
        "application across 31 items. The Meta AI Literacy Scale (MAILS; Carolus et al., 2023) "
        "uniquely includes psychological meta-competencies such as persuasion literacy and "
        "emotion regulation alongside AI knowledge. The AI Literacy Questionnaire (AILQ; Ng et "
        "al., 2024) adopts an Affective-Behavioural-Cognitive-Ethical framework validated with "
        "secondary students. More recently, performance-based assessments have emerged, including "
        "SAIL4ALL (Pinski et al., 2025), which measures AI literacy through knowledge tests "
        "rather than self-report.")

    add_body_text(doc,
        "Generative AI literacy scales have proliferated in response to the ChatGPT era. Multiple "
        "GAIL (Generative AI Literacy) scales have been proposed with varying structures, typically "
        "measuring awareness, usage, evaluation, and ethical dimensions of generative AI interaction.")

    add_body_text(doc,
        "Collaborative AI metacognition scales represent the closest existing construct to trust "
        "calibration readiness. Recent work on collaborative AI literacy and collaborative AI "
        "metacognition (2025) measures planning, monitoring, and evaluating in human-AI interaction "
        "through a distributed cognition lens.")

    add_body_text(doc,
        "Despite this proliferation, no existing scale addresses three critical elements that "
        "trust calibration requires: (a) awareness of one\u2019s own trust state in relation to AI "
        "(not just general AI awareness), (b) evaluation of whether that trust is appropriate "
        "for the specific situation, and (c) active behavioral regulation of trust through "
        "verification and adjustment. The Trust Calibration Readiness Scale (TCRS) is designed "
        "to fill this gap.")

    # >>> TABLE 2 <<<
    add_table_2(doc)

    # 1.4
    add_heading(doc, "Purpose of the Present Study", level=2)

    add_body_text(doc, "This study has two aims:")

    add_numbered(doc, 1,
        "**Propose a process model of trust calibration** that integrates metacognitive monitoring "
        "theory, trust in automation research, and human agency theory to explain how learners "
        "calibrate their trust in AI systems.")
    add_numbered(doc, 2,
        "**Develop and validate the TCRS**, a self-report instrument that measures learners\u2019 "
        "readiness to engage in each stage of the trust calibration process.")


def build_theoretical_framework(doc):
    """Theoretical Framework section."""
    add_heading(doc, "Theoretical Framework: A Process Model of Trust Calibration", level=1)

    # 2.1
    add_heading(doc, "Calibration as a Process, Not a Trait", level=2)

    add_body_text(doc,
        "We propose that trust calibration is not a stable individual characteristic but a "
        "dynamic process that learners engage in (or fail to engage in) when interacting with "
        "AI systems. This process involves three cognitive-behavioral components with a "
        "theorized ordering, each grounded in established theory.")

    add_body_text(doc,
        "This study adopts a post-positivist, critical realist position: we assume that trust "
        "calibration is a real cognitive-behavioral process with an underlying structure, while "
        "acknowledging that our measurement of it is fallible and theory-laden. This paradigmatic "
        "stance justifies the use of latent variable measurement for constructs that are real "
        "but not directly observable, and motivates the readiness framing\u2014measuring dispositional "
        "capacity rather than claiming direct access to the calibration process itself.")

    add_body_text(doc,
        "The model was developed through a deliberate theoretical construction that draws "
        "selectively from three established traditions not previously connected in the "
        "educational AI literature: metacognitive monitoring from self-regulated learning theory, "
        "trust dynamics from human-automation interaction research, and intentional "
        "self-regulation from social cognitive theory. Critically, we do not claim to import "
        "the full scope of any single framework. Each theory serves as an *explanatory resource* "
        "for a specific functional capacity required by the target construct, not as a structural "
        "pillar of the model. This approach\u2014assembling functional components from different "
        "theoretical traditions to explain a novel target construct\u2014follows established precedent "
        "in multi-source theoretical models such as the Technology Acceptance Model (Davis, 1989), "
        "the Theory of Planned Behavior (Ajzen, 1991), and the Transactional Model of Stress "
        "(Lazarus & Folkman, 1984).")

    add_body_text(doc,
        "The three functional capacities\u2014monitoring, evaluating, and regulating\u2014are not "
        "selected because they happen to map onto available theories. Rather, they are derived "
        "from the logical requirements of the target construct. Trust calibration, by definition, "
        "requires: (a) awareness of one\u2019s current trust state, because without knowing one\u2019s "
        "position, adjustment is impossible; (b) evaluation of whether that state is appropriate, "
        "because without a criterion, the direction of adjustment is unknown; and (c) intentional "
        "behavioral adjustment, because without action, awareness and evaluation remain inert. "
        "This necessity argument parallels the functional decomposition in other multi-source "
        "models: the Theory of Planned Behavior requires attitudinal, normative, and control "
        "components because behavioral intention logically requires all three (Ajzen, 1991); "
        "the Transactional Model of Stress requires appraisal and coping components because "
        "stress response logically requires both (Lazarus & Folkman, 1984).")

    # >>> TABLE 1 (Selective Borrowing) <<<
    add_table_selective_borrowing(doc)

    # 2.2
    add_heading(doc, "The Trust Calibration Process Model", level=2)

    add_body_text(doc,
        "The process model consists of four components arranged in a theorized sequential "
        "relationship (see Figure 1). While we propose a sequential ordering from awareness "
        "through evaluation to action, we acknowledge that in practice, these components may "
        "involve feedback loops and parallel processing. The theorized sequence represents the "
        "canonical pathway of deliberate calibration; recursive cycles in which calibration "
        "outcomes inform subsequent metacognitive monitoring are expected and consistent with "
        "the model.")

    add_body_text(doc,
        "**Monitoring Capacity (Foundation).** The calibration process begins with the "
        "learner\u2019s awareness of their own trust state when interacting with an AI system: "
        "recognizing how much they trust the AI, understanding what the AI system is designed "
        "to do, and attending to the task context in which the AI is being used. Without this "
        "monitoring foundation\u2014without knowing \u201Cwhere they are\u201D in the trust-reliability "
        "space\u2014calibration cannot begin. This monitoring function is informed by the concept "
        "of self-monitoring within self-regulated learning theory (Winne & Hadwin, 1998), "
        "specifically the monitoring processes central to Phases 1\u20133 of their SRL framework. "
        "We borrow the monitoring mechanism specifically; the broader SRL framework\u2019s planning, "
        "enacting, and adapting processes are beyond the scope of this model.")

    add_body_text(doc,
        "**Evaluative Capacity (Variable).** The second component is the cognitive operation "
        "of assessing the variable being calibrated: trust itself. Trust functions as a force "
        "that can push the learner away from the optimal point\u2014upward toward overtrust "
        "(leading to misuse) or downward toward distrust (leading to disuse; Lee & See, 2004). "
        "Critically, trust in this model is not the agent of calibration but the object of "
        "calibration. The learner must evaluate whether their current trust level is appropriate "
        "for the situation. This evaluative function is informed by trust dynamics research "
        "(Lee & See, 2004), which identifies the conditions under which trust aligns or "
        "misaligns with system reliability. We borrow the trust-reliability matching concept "
        "specifically; the broader trust literature\u2019s treatment of trust formation, repair, "
        "and dispositional antecedents is beyond the scope of this model.")

    add_body_text(doc,
        "**Regulatory Capacity (Mechanism).** The third component is the mechanism that "
        "enables calibration: intentional behavioral adjustment. This stage encompasses the "
        "learner\u2019s active behaviors to verify AI outputs, compare them against other sources, "
        "and adjust their reliance accordingly. Regulation transforms the learner from a "
        "passive trust holder into an active trust regulator. This regulatory function is "
        "informed by the self-reactiveness component of Bandura\u2019s (2001) theory of human "
        "agency\u2014the capacity for intentional self-regulation of behavior based on "
        "self-monitoring. We borrow the behavioral self-regulation mechanism specifically; "
        "Bandura\u2019s broader agentic framework\u2014including intentionality, forethought, "
        "collective agency, and proxy agency\u2014is beyond the scope of this model.")

    add_body_text(doc,
        "**Calibration (Outcome).** The outcome of this process is calibrated trust: a state "
        "in which the learner\u2019s trust level appropriately matches the AI system\u2019s actual "
        "reliability for the specific task and context. This is the \u201Coptimal point\u201D on the "
        "regression line where trust corresponds to reliability (Figure 1a).")

    add_body_text(doc,
        "It is important to note that the process model describes deliberate trust "
        "calibration\u2014what dual-process theory terms System 2 processing (Kahneman, 2011). "
        "Much of everyday trust in AI may operate through automatic, heuristic-based responses "
        "(System 1): a learner may habitually accept AI outputs without conscious evaluation, "
        "or reflexively reject them based on prior negative experience. These automatic trust "
        "responses can preempt the deliberate calibration process entirely\u2014a phenomenon "
        "well-documented in automation complacency research (Parasuraman & Manzey, 2010). "
        "The TCRS, therefore, measures readiness to engage in deliberate calibration, not "
        "the automatic trust processes that may bypass it. This boundary condition also helps "
        "explain why calibration readiness (as measured by self-report) may not perfectly "
        "predict calibration accuracy in practice: a learner may possess the requisite awareness "
        "and evaluative capacity yet still default to automatic trust responses under cognitive "
        "load or time pressure.")

    # 2.3
    add_heading(doc, "The Trust Calibration Space", level=2)

    add_body_text(doc,
        "The process model can be visualized through what we term the *trust calibration space* "
        "(Figure 1a). In this space, the x-axis represents AI reliability (actual system "
        "performance) and the y-axis represents the learner\u2019s trust level. The diagonal line "
        "represents perfect calibration\u2014the ideal where trust precisely matches reliability. "
        "Points above this line represent overtrust; points below represent distrust.")

    add_body_text(doc,
        "Within this space, monitoring capacity enables the learner to locate themselves (\u201CWhere "
        "am I?\u201D), evaluative capacity enables them to assess the distance from the calibration "
        "line (\u201CAm I over- or under-trusting?\u201D), and regulatory capacity provides the mechanism "
        "to move toward the line (\u201CAdjust\u201D). This visual metaphor makes explicit what the process "
        "model describes abstractly: calibration is the act of moving toward appropriate trust "
        "through awareness, evaluation, and intentional adjustment.")

    add_body_text(doc,
        "It is important to note that the TCRS measures learners\u2019 readiness to engage in "
        "the calibration process, not their actual position in the trust calibration space. "
        "Locating a learner\u2019s precise coordinates in this space would require an objective "
        "measure of both the learner\u2019s trust level and the AI system\u2019s reliability\u2014a "
        "task-specific, context-dependent assessment that is beyond the scope of a "
        "general-purpose survey instrument. The trust calibration space serves as a conceptual "
        "framework for understanding what calibration means; the TCRS serves as a practical "
        "tool for assessing whether learners are prepared to engage in it.")

    # 2.4
    add_heading(doc, "From Process Model to Measurement: Three Subscales", level=2)

    add_body_text(doc,
        "The process model yields three measurable constructs that correspond to the first "
        "three stages of the calibration process. The fourth component (calibration outcome) "
        "is not directly measured because the TCRS assesses readiness rather than accuracy. "
        "Table 3 summarizes the subscale structure and theoretical foundations.")

    add_body_text(doc,
        "**Calibration Awareness (CA-Aw)** captures the monitoring capacity: the learner\u2019s "
        "awareness of their own trust state, their understanding of the AI system\u2019s "
        "capabilities, and their recognition of task context factors that should influence "
        "trust. This subscale is informed by metacognitive monitoring theory (Winne & "
        "Hadwin, 1998; Efklides, 2006), particularly the concepts of self-monitoring and "
        "metacognitive experiences within the SRL framework\u2019s Phases 1\u20133. Importantly, "
        "CA-Aw measures trust state monitoring in the context of AI interaction; it does not "
        "claim to measure metacognitive awareness in general, nor does a low CA-Aw score "
        "indicate a general metacognitive deficit.")

    add_body_text(doc,
        "**Calibration Judgment (CA-Jd)** captures the evaluative capacity: the learner\u2019s "
        "tendency to evaluate whether their trust level is appropriate for the situation, "
        "to detect discrepancies between trust and reliability, and to consider contextual "
        "factors that should moderate their trust. This subscale is informed by trust dynamics "
        "research (Lee & See, 2004) and the concept of epistemic vigilance\u2014the capacity to "
        "evaluate the reliability of information sources and the plausibility of claims "
        "(Sperber et al., 2010). The judgment dimension is further strengthened by its "
        "connection to epistemic cognition research (Barzilai & Zohar, 2014), positioning "
        "CA-Jd as a domain-specific epistemic practice. Importantly, CA-Jd measures trust "
        "appropriateness evaluation in the AI context; it does not claim to measure general "
        "critical thinking or evaluative judgment.")

    add_body_text(doc,
        "**Calibration Action (CA-Ac)** captures the regulatory capacity: the learner\u2019s "
        "active behaviors to verify AI outputs, adjust their reliance based on evidence, and "
        "employ deliberate strategies for trust regulation. This subscale is informed by the "
        "self-reactiveness component of agency theory (Bandura, 2001) and the meta-strategic "
        "regulation component of SRL theory (Winne & Hadwin, 1998, Phase 4). Importantly, "
        "CA-Ac measures trust regulation behaviors in the AI context; it does not claim to "
        "measure general behavioral agency or self-regulation capacity.")

    add_body_text(doc,
        "It is important to note that these three subscales measure components of a theorized "
        "process sequence, yet they are assessed as parallel subscales within one instrument. "
        "This approach follows established precedent. The Motivated Strategies for Learning "
        "Questionnaire (MSLQ; Pintrich et al., 1993) measures planning, monitoring, and "
        "regulation\u2014theoretically sequential SRL stages\u2014as parallel subscales. The "
        "Metacognitive Awareness Inventory (MAI; Schraw & Dennison, 1994) measures knowledge "
        "of cognition and regulation of cognition as parallel subscales despite their theoretical "
        "sequentiality. In each case, the instrument captures the learner\u2019s readiness or "
        "capacity at each stage, not the real-time execution of the process.")

    # >>> TABLE 3 <<<
    add_table_3(doc)

    # >>> FIGURE 1 <<<
    if os.path.exists(FIGURE_FILE):
        add_figure_apa(doc, FIGURE_FILE, 1,
            "The Trust Calibration Process Model. Panel (a): the trust calibration space\u2014"
            "diagonal = perfect calibration; above = overtrust/misuse; below = distrust/disuse. "
            "Panel (b): the process model with recursive feedback: Monitoring Capacity "
            "(Winne & Hadwin, 1998) \u2192 Evaluative Capacity (Lee & See, 2004) \u2192 Regulatory "
            "Capacity (Bandura, 2001) \u2192 Calibration Outcome. Source theories serve as "
            "explanatory resources (see Table 1). The TCRS measures readiness to engage in "
            "this process, not the learner\u2019s position in the calibration space.")


def build_present_study(doc):
    """The Present Study section."""
    add_heading(doc, "The Present Study", level=1)

    # 3.1
    add_heading(doc, "Scale Development Overview", level=2)

    add_body_text(doc,
        "The Trust Calibration Readiness Scale (TCRS) was developed following DeVellis and "
        "Thorpe\u2019s (2022) guidelines for scale development. The development process consisted "
        "of five phases:")

    add_numbered(doc, 1,
        "**Construct definition and item generation:** A pool of 48 items was generated based "
        "on the process model, existing literature, and related scales. Items were distributed "
        "across three subscales (16 per subscale) covering nine facets.")
    add_numbered(doc, 2,
        "**Expert panel review:** [NUMBER] experts in AI literacy, educational measurement, "
        "trust in automation, and metacognition evaluated each item for content validity using "
        "a 4-point relevance rating scale.")
    add_numbered(doc, 3,
        "**Cognitive interviews:** [NUMBER] undergraduate and graduate students participated in "
        "think-aloud protocols to assess item comprehension and response process validity.")
    add_numbered(doc, 4,
        "**Study 1 (Exploratory Factor Analysis):** An independent sample was surveyed to "
        "determine the scale\u2019s factor structure empirically.")
    add_numbered(doc, 5,
        "**Study 2 (Confirmatory Factor Analysis and Validity):** A second independent sample "
        "was surveyed to confirm the factor structure and establish convergent, discriminant, "
        "and known-groups validity.")

    # 3.2
    add_heading(doc, "Item Development", level=2)

    add_heading(doc, "Item Generation", level=3)

    add_body_text(doc,
        "An initial pool of 48 items was generated through a systematic process:")

    add_bullet(doc,
        "**Theoretical derivation:** Items were written to reflect the three process stages "
        "and their facets, using language grounded in the source theories (Winne & Hadwin, 1998; "
        "Lee & See, 2004; Bandura, 2001).")
    add_bullet(doc,
        "**Related scale adaptation:** Items were informed by (but not copied from) existing "
        "measures including the MAI (Schraw & Dennison, 1994), MAILS (Carolus et al., 2023), "
        "and collaborative AI metacognition scales.")
    add_bullet(doc,
        "**Linguistic sharpening:** Following expert psychometric review, items were carefully "
        "crafted to maintain construct boundaries. Awareness items used monitoring verbs "
        "(e.g., \u201CI notice,\u201D \u201CI am aware,\u201D \u201CI pay attention to\u201D) to capture passive "
        "metacognitive processes. Judgment items used evaluative verbs (e.g., \u201CI evaluate,\u201D "
        "\u201CI weigh,\u201D \u201CI think about whether\u201D) without self-efficacy language (avoiding "
        "\u201CI can\u201D or \u201CI am able to\u201D). Action items used behavioral verbs (e.g., \u201CI check,\u201D "
        "\u201CI adjust,\u201D \u201CI discuss\u201D) with frequency response anchors.")

    add_body_text(doc,
        "The item pool covered nine facets across three subscales (see Table 4 for the full "
        "item pool):")

    add_body_text(doc,
        "**Calibration Awareness (16 items):** trust state awareness (awareness of one\u2019s trust "
        "level), AI system understanding (knowledge of AI capabilities), task context recognition "
        "(awareness of situational factors), affective monitoring (noticing emotional responses to "
        "AI), temporal drift awareness (recognizing trust changes over time), and social awareness "
        "(noticing others\u2019 trust levels).")

    add_body_text(doc,
        "**Calibration Judgment (16 items):** appropriateness evaluation (judging whether trust "
        "matches the situation), situational comparison (evaluating trust relative to prior "
        "knowledge and other sources), discrepancy detection (noticing when trust is misaligned), "
        "temporal drift judgment (evaluating habituation effects), proportional trust (assessing "
        "whether trust is proportional to evidence), and overgeneralization detection (recognizing "
        "inappropriate trust transfer across domains).")

    add_body_text(doc,
        "**Calibration Action (16 items):** verification behaviors (checking AI outputs against "
        "other sources), trust adjustment (modifying reliance based on evidence), strategic "
        "regulation (employing deliberate evaluation strategies), social verification (discussing "
        "AI outputs with peers), proactive calibration (deliberately practicing critical evaluation), "
        "and meta-strategic regulation (revising one\u2019s evaluation approach).")

    # >>> TABLE 4 <<<
    add_table_4(doc)

    add_heading(doc, "Response Format", level=3)

    add_body_text(doc,
        "Two response formats were used. Calibration Awareness and Calibration Judgment items "
        "used a 7-point agreement scale (1 = *Strongly Disagree* to 7 = *Strongly Agree*) because "
        "these subscales assess self-perceptions of awareness and evaluative tendencies. "
        "Calibration Action items used a 7-point frequency scale (1 = *Never* to 7 = *Always*) "
        "because this subscale assesses behavioral occurrence, and frequency anchors reduce "
        "acquiescence bias for behavioral items (Krosnick, 1999). This mixed-format approach "
        "follows the precedent of the MSLQ, which uses different response formats across "
        "subscales.")

    add_heading(doc, "Reverse-Coded Items", level=3)

    add_body_text(doc,
        "Nine items (19%) were reverse-coded to control for acquiescence bias (three per "
        "subscale). Reverse items were constructed as natural negations or behavioral absences "
        "(e.g., \u201CI rarely think about whether I am trusting AI too much or too little\u201D) "
        "rather than awkward grammatical reversals.")

    # 3.3
    add_heading(doc, "Expert Panel Review", level=2)

    add_body_text(doc,
        "[NUMBER] experts participated in a content validity review. Expert specializations "
        "included AI in education (n = [TBC]), trust in automation (n = [TBC]), psychometrics "
        "(n = [TBC]), metacognition and self-regulated learning (n = [TBC]), and AI literacy "
        "(n = [TBC]).")

    add_body_text(doc,
        "Each expert rated every item on a 4-point relevance scale (1 = *not relevant*, "
        "2 = *somewhat relevant*, 3 = *quite relevant*, 4 = *highly relevant*) and indicated "
        "which subscale each item best fit. Items with an item-level content validity index "
        "(I-CVI) below .78 were revised or removed. The scale-level CVI average (S-CVI/Ave) "
        "target was .90 or above.")

    add_body_text(doc, "[RESULTS TO BE COMPLETED]")

    # 3.4
    add_heading(doc, "Cognitive Interviews", level=2)

    add_body_text(doc,
        "[NUMBER] undergraduate and graduate students (STEM: n = [TBC]; humanities: n = [TBC]; "
        "varied AI experience) participated in think-aloud cognitive interviews. Participants "
        "read each item aloud, paraphrased its meaning, described a relevant personal experience, "
        "and identified any confusing language. Interviews focused on items flagged as potentially "
        "complex, new items without precedent in existing scales, and reverse-coded items.")

    add_body_text(doc, "[RESULTS TO BE COMPLETED]")


def build_study1(doc):
    """Study 1: EFA."""
    add_heading(doc, "Study 1: Exploratory Factor Analysis", level=1)

    add_heading(doc, "Method", level=2)

    add_heading(doc, "Participants", level=3)

    add_body_text(doc, "[TO BE COMPLETED]")
    add_body_text(doc,
        "Target: N = 300\u2013350 undergraduate and graduate students from [university/online "
        "platform]. Inclusion criteria: (a) enrolled in higher education, (b) experience using "
        "at least one AI tool for learning (e.g., ChatGPT, Grammarly, AI tutoring systems), "
        "(c) aged 18 or older.")

    add_heading(doc, "Procedure", level=3)

    add_body_text(doc,
        "Participants completed an online survey consisting of: (a) demographic questions "
        "(age, gender, field of study, year of study), (b) AI use questions (types of AI tools "
        "used, frequency of use, duration of use), (c) the TCRS item pool ([NUMBER] items after "
        "expert review and cognitive interviews), and (d) an attention check item.")

    add_heading(doc, "Data Analysis", level=3)

    add_body_text(doc,
        "Exploratory factor analysis was conducted following best practices (Watkins, 2018):")

    add_numbered(doc, 1,
        "**Data screening:** Multivariate normality, outliers, and missing data were assessed. "
        "Items with excessive skewness (>|2.0|) or kurtosis (>|7.0|) were flagged.")
    add_numbered(doc, 2,
        "**Factorability:** The Kaiser-Meyer-Olkin (KMO) measure of sampling adequacy "
        "(target > .80) and Bartlett\u2019s test of sphericity were assessed.")
    add_numbered(doc, 3,
        "**Factor extraction:** Principal axis factoring was used as the extraction method. "
        "The number of factors was determined using multiple criteria: parallel analysis "
        "(Horn, 1965), the scree plot, eigenvalue > 1, and theoretical coherence.")
    add_numbered(doc, 4,
        "**Factor rotation:** Oblique rotation (Promax, kappa = 4) was used because subscales "
        "were expected to be moderately correlated (r = .40\u2013.60).")
    add_numbered(doc, 5,
        "**Item retention criteria:** Items were retained if they loaded \u2265 .40 on the primary "
        "factor, loaded < .30 on non-target factors, and had a communality \u2265 .30.")
    add_numbered(doc, 6,
        "**Internal consistency:** Cronbach\u2019s alpha and McDonald\u2019s omega were computed for "
        "each retained subscale.")

    add_heading(doc, "Planned Comparisons", level=3)

    add_body_text(doc, "Three factor models were compared:")

    add_bullet(doc, "**Three-factor model:** As theorized (Awareness, Judgment, Action)")
    add_bullet(doc, "**Two-factor model:** Cognition (Awareness + Judgment) and Action")
    add_bullet(doc, "**Single-factor model:** Unidimensional trust calibration readiness")

    add_body_text(doc,
        "The factor structure was determined empirically; the theoretical model was treated as "
        "one candidate rather than a foregone conclusion.")

    add_heading(doc, "Results", level=2)

    add_body_text(doc, "[TO BE COMPLETED AFTER DATA COLLECTION]")


def build_study2(doc):
    """Study 2: CFA and Validity."""
    add_heading(doc, "Study 2: Confirmatory Factor Analysis and Validity Evidence", level=1)

    add_heading(doc, "Method", level=2)

    add_heading(doc, "Participants", level=3)

    add_body_text(doc, "[TO BE COMPLETED]")
    add_body_text(doc,
        "Target: N = 300\u2013350 undergraduate and graduate students from [university/online "
        "platform], independent of Study 1 sample. Same inclusion criteria as Study 1.")

    add_heading(doc, "Instruments", level=3)

    add_body_text(doc,
        "**Trust Calibration Readiness Scale (TCRS).** The items retained from Study 1 EFA.")

    add_body_text(doc,
        "In addition to the TCRS, participants completed five instruments for convergent "
        "validity, discriminant validity, and response bias control (see Table 5).")

    # >>> TABLE 5 <<<
    add_table_5(doc)

    add_heading(doc, "Procedure", level=3)

    add_body_text(doc,
        "Participants completed an online survey including: (a) demographics and AI use questions, "
        "(b) TCRS, (c) convergent and discriminant validity instruments (counterbalanced), "
        "(d) social desirability measure, and (e) attention checks.")

    add_heading(doc, "Data Analysis", level=3)

    add_body_text(doc,
        "**Confirmatory factor analysis.** CFA was conducted using maximum likelihood estimation "
        "with robust standard errors (MLR) in [software]. Four models were compared (see Table 6).")

    # >>> TABLE 6 <<<
    add_table_6(doc)

    add_body_text(doc,
        "If the bifactor model was supported, omega hierarchical (\u03C9h) and omega subscale "
        "(\u03C9s) coefficients were computed to assess the proportion of variance attributable to "
        "the general factor versus specific factors (Rodriguez et al., 2016).")

    add_body_text(doc,
        "**Method factor modeling.** Because the TCRS uses mixed response formats (agreement "
        "for Awareness/Judgment, frequency for Action), a method factor for response format was "
        "modeled to ensure that the factor structure reflects substantive constructs rather than "
        "method variance (Podsakoff et al., 2003).")

    add_body_text(doc,
        "**Convergent and discriminant validity.** Bivariate correlations between TCRS subscales "
        "and validity instruments were computed. Convergent validity was supported if TCRS "
        "subscales showed the differentiated pattern of correlations described in Table 5. "
        "Discriminant validity was supported if TCRS subscales correlated r < .30 with S-TIAS "
        "and General Self-Efficacy.")

    add_body_text(doc,
        "**Known-groups validity.** Group comparisons were planned based on:")
    add_bullet(doc,
        "*AI use experience:* High-experience users (daily use, multiple tools) were expected "
        "to score higher than low-experience users on all TCRS subscales.")
    add_bullet(doc,
        "*Field of study:* STEM students were expected to score higher on Calibration Awareness "
        "(greater familiarity with AI system properties) than humanities students, but similar "
        "on Calibration Action (behavioral regulation is domain-general).")

    add_body_text(doc,
        "**Falsification test for Calibration Judgment.** If the correlation between CA-Jd and "
        "the Critical Thinking Disposition Scale exceeded r = .60, this would suggest that "
        "Calibration Judgment does not provide incremental measurement beyond critical thinking. "
        "In this case, we would report a two-subscale solution (Awareness + Action) and discuss "
        "the implications honestly.")

    add_heading(doc, "Results", level=2)

    add_body_text(doc, "[TO BE COMPLETED AFTER DATA COLLECTION]")


def build_discussion(doc):
    """General Discussion."""
    add_heading(doc, "General Discussion", level=1)

    add_body_text(doc,
        "[TO BE COMPLETED AFTER DATA COLLECTION. The following sections outline the planned "
        "discussion structure.]")

    add_heading(doc, "Summary of Findings", level=2)
    add_body_text(doc, "[Report factor structure, reliability, and validity evidence]")

    add_heading(doc, "Theoretical Contributions", level=2)

    add_body_text(doc,
        "The trust calibration process model makes three theoretical contributions.")

    add_body_text(doc,
        "First, it identifies three functional capacities required for trust calibration\u2014"
        "monitoring, evaluating, and regulating\u2014and assembles them into a coherent process "
        "model. Each capacity is informed by an established theoretical tradition (metacognitive "
        "monitoring, trust dynamics, and behavioral self-regulation, respectively), but the "
        "organizing principle of the model is the target construct\u2019s internal logic, not the "
        "source theories themselves. This functional decomposition approach parallels other "
        "successful multi-source models in social science, including the Technology Acceptance "
        "Model (Davis, 1989), the Theory of Planned Behavior (Ajzen, 1991), and the "
        "Transactional Model of Stress (Lazarus & Folkman, 1984). Each source theory serves "
        "as an explanatory resource for one functional component, not as a structural pillar.")

    add_body_text(doc,
        "Second, the model distinguishes trust calibration from trust itself. This distinction "
        "has important implications for intervention design. Interventions targeting calibration "
        "should not simply aim to increase or decrease trust; they should support learners\u2019 "
        "monitoring capacity, evaluative capacity, and regulatory capacity\u2014the functional "
        "process by which appropriate trust is achieved.")

    add_body_text(doc,
        "Third, the process model provides a framework for understanding why learners "
        "miscalibrate. A learner may overtrust not because of any single deficit, but because "
        "of failures at different functional stages: failing to monitor that trust is high "
        "(monitoring failure), failing to evaluate whether high trust is warranted (evaluation "
        "failure), or knowing that trust should be lower but not acting on that knowledge "
        "(regulation failure). Importantly, these stage-specific diagnoses are domain-specific: "
        "they identify where the calibration process breaks down in the AI interaction context, "
        "not general deficits in metacognition, judgment, or agency.")

    add_heading(doc, "Practical Implications", level=2)

    add_body_text(doc, "The TCRS offers practical value for three audiences:")

    add_body_text(doc,
        "**Educators** can use the TCRS as a screening instrument to identify students who "
        "may need additional support before engaging with AI tools in learning activities. "
        "Students who score low on Calibration Awareness may benefit from structured prompts "
        "that direct attention to their trust state during AI interaction (e.g., \u201CHow much "
        "are you trusting this AI output right now?\u201D). Students who score low on Calibration "
        "Judgment may benefit from guided comparison exercises that model the evaluation of AI "
        "output quality. Students who score low on Calibration Action may benefit from structured "
        "verification protocols embedded in AI-assisted learning tasks. These recommendations "
        "target calibration-specific capacities, not general metacognitive, evaluative, or "
        "behavioral skills.")

    add_body_text(doc,
        "**Researchers** can use the TCRS as a dependent variable in intervention studies, as a "
        "moderator in studies of AI-enhanced learning outcomes, or as a screening tool in "
        "studies of AI literacy development. The TCRS should be interpreted as a screening "
        "instrument for identifying low-readiness learners, not as a diagnostic tool that "
        "pinpoints the source of calibration failure.")

    add_body_text(doc,
        "**Instructional designers** can use the process model to inform the design of "
        "AI-integrated learning environments that scaffold the calibration process\u2014for example, "
        "by including trust state reflection prompts (monitoring support), reliability indicators "
        "and comparison tools (evaluation support), and structured verification workflows "
        "(regulation support).")

    add_heading(doc, "Limitations and Future Directions", level=2)

    add_body_text(doc, "Several limitations should be acknowledged.")

    add_body_text(doc,
        "First, the TCRS measures readiness through self-report, not actual calibration "
        "accuracy. Self-report measures of metacognitive processes are subject to known validity "
        "ceilings, as individuals with poor metacognition tend to overestimate their metacognitive "
        "abilities (Dunning et al., 2003). Importantly, this validity ceiling is not uniform "
        "across the three subscales. Self-report validity is expected to be strongest for "
        "Calibration Action (CA-Ac), which asks about observable behavioral frequencies\u2014"
        "behaviors that are concrete and relatively easy to report accurately. Validity is "
        "expected to be moderate for Calibration Judgment (CA-Jd), which asks about evaluative "
        "tendencies that are somewhat introspectable but subject to social desirability. Validity "
        "is expected to be weakest for Calibration Awareness (CA-Aw), which asks about "
        "metacognitive awareness\u2014precisely the construct most vulnerable to the Dunning-Kruger "
        "effect. Future research should examine the relationship between TCRS scores and "
        "behavioral measures of trust calibration accuracy.")

    add_body_text(doc,
        "Second, the TCRS is likely most reliable as a screening tool for identifying "
        "low-readiness learners rather than as a confirmatory tool for high-readiness learners. "
        "Learners who report low calibration readiness are almost certainly low; however, "
        "learners who report high readiness may be overestimating their capacity. This asymmetric "
        "interpretive value should inform how the TCRS is used in practice.")

    add_body_text(doc,
        "Third, the cross-sectional design captures a snapshot of readiness at each process "
        "stage but cannot confirm the theorized sequential relationship among the stages. "
        "Longitudinal or experience-sampling studies are needed to test whether awareness "
        "temporally precedes judgment, and judgment precedes action, during actual AI-assisted "
        "learning episodes.")

    add_body_text(doc,
        "Fourth, the process model theorizes calibration as a primarily sequential process with "
        "recursive feedback, but the balance between sequential and parallel processing in "
        "practice remains an empirical question.")

    add_body_text(doc,
        "Fifth, the TCRS subscale scores should not be interpreted as indicators of general "
        "metacognitive ability (CA-Aw), general critical thinking (CA-Jd), or general behavioral "
        "agency (CA-Ac). Each subscale measures a domain-specific functional capacity within the "
        "trust calibration context. The process model draws selectively from three theoretical "
        "traditions as explanatory resources, not as constructs being measured in their entirety. "
        "Practitioners should avoid over-generalizing low subscale scores to broader psychological "
        "constructs: a low CA-Aw score identifies a learner who may not attend to their trust "
        "state when using AI, not a learner with deficient metacognition.")

    add_body_text(doc,
        "Sixth, the sample [describe demographic limitations]. Cross-cultural validation is a "
        "particularly important next step, as the theoretical traditions informing the TCRS "
        "carry unequal cultural assumptions. Bandura\u2019s (2001) self-regulation concept\u2014which "
        "informs the Calibration Action subscale\u2014is rooted in a Western, individualistic "
        "conception of self-directed action. Trust calibration in collectivist educational contexts may "
        "involve qualitatively different processes.")

    add_body_text(doc,
        "Sixth, the rapid evolution of AI capabilities means that the specific AI behaviors "
        "learners encounter will change over time. The TCRS is designed at the construct level "
        "(trust calibration readiness) rather than the technology level (specific AI tools), "
        "which should provide some stability, but ongoing monitoring of the scale\u2019s relevance "
        "is warranted.")

    add_heading(doc, "Conclusion", level=2)

    add_body_text(doc,
        "Trust calibration in educational AI is a process, not a trait. Learners who can "
        "appropriately calibrate their trust in AI systems\u2014monitoring their trust state, "
        "evaluating its appropriateness, and actively adjusting their reliance\u2014are better "
        "positioned to benefit from AI-enhanced learning without falling into the traps of "
        "overtrust or distrust. The Trust Calibration Readiness Scale provides a first "
        "instrument for measuring these three functional capacities, grounded in a process "
        "model that draws selectively on metacognitive monitoring, trust dynamics, and "
        "behavioral self-regulation as explanatory resources. "
        "As AI becomes an increasingly routine part of educational experience, the ability to "
        "calibrate one\u2019s trust in these systems may prove to be one of the most important "
        "literacies of the coming decade.")


def build_references(doc):
    """References section."""
    doc.add_page_break()
    add_heading(doc, "References", level=1)

    refs = [
        "Ajzen, I. (1991). The theory of planned behavior. *Organizational Behavior and Human Decision Processes, 50*(2), 179\u2013211. https://doi.org/10.1016/0749-5978(91)90020-T",
        "Bandura, A. (2001). Social cognitive theory: An agentic perspective. *Annual Review of Psychology, 52*, 1\u201326. https://doi.org/10.1146/annurev.psych.52.1.1",
        "Barzilai, S., & Zohar, A. (2014). Reconsidering personal epistemology as metacognition: An integrative review. *Educational Psychologist, 49*(3), 141\u2013167. https://doi.org/10.1080/00461520.2014.914823",
        "Carolus, A., Koch, M. J., Straka, S., Latoschik, M. E., & Wienrich, C. (2023). MAILS - Meta AI literacy scale: Development and testing of an AI literacy questionnaire based on well-founded competency models and psychological change- and meta-competencies. *Computers in Human Behavior: Artificial Humans, 1*(2), 100014. https://doi.org/10.1016/j.chbah.2023.100014",
        "Carolus, A., Augustin, Y., Markus, A., & Wienrich, C. (2024). Meta AI literacy scale: Further validation and development of a short version. *Heliyon, 10*(21), e39717. https://doi.org/10.1016/j.heliyon.2024.e39717",
        "Crowne, D. P., & Marlowe, D. (1960). A new scale of social desirability independent of psychopathology. *Journal of Consulting Psychology, 24*(4), 349\u2013354. https://doi.org/10.1037/h0047358",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. *MIS Quarterly, 13*(3), 319\u2013340. https://doi.org/10.2307/249008",
        "DeVellis, R. F., & Thorpe, C. T. (2022). *Scale development: Theory and applications* (5th ed.). Sage.",
        "Dunning, D., Johnson, K., Ehrlinger, J., & Kruger, J. (2003). Why people fail to recognize their own incompetence. *Current Directions in Psychological Science, 12*(3), 83\u201387. https://doi.org/10.1111/1467-8721.01235",
        "Efklides, A. (2006). Metacognition and affect: What can metacognitive experiences tell us about the learning process? *Educational Research Review, 1*(1), 3\u201314. https://doi.org/10.1016/j.edurev.2005.11.001",
        "Horn, J. L. (1965). A rationale and test for the number of factors in factor analysis. *Psychometrika, 30*, 179\u2013185. https://doi.org/10.1007/BF02289447",
        "Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. *Structural Equation Modeling, 6*(1), 1\u201355. https://doi.org/10.1080/10705519909540118",
        "Jian, J. Y., Bisantz, A. M., & Drury, C. G. (2000). Foundations for an empirically determined scale of trust in automated systems. *International Journal of Cognitive Ergonomics, 4*(1), 53\u201371. https://doi.org/10.1207/S15327566IJCE0401_04",
        "Kahneman, D. (2011). *Thinking, fast and slow.* Farrar, Straus and Giroux.",
        "Krosnick, J. A. (1999). Survey research. *Annual Review of Psychology, 50*, 537\u2013567. https://doi.org/10.1146/annurev.psych.50.1.537",
        "Lazarus, R. S., & Folkman, S. (1984). *Stress, appraisal, and coping*. Springer.",
        "Laupichler, M. C., Aster, A., Perschewski, J.-O., & Schleiss, J. (2023). Development of the Scale for the Assessment of Non-experts\u2019 AI Literacy (SNAIL). *Computers in Human Behavior Reports, 12*, 100338. https://doi.org/10.1016/j.chbr.2023.100338",
        "Lee, J. D., & See, K. A. (2004). Trust in automation: Designing for appropriate reliance. *Human Factors, 46*(1), 50\u201380. https://doi.org/10.1518/hfes.46.1.50.30392",
        "Lintner, C. (2024). A systematic review of AI literacy scales. *npj Science of Learning, 9*, 55. https://doi.org/10.1038/s41539-024-00267-7",
        "Mayer, R. C., Davis, J. H., & Schoorman, F. D. (1995). An integrative model of organizational trust. *Academy of Management Review, 20*(3), 709\u2013734. https://doi.org/10.5465/amr.1995.9508080335",
        "Ng, D. T. K., Wu, W., Leung, J. K. L., Chiu, T. K. F., & Chu, S. K. W. (2024). Design and validation of the AI literacy questionnaire: The quest to measure human literacy in artificial intelligence. *British Journal of Educational Technology, 55*(3), 1082\u20131104. https://doi.org/10.1111/bjet.13411",
        "Parasuraman, R., & Manzey, D. H. (2010). Complacency and bias in human use of automation: An attentional integration. *Human Factors, 52*(3), 381\u2013410. https://doi.org/10.1518/001872010X12829765906155",
        "Parasuraman, R., & Riley, V. (1997). Humans and automation: Use, misuse, disuse, abuse. *Human Factors, 39*(2), 230\u2013253. https://doi.org/10.1518/001872097778543886",
        "Pintrich, P. R., Smith, D. A. F., Garcia, T., & McKeachie, W. J. (1993). Reliability and predictive validity of the Motivated Strategies for Learning Questionnaire (MSLQ). *Educational and Psychological Measurement, 53*(3), 801\u2013813. https://doi.org/10.1177/0013164493053003024",
        "Pinski, M., et al. (2025). SAIL4ALL: Scale of AI literacy for all. *Humanities and Social Sciences Communications, 12*. https://doi.org/10.1057/s41599-025-04435-7",
        "Podsakoff, P. M., MacKenzie, S. B., Lee, J.-Y., & Podsakoff, N. P. (2003). Common method biases in behavioral research: A critical review of the literature and recommended remedies. *Journal of Applied Psychology, 88*(5), 879\u2013903. https://doi.org/10.1037/0021-9010.88.5.879",
        "Rodriguez, A., Reise, S. P., & Haviland, M. G. (2016). Evaluating bifactor models: Calculating and interpreting statistical indices. *Psychological Methods, 21*(2), 137\u2013150. https://doi.org/10.1037/met0000045",
        "Schraw, G., & Dennison, R. S. (1994). Assessing metacognitive awareness. *Contemporary Educational Psychology, 19*(4), 460\u2013475. https://doi.org/10.1006/ceps.1994.1033",
        "Schwarzer, R., & Jerusalem, M. (1995). Generalized Self-Efficacy Scale. In J. Weinman, S. Wright, & M. Johnston (Eds.), *Measures in health psychology: A user\u2019s portfolio. Causal and control beliefs* (pp. 35\u201337). NFER-Nelson.",
        "Sosu, E. M. (2013). The development and psychometric validation of a Critical Thinking Disposition Scale. *Thinking Skills and Creativity, 9*, 107\u2013119. https://doi.org/10.1016/j.tsc.2012.09.002",
        "Sperber, D., Cl\u00E9ment, F., Heintz, C., Mascaro, O., Mercier, H., Origgi, G., & Wilson, D. (2010). Epistemic vigilance. *Mind & Language, 25*(4), 359\u2013393. https://doi.org/10.1111/j.1468-0017.2010.01394.x",
        "Stone, N. J. (2000). Exploring the relationship between calibration and self-regulated learning. *Educational Psychology Review, 12*(4), 437\u2013475. https://doi.org/10.1023/A:1009084430926",
        "Watkins, M. W. (2018). Exploratory factor analysis: A guide to best practice. *Journal of Black Psychology, 44*(3), 219\u2013246. https://doi.org/10.1177/0095798418766959",
        "Winne, P. H., & Hadwin, A. F. (1998). Studying as self-regulated learning. In D. J. Hacker, J. Dunlosky, & A. C. Graesser (Eds.), *Metacognition in educational theory and practice* (pp. 277\u2013304). Erlbaum.",
    ]

    import re
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 2.0

        parts = re.split(r'(\*.*?\*)', ref)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                add_run(p, part[1:-1], italic=True)
            else:
                add_run(p, part)


def build_appendices(doc):
    """Appendix placeholders."""
    doc.add_page_break()
    add_heading(doc, "Appendix A: Trust Calibration Readiness Scale (TCRS) \u2014 Final Items", level=1)
    add_body_text(doc, "[TO BE INCLUDED: Final items after expert panel review and Study 1 EFA]",
                  first_line_indent=False)

    doc.add_page_break()
    add_heading(doc, "Appendix B: Supplementary Tables", level=1)
    add_body_text(doc, "[TO BE INCLUDED: EFA pattern matrix, CFA fit indices, correlation matrices]",
                  first_line_indent=False)


# ─── Main ───

def main():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Build document
    create_title_page(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_theoretical_framework(doc)
    build_present_study(doc)
    build_study1(doc)
    build_study2(doc)
    build_discussion(doc)
    build_references(doc)
    build_appendices(doc)

    # Save
    doc.save(OUTPUT_PATH)

    # Stats
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    img_count = sum(1 for p in doc.paragraphs for r in p.runs
                    if r._element.findall(qn('w:drawing')))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"Paragraphs: {para_count}, Tables: {table_count}, Images: {img_count}")


if __name__ == "__main__":
    main()
