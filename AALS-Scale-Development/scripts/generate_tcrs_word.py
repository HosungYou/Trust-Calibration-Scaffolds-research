#!/usr/bin/env python3
"""
Generate APA 7th Edition Word document for:
"Trust Calibration in Educational AI: A Process Model and Scale Development"
(TCRS manuscript draft v2)

Reads from draft_v2_TCRS.md and produces TCRS_Theoretical_Background_APA7.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
MANUSCRIPT_DIR = os.path.join(PROJECT_DIR, "manuscript")
FIGURES_DIR = os.path.join(PROJECT_DIR, "figures")
SOURCE_MD = os.path.join(MANUSCRIPT_DIR, "draft_v2_TCRS.md")
OUTPUT_PATH = os.path.join(MANUSCRIPT_DIR, "TCRS_Theoretical_Background_APA7.docx")

FIGURE_FILE = "Figure_1_Trust_Calibration_Process_Model_v2.png"


# ─── Helpers ───

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_styled_run(paragraph, text, bold=False, italic=False, size=12,
                   font_name="Times New Roman"):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def process_inline_formatting(paragraph, text, size=12):
    """Parse bold (**), italic (*), and dash formatting."""
    text = text.replace("---", "\u2014")
    text = text.replace("--", "\u2013")

    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)

    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            add_styled_run(paragraph, part[3:-3], bold=True, italic=True, size=size)
        elif part.startswith('**') and part.endswith('**'):
            add_styled_run(paragraph, part[2:-2], bold=True, size=size)
        elif part.startswith('*') and part.endswith('*'):
            add_styled_run(paragraph, part[1:-1], italic=True, size=size)
        else:
            add_styled_run(paragraph, part, size=size)


def add_body_paragraph(doc, text, first_line_indent=True):
    """Add an APA-style body paragraph with double spacing."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    process_inline_formatting(p, text)
    return p


def add_heading_apa(doc, text, level=1):
    """Add an APA-style heading."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=2.0)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_run(p, text, bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_styled_run(p, text, bold=True)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_styled_run(p, text, bold=True, italic=True)
    return p


def add_figure(doc, image_path, figure_num, caption_text):
    """Add a figure with APA 7th formatting."""
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_label, before=12, after=0, line_spacing=2.0)
    add_styled_run(p_label, f"Figure {figure_num}", bold=True, italic=True)

    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_caption, before=0, after=6, line_spacing=2.0)
    process_inline_formatting(p_caption, caption_text)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_img, before=0, after=12, line_spacing=1.0)
    run = p_img.add_run()
    run.add_picture(image_path, width=Inches(6.5))

    return p_img


# ─── Parse Markdown ───

def parse_markdown(filepath):
    """Parse the markdown file into structured sections."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content


def process_markdown_to_doc(doc, content):
    """Process markdown content into Word document paragraphs."""
    lines = content.strip().split('\n')
    i = 0
    figure_inserted = False
    in_references = False

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip YAML-like frontmatter and title lines
        if line == '---':
            i += 1
            continue

        # Skip the main title (# heading) - handled by title page
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue

        # Skip working title block
        if line.startswith('**Working Title'):
            while i < len(lines) and lines[i].strip():
                i += 1
            continue

        # Handle ## headings (Level 1 APA)
        if line.startswith('## '):
            heading_text = line[3:].strip()
            # Clean up numbering
            heading_text = re.sub(r'^\d+\.\s+', '', heading_text)

            if heading_text == 'References':
                in_references = True
                doc.add_page_break()
                add_heading_apa(doc, 'References', level=1)
                i += 1
                continue

            if heading_text == 'Abstract':
                # Abstract handled specially
                add_heading_apa(doc, 'Abstract', level=1)
                i += 1
                continue

            # Regular section heading
            add_heading_apa(doc, heading_text, level=1)
            i += 1
            continue

        # Handle ### headings (Level 2 APA)
        if line.startswith('### '):
            heading_text = line[4:].strip()
            heading_text = re.sub(r'^\d+\.\d+\s+', '', heading_text)
            add_heading_apa(doc, heading_text, level=2)
            i += 1
            continue

        # Handle #### headings (Level 3 APA)
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_text = re.sub(r'^\d+\.\d+\.\d+\s+', '', heading_text)
            add_heading_apa(doc, heading_text, level=3)
            i += 1
            continue

        # Handle figure placeholder
        if '[INSERT FIGURE 1 HERE]' in line:
            fig_path = os.path.join(FIGURES_DIR, FIGURE_FILE)
            if os.path.exists(fig_path):
                add_figure(doc, fig_path, 1,
                    "The Trust Calibration Process Model. Panel (a) depicts the trust "
                    "calibration space where the diagonal line represents perfect calibration "
                    "(trust matches AI reliability). Over-trust (above line) leads to misuse; "
                    "distrust (below line) leads to disuse. Metacognition enables learners to "
                    "recognize their position in this space, while agency provides the mechanism "
                    "to adjust toward calibration. Note: The TCRS measures readiness to engage "
                    "in this calibration process, not the learner\u2019s position in this space. "
                    "Panel (b) shows the process model with recursive feedback: Metacognition "
                    "(foundation) \u2192 Trust Evaluation (variable being calibrated) \u2192 Human "
                    "Agency (regulation mechanism) \u2192 Calibration (outcome), with a feedback "
                    "loop from Calibration back to Metacognition. Corresponding scale subscales "
                    "are mapped to each component.")
                figure_inserted = True
            i += 1
            continue

        # Skip figure caption in markdown (already handled above)
        if line.startswith('*Figure 1.*') or line.startswith('*Figure 1.'):
            # Skip entire caption paragraph
            while i < len(lines) and lines[i].strip():
                i += 1
            continue

        # Handle Keywords line
        if line.startswith('**Keywords:**'):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
            add_styled_run(p, "Keywords: ", italic=True)
            kw_text = line.replace('**Keywords:**', '').strip()
            add_styled_run(p, kw_text)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s+', line):
            bullet_text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

            # Get the number
            num_match = re.match(r'^(\d+)\.\s+', line)
            num = num_match.group(1) if num_match else ''
            add_styled_run(p, f"{num}. ")
            process_inline_formatting(p, bullet_text)
            i += 1
            continue

        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = re.sub(r'^[-*]\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
            add_styled_run(p, "\u2022 ")
            process_inline_formatting(p, bullet_text)
            i += 1
            continue

        # Handle references (hanging indent)
        if in_references:
            # Reference entries
            if line and not line.startswith('#'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
                process_inline_formatting(p, line)
            i += 1
            continue

        # Handle Appendix headings
        if line.startswith('## Appendix'):
            doc.add_page_break()
            heading_text = line[3:].strip()
            add_heading_apa(doc, heading_text, level=1)
            i += 1
            continue

        # Skip markdown table lines
        if '|' in line and line.startswith('|'):
            while i < len(lines) and '|' in lines[i].strip():
                i += 1
            continue

        # Regular paragraph
        if line and not line.startswith('#'):
            # Skip placeholder lines
            if line.startswith('[TO BE') or line.startswith('[RESULTS TO BE') or \
               line.startswith('[NUMBER]') or line.startswith('[Report') or \
               line.startswith('[describe'):
                add_body_paragraph(doc, line, first_line_indent=True)
                i += 1
                continue

            add_body_paragraph(doc, line, first_line_indent=True)

        i += 1


# ─── Title Page ───

def create_title_page(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Trust Calibration in Educational AI:", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "A Process Model and Scale Development", bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Hosung You")

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Department of Artificial Intelligence Convergence Education")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Sungkyunkwan University")

    # Author note
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    add_styled_run(p, "Author Note", bold=True)

    add_body_paragraph(doc,
        "Correspondence concerning this article should be addressed to "
        "Hosung You, Department of Artificial Intelligence Convergence Education, "
        "Sungkyunkwan University, Seoul, South Korea.")

    doc.add_page_break()


# ─── Main Document Builder ───

def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title page
    create_title_page(doc)

    # Read and process markdown
    content = parse_markdown(SOURCE_MD)
    process_markdown_to_doc(doc, content)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

    # Report stats
    para_count = len(doc.paragraphs)
    img_count = sum(1 for p in doc.paragraphs for r in p.runs
                    if r._element.findall(qn('w:drawing')))
    print(f"Paragraphs: {para_count}, Images: {img_count}")


if __name__ == "__main__":
    build_document()
