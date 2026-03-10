#!/usr/bin/env python3
"""
G6 Academic Style Humanizer - Balanced Mode
Transforms AI patterns to natural academic prose in the TCRS manuscript.
Applies all Priority 1-3 fixes from G5 audit results.

CRITICAL RULES:
- ZERO em dashes in output
- Preserve all citations exactly
- Preserve all technical terms (CA-Aw, CA-Jd, CA-Ac, TCRS, etc.)
- Preserve en dashes in compound terms
"""

import os
import sys
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
MANUSCRIPT_DIR = os.path.join(PROJECT_DIR, "manuscript")
INPUT_PATH = os.path.join(MANUSCRIPT_DIR, "TCRS_Theoretical_Background_APA7.docx")
OUTPUT_PATH = INPUT_PATH  # Overwrite in place

EM_DASH = "\u2014"
EN_DASH = "\u2013"

changes_log = []


def log_change(fix_id, para_idx, before_snippet, after_snippet):
    changes_log.append({
        "fix_id": fix_id,
        "para_idx": para_idx,
        "before": before_snippet[:150] + ("..." if len(before_snippet) > 150 else ""),
        "after": after_snippet[:150] + ("..." if len(after_snippet) > 150 else ""),
    })


def get_para_text(para):
    """Get full text from a paragraph, combining all runs."""
    return "".join(run.text for run in para.runs)


def replace_in_runs(para, old_text, new_text):
    """
    Replace old_text with new_text in a paragraph's runs.

    Strategy: Build a map of character positions to runs. Find where old_text
    starts and ends. Replace only the affected runs, preserving formatting
    on runs outside the replacement zone.

    Returns True if replacement was made.
    """
    full = get_para_text(para)
    pos = full.find(old_text)
    if pos == -1:
        return False

    end_pos = pos + len(old_text)

    # Build character-to-run mapping
    runs = list(para.runs)
    if not runs:
        return False

    # Find which runs are affected
    char_offset = 0
    start_run = None
    start_char_in_run = None
    end_run = None
    end_char_in_run = None

    for i, run in enumerate(runs):
        run_start = char_offset
        run_end = char_offset + len(run.text)

        if start_run is None and run_end > pos:
            start_run = i
            start_char_in_run = pos - run_start

        if end_run is None and run_end >= end_pos:
            end_run = i
            end_char_in_run = end_pos - run_start
            break

        char_offset = run_end

    if start_run is None or end_run is None:
        return False

    if start_run == end_run:
        # Replacement is within a single run - simple case
        run = runs[start_run]
        run.text = run.text[:start_char_in_run] + new_text + run.text[end_char_in_run:]
    else:
        # Replacement spans multiple runs
        # Put replacement text in start run (keeping prefix), clear middle runs,
        # keep suffix of end run
        prefix = runs[start_run].text[:start_char_in_run]
        suffix = runs[end_run].text[end_char_in_run:]

        runs[start_run].text = prefix + new_text

        # Clear middle runs
        for j in range(start_run + 1, end_run):
            runs[j].text = ""

        # Set end run to just the suffix
        if end_run != start_run:
            runs[end_run].text = suffix

    return True


def find_para(paragraphs, fragment, start=0):
    """Find index of paragraph containing text fragment."""
    for i in range(start, len(paragraphs)):
        if fragment in get_para_text(paragraphs[i]):
            return i
    return None


def main():
    print(f"Reading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        print(f"ERROR: File not found: {INPUT_PATH}")
        sys.exit(1)

    doc = Document(INPUT_PATH)
    paras = doc.paragraphs

    print(f"Total paragraphs: {len(paras)}")

    # Debug: print first 20 paragraphs to understand structure
    print("\n--- First 30 paragraph previews ---")
    for i in range(min(30, len(paras))):
        t = get_para_text(paras[i])
        if t.strip():
            print(f"  [{i:3d}] {t[:100]}...")
    print("---\n")

    # =====================================================================
    # FIX P1: Remove "It is important to note" (all instances)
    # =====================================================================
    print("=" * 60)
    print("FIX P1: Removing 'It is important to note'")
    print("=" * 60)

    # P1-1: section 2.2 "It is important to note that the process model describes deliberate trust calibration"
    idx = find_para(paras, "It is important to note that the process model describes deliberate trust")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "It is important to note that the process model describes deliberate trust "
            "calibration",
            "The process model describes deliberate trust calibration only")
        if ok:
            log_change("P1-1", idx,
                "It is important to note that the process model describes...",
                "The process model describes deliberate trust calibration only...")
            print(f"  [P1-1] Para {idx}: OK")
    else:
        print("  [P1-1] NOT FOUND")

    # P1-2: section 2.3 "It is important to note that the TCRS measures learners' readiness"
    idx = find_para(paras, "It is important to note that the TCRS measures learners")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "It is important to note that the TCRS measures learners\u2019 readiness to engage in "
            "the calibration process, not their actual position in the trust calibration space.",
            "The TCRS measures learners\u2019 readiness to engage in "
            "the calibration process, not their actual position in the trust calibration space.")
        if ok:
            log_change("P1-2", idx,
                "It is important to note that the TCRS measures...",
                "The TCRS measures learners\u2019 readiness...")
            print(f"  [P1-2] Para {idx}: OK")
    else:
        print("  [P1-2] NOT FOUND")

    # P1-3: section 2.4 "It is important to note that these three subscales"
    idx = find_para(paras, "It is important to note that these three subscales measure components")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "It is important to note that these three subscales measure components of a theorized "
            "process sequence, yet they are assessed as parallel subscales within one instrument.",
            "These three subscales measure components of a theorized "
            "process sequence, yet they are assessed as parallel subscales within one instrument.")
        if ok:
            log_change("P1-3", idx,
                "It is important to note that these three subscales...",
                "These three subscales measure components...")
            print(f"  [P1-3] Para {idx}: OK")
    else:
        print("  [P1-3] NOT FOUND")

    # P1-4,5,6: "Importantly, CA-Aw/CA-Jd/CA-Ac measures..." (these are inside formatted paragraphs)
    # CA-Aw
    idx = find_para(paras, "Importantly, CA-Aw measures trust state monitoring")
    if idx is not None:
        # The text after bold run starts with " captures the monitoring capacity..."
        # "Importantly," appears later. Find it in runs.
        ok = replace_in_runs(paras[idx],
            "Importantly, "
            "CA-Aw measures trust state monitoring in the context of AI interaction; "
            "it does not claim to measure metacognitive awareness in general, nor does a low CA-Aw "
            "score indicate a general metacognitive deficit.",
            "CA-Aw measures trust state monitoring in the context of AI interaction, not "
            "metacognitive awareness in general.")
        if ok:
            log_change("P1-4", idx,
                "Importantly, CA-Aw measures...it does not claim...",
                "CA-Aw measures...not metacognitive awareness in general.")
            print(f"  [P1-4] Para {idx}: OK")
    else:
        # Try alternative: "Importantly," might be at the start of a run after bold
        idx = find_para(paras, "CA-Aw measures trust state monitoring")
        if idx is not None:
            t = get_para_text(paras[idx])
            if "Importantly" in t:
                ok = replace_in_runs(paras[idx],
                    "Importantly, "
                    "CA-Aw measures trust state monitoring in the context of AI interaction; "
                    "it does not claim to measure metacognitive awareness in general, nor does a low CA-Aw "
                    "score indicate a general metacognitive deficit.",
                    "CA-Aw measures trust state monitoring in the context of AI interaction, not "
                    "metacognitive awareness in general.")
                if ok:
                    log_change("P1-4", idx, "Importantly, CA-Aw...", "CA-Aw measures...")
                    print(f"  [P1-4] Para {idx}: OK (alt)")
            else:
                print(f"  [P1-4] Found para {idx} but no 'Importantly'")
        else:
            print("  [P1-4] NOT FOUND")

    # CA-Jd
    idx = find_para(paras, "Importantly, CA-Jd measures trust")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "Importantly, CA-Jd measures trust "
            "appropriateness evaluation in the AI context; it does not claim to measure general "
            "critical thinking or evaluative judgment.",
            "CA-Jd measures trust "
            "appropriateness evaluation in the AI context, not general "
            "critical thinking or evaluative judgment.")
        if ok:
            log_change("P1-5", idx,
                "Importantly, CA-Jd measures...it does not claim...",
                "CA-Jd measures...not general critical thinking...")
            print(f"  [P1-5] Para {idx}: OK")
    else:
        idx = find_para(paras, "CA-Jd measures trust appropriateness evaluation")
        if idx is not None:
            t = get_para_text(paras[idx])
            if "Importantly" in t:
                ok = replace_in_runs(paras[idx],
                    "Importantly, CA-Jd measures trust "
                    "appropriateness evaluation in the AI context; it does not claim to measure general "
                    "critical thinking or evaluative judgment.",
                    "CA-Jd measures trust "
                    "appropriateness evaluation in the AI context, not general "
                    "critical thinking or evaluative judgment.")
                if ok:
                    log_change("P1-5", idx, "Importantly, CA-Jd...", "CA-Jd measures...")
                    print(f"  [P1-5] Para {idx}: OK (alt)")
        else:
            print("  [P1-5] NOT FOUND")

    # CA-Ac
    idx = find_para(paras, "Importantly, CA-Ac measures trust regulation")
    if idx is not None:
        # The exact text from the generation script has "Importantly, \nCA-Ac" because
        # the "Importantly, " was at end of one line in the script string.
        # In the actual doc it's continuous.
        ok = replace_in_runs(paras[idx],
            "Importantly, "
            "CA-Ac measures trust regulation behaviors in the AI context; it does not claim to "
            "measure general behavioral agency or self-regulation capacity.",
            "CA-Ac measures trust regulation behaviors in the AI context, not "
            "general behavioral agency or self-regulation capacity.")
        if ok:
            log_change("P1-6", idx,
                "Importantly, CA-Ac measures...it does not claim...",
                "CA-Ac measures...not general behavioral agency...")
            print(f"  [P1-6] Para {idx}: OK")
    else:
        idx = find_para(paras, "CA-Ac measures trust regulation behaviors")
        if idx is not None:
            t = get_para_text(paras[idx])
            if "Importantly" in t:
                ok = replace_in_runs(paras[idx],
                    "Importantly, "
                    "CA-Ac measures trust regulation behaviors in the AI context; it does not claim to "
                    "measure general behavioral agency or self-regulation capacity.",
                    "CA-Ac measures trust regulation behaviors in the AI context, not "
                    "general behavioral agency or self-regulation capacity.")
                if ok:
                    log_change("P1-6", idx, "Importantly, CA-Ac...", "CA-Ac measures...")
                    print(f"  [P1-6] Para {idx}: OK (alt)")
        else:
            print("  [P1-6] NOT FOUND")

    print()

    # =====================================================================
    # FIX X1: Add consolidated domain-specificity sentence
    # =====================================================================
    print("=" * 60)
    print("FIX X1: Adding consolidated domain-specificity sentence")
    print("=" * 60)

    # Find the "These three subscales" paragraph (after P1-3 fix)
    idx = find_para(paras, "three subscales measure components of a theorized")
    if idx is not None:
        t = get_para_text(paras[idx])
        if "Each subscale is domain-specific" not in t:
            consolidation = (
                " Each subscale is domain-specific: low scores on CA-Aw, CA-Jd, or CA-Ac "
                "identify calibration-specific deficits, not general deficits in "
                "metacognition, critical thinking, or self-regulation, respectively."
            )
            # Append to last run
            last_run = paras[idx].runs[-1]
            last_run.text += consolidation
            log_change("X1", idx, "[appended]",
                "...Each subscale is domain-specific: low scores on CA-Aw, CA-Jd, or CA-Ac identify calibration-specific deficits...")
            print(f"  [X1] Para {idx}: Appended consolidated sentence")
    else:
        print("  [X1] Target paragraph NOT FOUND")
    print()

    # =====================================================================
    # FIX X3: Remove "follows established precedent" / "follows the precedent"
    # =====================================================================
    print("=" * 60)
    print("FIX X3: Removing 'established precedent' phrases")
    print("=" * 60)

    # X3-1: section 2.1 - long em-dash parenthetical with "follows established precedent"
    idx = find_para(paras, "follows established precedent")
    if idx is not None:
        t = get_para_text(paras[idx])
        if "multi-source theoretical models" in t:
            # This has em dashes: "This approach\u2014assembling...construct\u2014follows established precedent in multi-source..."
            ok = replace_in_runs(paras[idx],
                "This approach" + EM_DASH + "assembling functional components from different "
                "theoretical traditions to explain a novel target construct" + EM_DASH + "follows established precedent "
                "in multi-source theoretical models such as the Technology Acceptance Model (Davis, 1989), "
                "the Theory of Planned Behavior (Ajzen, 1991), and the Transactional Model of Stress "
                "(Lazarus & Folkman, 1984).",
                "Assembling functional components from different "
                "theoretical traditions to explain a novel target construct is a strategy shared by "
                "the Technology Acceptance Model (Davis, 1989), "
                "the Theory of Planned Behavior (Ajzen, 1991), and the Transactional Model of Stress "
                "(Lazarus & Folkman, 1984).")
            if ok:
                log_change("X3-1", idx,
                    "This approach\u2014assembling...follows established precedent in multi-source...",
                    "Assembling functional components...is a strategy shared by...")
                print(f"  [X3-1] Para {idx}: OK")
            else:
                print(f"  [X3-1] Para {idx}: replace_in_runs failed. Text preview: {t[:200]}")
    else:
        print("  [X3-1] NOT FOUND")

    # X3-2: section 2.4 - "This approach follows established precedent. The Motivated Strategies..."
    idx = find_para(paras, "This approach follows established precedent")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "This approach follows established precedent. The Motivated Strategies for Learning "
            "Questionnaire",
            "The Motivated Strategies for Learning "
            "Questionnaire")
        if ok:
            log_change("X3-2", idx,
                "This approach follows established precedent. The Motivated Strategies...",
                "The Motivated Strategies for Learning Questionnaire...")
            print(f"  [X3-2] Para {idx}: OK")
        else:
            print(f"  [X3-2] Para {idx}: replace failed")
    else:
        print("  [X3-2] NOT FOUND")

    # X3-3: section 3.2 - "follows the precedent of the MSLQ"
    idx = find_para(paras, "follows the precedent of the MSLQ")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "This mixed-format approach "
            "follows the precedent of the MSLQ, which uses different response formats across "
            "subscales.",
            "The MSLQ similarly uses different response formats across "
            "subscales.")
        if ok:
            log_change("X3-3", idx,
                "This mixed-format approach follows the precedent of the MSLQ...",
                "The MSLQ similarly uses different response formats across subscales.")
            print(f"  [X3-3] Para {idx}: OK")
    else:
        print("  [X3-3] NOT FOUND")
    print()

    # =====================================================================
    # FIX P5: Replace weak conclusion
    # =====================================================================
    print("=" * 60)
    print("FIX P5: Replacing weak conclusion")
    print("=" * 60)

    idx = find_para(paras, "may prove to be one of the most important literacies")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "the ability to "
            "calibrate one\u2019s trust in these systems may prove to be one of the most important "
            "literacies of the coming decade.",
            "calibrating one\u2019s trust in these systems constitutes a core literacy "
            "that educators can no longer afford to leave unmeasured.")
        if ok:
            log_change("P5", idx,
                "...may prove to be one of the most important literacies of the coming decade.",
                "...constitutes a core literacy that educators can no longer afford to leave unmeasured.")
            print(f"  [P5] Para {idx}: OK")
        else:
            print(f"  [P5] Para {idx}: replace failed")
    else:
        print("  [P5] NOT FOUND")
    print()

    # =====================================================================
    # FIX S2: Insert short sentences (7 punchy sentences)
    # =====================================================================
    print("=" * 60)
    print("FIX S2: Inserting short sentences for length variation")
    print("=" * 60)

    # S2-1: After "83.5% fell below a calibration threshold"
    idx = find_para(paras, "83.5% fell below a calibration threshold")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "This \u201Ccalibration gap\u201D is not merely an "
            "academic concern; it has direct implications",
            "The gap is pervasive. It has direct implications")
        if ok:
            log_change("S2-1", idx,
                'This "calibration gap" is not merely an academic concern...',
                "The gap is pervasive. It has direct implications...")
            print(f"  [S2-1] Para {idx}: OK")
    else:
        print("  [S2-1] NOT FOUND")

    # S2-2: "the field lacks a measurement instrument"
    idx = find_para(paras, "the field lacks a measurement instrument")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "do not address the specific cognitive and behavioral processes "
            "involved in trust calibration. Trust measurement instruments",
            "do not address the specific cognitive and behavioral processes "
            "involved in trust calibration. No such instrument exists. Trust measurement instruments")
        if ok:
            log_change("S2-2", idx, "...do not address...Trust measurement...",
                "...do not address...No such instrument exists. Trust measurement...")
            print(f"  [S2-2] Para {idx}: OK")
    else:
        print("  [S2-2] NOT FOUND")

    # S2-3: "no existing scale addresses three critical elements"
    idx = find_para(paras, "no existing scale addresses three critical elements")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "Despite this proliferation, no existing scale addresses",
            "Despite this proliferation, none measures trust calibration. No existing scale addresses")
        if ok:
            log_change("S2-3", idx, "Despite this proliferation, no existing scale...",
                "Despite this proliferation, none measures trust calibration. No existing scale...")
            print(f"  [S2-3] Para {idx}: OK")
    else:
        print("  [S2-3] NOT FOUND")

    # S2-4: "calibration cannot begin" in Monitoring Capacity
    idx = find_para(paras, "calibration cannot begin")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "calibration cannot begin. This monitoring function",
            "calibration cannot begin. Awareness comes first. This monitoring function")
        if ok:
            log_change("S2-4", idx, "...cannot begin. This monitoring function...",
                "...cannot begin. Awareness comes first. This monitoring function...")
            print(f"  [S2-4] Para {idx}: OK")
    else:
        print("  [S2-4] NOT FOUND")

    # S2-5: "passive trust holder into an active trust regulator"
    idx = find_para(paras, "passive trust holder into an active trust regulator")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "Regulation transforms the learner from a "
            "passive trust holder into an active trust regulator. This regulatory function",
            "Regulation transforms the learner from a "
            "passive trust holder into an active trust regulator. Without action, awareness stays inert. This regulatory function")
        if ok:
            log_change("S2-5", idx, "...active trust regulator. This regulatory function...",
                "...active trust regulator. Without action, awareness stays inert. This regulatory function...")
            print(f"  [S2-5] Para {idx}: OK")
    else:
        print("  [S2-5] NOT FOUND")

    # S2-6: "it is about having the right amount of trust"
    idx = find_para(paras, "it is about having the")
    if idx is not None:
        t = get_para_text(paras[idx])
        if "right amount" in t:
            ok = replace_in_runs(paras[idx],
                "for the situation. A learner",
                "for the situation. Direction matters, not magnitude. A learner")
            if ok:
                log_change("S2-6", idx, "...for the situation. A learner...",
                    "...for the situation. Direction matters, not magnitude. A learner...")
                print(f"  [S2-6] Para {idx}: OK")
    else:
        print("  [S2-6] NOT FOUND")

    # S2-7: Conclusion "Trust calibration...is a process, not a trait"
    idx = find_para(paras, "Trust calibration in educational AI is a process, not a trait")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "Trust calibration in educational AI is a process, not a trait. Learners",
            "Trust calibration in educational AI is a process, not a trait. The distinction matters. Learners")
        if ok:
            log_change("S2-7", idx, "...is a process, not a trait. Learners...",
                "...is a process, not a trait. The distinction matters. Learners...")
            print(f"  [S2-7] Para {idx}: OK")
    else:
        print("  [S2-7] NOT FOUND")
    print()

    # =====================================================================
    # FIX S6/S8: Vary subscale paragraph structure in section 2.4
    # =====================================================================
    print("=" * 60)
    print("FIX S6/S8: Varying subscale paragraph structure")
    print("=" * 60)

    # CA-Jd paragraph: change "captures the evaluative capacity:" to different structure
    idx = find_para(paras, "Calibration Judgment (CA-Jd)")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "captures the evaluative capacity: the learner\u2019s "
            "tendency to evaluate whether their trust level is appropriate for the situation, "
            "to detect discrepancies between trust and reliability, and to consider contextual "
            "factors that should moderate their trust. This subscale is informed by trust dynamics "
            "research (Lee & See, 2004) and the concept of epistemic vigilance",
            "addresses a different question: Is my trust appropriate here? This subscale taps the learner\u2019s "
            "tendency to evaluate whether their trust level fits the situation, "
            "to detect discrepancies between trust and reliability, and to weigh contextual "
            "factors that should moderate their trust. Its theoretical grounding lies in trust dynamics "
            "research (Lee & See, 2004) and the concept of epistemic vigilance")
        if ok:
            log_change("S6-1", idx,
                "...captures the evaluative capacity: the learner\u2019s tendency...",
                "...addresses a different question: Is my trust appropriate here?...")
            print(f"  [S6-1] Para {idx}: OK (CA-Jd varied)")
        else:
            print(f"  [S6-1] Para {idx}: replace failed")

    # CA-Ac paragraph: change structure
    idx = find_para(paras, "Calibration Action (CA-Ac)")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "captures the regulatory capacity: the learner\u2019s "
            "active behaviors to verify AI outputs, adjust their reliance based on evidence, and "
            "employ deliberate strategies for trust regulation. This subscale is informed by the",
            "shifts from cognition to behavior. It captures the learner\u2019s "
            "active efforts to verify AI outputs, adjust reliance based on evidence, and "
            "employ deliberate strategies for trust regulation. The subscale draws on the")
        if ok:
            log_change("S6-2", idx,
                "...captures the regulatory capacity: the learner\u2019s active behaviors...",
                "...shifts from cognition to behavior. It captures the learner\u2019s active efforts...")
            print(f"  [S6-2] Para {idx}: OK (CA-Ac varied)")
        else:
            print(f"  [S6-2] Para {idx}: replace failed")
    print()

    # =====================================================================
    # FIX S9: Section 1.3 opening - taxonomy to argument
    # =====================================================================
    print("=" * 60)
    print("FIX S9: Section 1.3 opening - taxonomy to argument")
    print("=" * 60)

    idx = find_para(paras, "The field of AI literacy measurement has grown rapidly")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The field of AI literacy measurement has grown rapidly, with a systematic review "
            "identifying 16 distinct scales validated across 22 studies (Lintner, 2024). These "
            "instruments represent important advances but leave the trust calibration domain "
            "unaddressed. Table 2 summarizes the major existing scales and their coverage of trust "
            "calibration elements.",
            "Sixteen validated AI literacy scales now exist (Lintner, 2024), yet none addresses "
            "what learners must do to calibrate their trust. "
            "Table 2 summarizes these instruments and their coverage of trust "
            "calibration elements.")
        if ok:
            log_change("S9", idx,
                "The field of AI literacy measurement has grown rapidly...",
                "Sixteen validated AI literacy scales now exist (Lintner, 2024), yet none addresses...")
            print(f"  [S9] Para {idx}: OK")
        else:
            print(f"  [S9] Para {idx}: replace failed")
    else:
        print("  [S9] NOT FOUND")
    print()

    # =====================================================================
    # FIX L1: Reduce hedging in Limitations
    # =====================================================================
    print("=" * 60)
    print("FIX L1: Reducing hedging in Limitations")
    print("=" * 60)

    # L1-1: "Importantly, this validity ceiling is not uniform"
    idx = find_para(paras, "Self-report measures of metacognitive processes")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "Importantly, this validity ceiling is not uniform",
            "This validity ceiling varies")
        if ok:
            log_change("L1-1", idx, "Importantly, this validity ceiling is not uniform", "This validity ceiling varies")
            print(f"  [L1-1] Para {idx}: OK")

        ok = replace_in_runs(paras[idx],
            "Self-report validity is expected to be strongest for",
            "Self-report validity is strongest for")
        if ok:
            log_change("L1-2", idx, "...is expected to be strongest for", "...is strongest for")
            print(f"  [L1-2] Para {idx}: OK")

        ok = replace_in_runs(paras[idx],
            "Validity is "
            "expected to be moderate for",
            "Validity is moderate for")
        if ok:
            log_change("L1-3", idx, "Validity is expected to be moderate for", "Validity is moderate for")
            print(f"  [L1-3] Para {idx}: OK")

        ok = replace_in_runs(paras[idx],
            "Validity "
            "is expected to be weakest for",
            "Validity "
            "is weakest for")
        if ok:
            log_change("L1-4", idx, "Validity is expected to be weakest for", "Validity is weakest for")
            print(f"  [L1-4] Para {idx}: OK")

    # L1-5: "the TCRS is likely most reliable as a screening tool"
    idx = find_para(paras, "TCRS is likely most reliable as a screening tool")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "the TCRS is likely most reliable as a screening tool",
            "the TCRS functions best as a screening tool")
        if ok:
            log_change("L1-5", idx, "...is likely most reliable as...", "...functions best as...")
            print(f"  [L1-5] Para {idx}: OK")

        ok = replace_in_runs(paras[idx],
            "Learners who report low calibration readiness are almost certainly low; however, "
            "learners who report high readiness may be overestimating",
            "Learners who report low calibration readiness are low; "
            "learners who report high readiness may be overestimating")
        if ok:
            log_change("L1-6", idx, "...are almost certainly low; however, learners...", "...are low; learners...")
            print(f"  [L1-6] Para {idx}: OK")

    # L1-7: "captures a snapshot of readiness"
    idx = find_para(paras, "captures a snapshot of readiness")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "the cross-sectional design captures a snapshot of readiness at each process "
            "stage but cannot confirm the theorized sequential relationship among the stages.",
            "the cross-sectional design captures readiness at each process "
            "stage but cannot confirm the theorized sequence.")
        if ok:
            log_change("L1-7", idx, "...captures a snapshot...theorized sequential relationship among the stages.",
                "...captures readiness...theorized sequence.")
            print(f"  [L1-7] Para {idx}: OK")
    print()

    # =====================================================================
    # FIX D2: Vary paragraph openings
    # =====================================================================
    print("=" * 60)
    print("FIX D2: Varying paragraph openings")
    print("=" * 60)

    # D2-1
    idx = find_para(paras, "The integration of artificial intelligence into education has accelerated")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The integration of artificial intelligence into education has accelerated dramatically, "
            "with AI tools now supporting",
            "Artificial intelligence has entered education at speed. AI tools now support")
        if ok:
            log_change("D2-1", idx, "The integration of artificial intelligence...", "Artificial intelligence has entered education at speed...")
            print(f"  [D2-1] Para {idx}: OK")

    # D2-2
    idx = find_para(paras, "The scope of this problem is substantial")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The scope of this problem is substantial. A systematic review",
            "How widespread is the problem? A systematic review")
        if ok:
            log_change("D2-2", idx, "The scope of this problem is substantial...", "How widespread is the problem?...")
            print(f"  [D2-2] Para {idx}: OK")

    # D2-3
    idx = find_para(paras, "The process model consists of four components")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The process model consists of four components arranged in a theorized sequential "
            "relationship (see Figure 1).",
            "Four components, arranged in a theorized sequence, comprise the process model "
            "(see Figure 1).")
        if ok:
            log_change("D2-3", idx, "The process model consists of four components...", "Four components, arranged in a theorized sequence...")
            print(f"  [D2-3] Para {idx}: OK")

    # D2-4
    idx = find_para(paras, "The process model can be visualized through what we term")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The process model can be visualized through what we term the",
            "We represent the model spatially through the")
        if ok:
            log_change("D2-4", idx, "The process model can be visualized...", "We represent the model spatially...")
            print(f"  [D2-4] Para {idx}: OK")

    # D2-5
    idx = find_para(paras, "The process model yields three measurable constructs")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The process model yields three measurable constructs that correspond to the first "
            "three stages of the calibration process.",
            "Three measurable constructs follow from the first "
            "three stages of the calibration process.")
        if ok:
            log_change("D2-5", idx, "The process model yields three measurable...", "Three measurable constructs follow from...")
            print(f"  [D2-5] Para {idx}: OK")

    # D2-6
    idx = find_para(paras, "The trust calibration process model makes three theoretical contributions")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The trust calibration process model makes three theoretical contributions.",
            "Three theoretical contributions emerge from the trust calibration process model.")
        if ok:
            log_change("D2-6", idx, "The trust calibration process model makes...", "Three theoretical contributions emerge...")
            print(f"  [D2-6] Para {idx}: OK")

    # D2-7
    idx = find_para(paras, "The TCRS offers practical value for three audiences")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The TCRS offers practical value for three audiences:",
            "Three audiences stand to benefit from the TCRS:")
        if ok:
            log_change("D2-7", idx, "The TCRS offers practical value...", "Three audiences stand to benefit...")
            print(f"  [D2-7] Para {idx}: OK")

    # D2-8
    idx = find_para(paras, "The three functional capacities")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "The three functional capacities" + EM_DASH + "monitoring, evaluating, and regulating" + EM_DASH + "are not "
            "selected because they happen to map onto available theories.",
            "Why these three capacities? Monitoring, evaluating, and regulating are not "
            "selected because they happen to map onto available theories.")
        if ok:
            log_change("D2-8", idx, "The three functional capacities...are not selected...", "Why these three capacities? Monitoring, evaluating...")
            print(f"  [D2-8] Para {idx}: OK")
    print()

    # =====================================================================
    # FIX S3: Comma splice - "away from the optimal point" em dash to colon
    # =====================================================================
    print("=" * 60)
    print("FIX S3: Fixing comma splices / em-dash transitions")
    print("=" * 60)

    idx = find_para(paras, "away from the optimal point")
    if idx is not None:
        ok = replace_in_runs(paras[idx],
            "away from the optimal point" + EM_DASH + "upward toward",
            "away from the optimal point: upward toward")
        if ok:
            log_change("S3-1", idx, "...optimal point\u2014upward...", "...optimal point: upward...")
            print(f"  [S3-1] Para {idx}: OK")
        else:
            # Try without em dash (might have comma already)
            ok = replace_in_runs(paras[idx],
                "away from the optimal point, upward toward",
                "away from the optimal point: upward toward")
            if ok:
                print(f"  [S3-1] Para {idx}: OK (comma variant)")
    print()

    # =====================================================================
    # FIX M3: Remove redundant "see Table 1" (keep first, remove up to 3)
    # =====================================================================
    print("=" * 60)
    print("FIX M3: Removing redundant 'see Table 1' references")
    print("=" * 60)

    table1_refs = []
    for i, p in enumerate(paras):
        t = get_para_text(p)
        if "see Table 1" in t:
            table1_refs.append(i)

    print(f"  Found {len(table1_refs)} 'see Table 1' references at paragraphs: {table1_refs}")

    removed = 0
    for ref_idx in table1_refs[1:]:
        if removed >= 3:
            break
        ok = replace_in_runs(paras[ref_idx], " (see Table 1)", "")
        if ok:
            removed += 1
            log_change(f"M3-{removed}", ref_idx, "...(see Table 1)...", "...[removed]...")
            print(f"  [M3-{removed}] Para {ref_idx}: Removed")
        else:
            # Try without space before
            ok = replace_in_runs(paras[ref_idx], "(see Table 1)", "")
            if ok:
                removed += 1
                log_change(f"M3-{removed}", ref_idx, "...(see Table 1)...", "...[removed]...")
                print(f"  [M3-{removed}] Para {ref_idx}: Removed (no leading space)")
    print()

    # =====================================================================
    # GLOBAL: Remove ALL em dashes from body text
    # =====================================================================
    print("=" * 60)
    print("GLOBAL: Removing all em dashes from paragraphs")
    print("=" * 60)

    em_dash_total = 0
    em_dash_fixed = 0

    for i, p in enumerate(paras):
        for run in p.runs:
            count = run.text.count(EM_DASH)
            if count > 0:
                em_dash_total += count

                # Context-sensitive replacements
                txt = run.text

                # Parenthetical insertions: "X\u2014Y\u2014Z" -> "X (Y) Z" or "X, Y, Z"
                # Handle known patterns first

                # Pattern: "Trust calibration\u2014the correspondence...reliability\u2014is critical"
                txt = txt.replace(
                    "Trust calibration" + EM_DASH + "the correspondence",
                    "Trust calibration (the correspondence")
                txt = txt.replace(
                    "reliability" + EM_DASH + "is critical",
                    "reliability) is critical")

                # Pattern: "Appropriate trust\u2014calibrated trust"
                txt = txt.replace("Appropriate trust" + EM_DASH, "Appropriate trust, ")
                txt = txt.replace("trust" + EM_DASH + "calibrated trust that matches", "trust, calibrated trust that matches")

                # Pattern: "monitoring foundation\u2014without knowing"
                txt = txt.replace("monitoring foundation" + EM_DASH + "without knowing", "monitoring foundation, without knowing")

                # Pattern: "\u2014a phenomenon" -> ", a phenomenon"
                txt = txt.replace(EM_DASH + "a phenomenon", ", a phenomenon")

                # Pattern: "\u2014a design" -> ", a design"
                txt = txt.replace(EM_DASH + "a design", ", a design")

                # Pattern: "readiness framing\u2014measuring" -> "readiness framing: measuring"
                txt = txt.replace("readiness framing" + EM_DASH + "measuring", "readiness framing: measuring")

                # Pattern: "calibration\u2014what dual-process" -> "calibration, what dual-process"
                txt = txt.replace("calibration" + EM_DASH + "what dual-process", "calibration, what dual-process")

                # Pattern: "framework\u2014including...agency\u2014is"
                txt = txt.replace("framework" + EM_DASH + "including", "framework (including")
                txt = txt.replace("agency" + EM_DASH + "is beyond", "agency) is beyond")

                # Pattern: "process\u2014for example" -> "process, for example,"
                txt = txt.replace("process" + EM_DASH + "for example", "process, for example,")

                # Pattern: "behaviors\u2014behaviors" (clarifying appositive)
                txt = txt.replace("behaviors" + EM_DASH + "behaviors", "behaviors, behaviors")

                # Pattern: "awareness\u2014precisely" -> "awareness, precisely"
                txt = txt.replace("awareness" + EM_DASH + "precisely", "awareness, precisely")

                # Pattern: capacities\u2014monitoring -> capacities, monitoring (if not already handled)
                txt = txt.replace("capacities" + EM_DASH + "monitoring", "capacities (monitoring")
                txt = txt.replace("regulating" + EM_DASH + "are not", "regulating) are not")

                # Pattern: capacities\u2014 in general
                # Catch any "noun\u2014noun" with a generic comma replacement
                # But be careful not to break en dashes

                # Generic fallback: any remaining em dashes become commas
                remaining = txt.count(EM_DASH)
                if remaining > 0:
                    # Handle paired em dashes (parenthetical): find pairs
                    parts = txt.split(EM_DASH)
                    if len(parts) == 3:
                        # Likely a parenthetical: "A\u2014B\u2014C" -> "A (B) C"
                        txt = parts[0] + " (" + parts[1] + ") " + parts[2]
                        # Clean up double spaces
                        txt = txt.replace("  ", " ").replace(" )", ")").replace("( ", "(")
                    elif len(parts) == 2:
                        # Single em dash: use comma
                        txt = parts[0] + ", " + parts[1]
                    else:
                        # Multiple: just replace all with commas
                        txt = txt.replace(EM_DASH, ", ")

                new_count = txt.count(EM_DASH)
                em_dash_fixed += (count - new_count)
                run.text = txt

    # Final aggressive sweep
    final_em = 0
    for p in paras:
        for run in p.runs:
            if EM_DASH in run.text:
                final_em += run.text.count(EM_DASH)
                run.text = run.text.replace(EM_DASH, ", ")

    print(f"  Em dashes found: {em_dash_total}")
    print(f"  Em dashes fixed (context-aware): {em_dash_fixed}")
    print(f"  Em dashes remaining after sweep: {final_em} (all converted to commas)")

    # Also check table cells for em dashes
    table_em = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if EM_DASH in run.text:
                            table_em += run.text.count(EM_DASH)
                            run.text = run.text.replace(EM_DASH, ", ")
    if table_em:
        print(f"  Table cell em dashes removed: {table_em}")
    print()

    # =====================================================================
    # SAVE
    # =====================================================================
    print("=" * 60)
    print("SAVING DOCUMENT")
    print("=" * 60)

    doc.save(OUTPUT_PATH)
    print(f"  Saved to: {OUTPUT_PATH}")
    print()

    # =====================================================================
    # VERIFICATION
    # =====================================================================
    print("=" * 60)
    print("PRESERVATION VERIFICATION")
    print("=" * 60)

    doc2 = Document(OUTPUT_PATH)
    full_text = " ".join(get_para_text(p) for p in doc2.paragraphs)

    # Check for remaining em dashes
    em_remaining = full_text.count(EM_DASH)
    print(f"\n  Em dashes in body text: {em_remaining}")
    if em_remaining == 0:
        print("    [PASS] Zero em dashes")
    else:
        print("    [FAIL] Em dashes still present!")
        # Find where
        for i, p in enumerate(doc2.paragraphs):
            t = get_para_text(p)
            if EM_DASH in t:
                pos = t.find(EM_DASH)
                print(f"    Para {i}: ...{t[max(0,pos-30):pos+30]}...")

    # Check citations
    citations = [
        "Lee & See, 2004", "Winne & Hadwin, 1998", "Bandura, 2001",
        "Parasuraman & Riley, 1997", "Kahneman, 2011", "Ajzen, 1991",
        "Davis, 1989", "Lazarus & Folkman, 1984", "Schraw & Dennison, 1994",
        "Pintrich et al., 1993", "Lintner, 2024", "DeVellis",
        "Carolus et al., 2023", "Laupichler et al., 2023",
        "Sperber et al., 2010", "Efklides, 2006",
    ]

    print("\n  Citations check:")
    cit_ok = True
    for c in citations:
        if c in full_text:
            print(f"    [PASS] {c}")
        else:
            print(f"    [FAIL] {c}")
            cit_ok = False

    # Check technical terms
    terms = ["CA-Aw", "CA-Jd", "CA-Ac", "TCRS", "SRL", "MSLQ", "MAI",
             "SNAIL", "MAILS", "AILQ"]

    print("\n  Technical terms check:")
    term_ok = True
    for t in terms:
        if t in full_text:
            print(f"    [PASS] {t}")
        else:
            print(f"    [FAIL] {t}")
            term_ok = False

    # Check "It is important to note" is gone
    iitn_count = full_text.lower().count("it is important to note")
    print(f"\n  'It is important to note' remaining: {iitn_count}")
    if iitn_count == 0:
        print("    [PASS]")
    else:
        print("    [FAIL]")

    # Check "established precedent" is gone
    ep_count = full_text.lower().count("established precedent")
    print(f"  'established precedent' remaining: {ep_count}")
    if ep_count == 0:
        print("    [PASS]")
    else:
        print("    [FAIL]")

    # Check "Importantly," is gone from subscale paras
    imp_count = full_text.count("Importantly, CA-")
    print(f"  'Importantly, CA-' remaining: {imp_count}")
    if imp_count == 0:
        print("    [PASS]")
    else:
        print("    [FAIL]")

    # Check weak conclusion is gone
    weak_count = full_text.count("may prove to be one of the most important")
    print(f"  Weak conclusion remaining: {weak_count}")
    if weak_count == 0:
        print("    [PASS]")
    else:
        print("    [FAIL]")

    # Summary
    print("\n" + "=" * 60)
    print("CHANGE REPORT")
    print("=" * 60)
    print(f"Total changes logged: {len(changes_log)}")
    for change in changes_log:
        print(f"\n  [{change['fix_id']}] Paragraph {change['para_idx']}:")
        print(f"    BEFORE: {change['before']}")
        print(f"    AFTER:  {change['after']}")

    all_pass = (em_remaining == 0 and cit_ok and term_ok and
                iitn_count == 0 and ep_count == 0 and imp_count == 0 and weak_count == 0)
    print("\n" + "=" * 60)
    if all_pass:
        print("ALL CHECKS PASSED")
    else:
        print("SOME CHECKS FAILED - review output above")
    print("=" * 60)


if __name__ == "__main__":
    main()
